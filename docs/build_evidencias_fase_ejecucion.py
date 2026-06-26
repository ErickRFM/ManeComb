from __future__ import annotations

import math
import shutil
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "generated-execution-assets"
OUT_DOCX = DOCS_DIR / "evidencias-fase-ejecucion-manecomb.docx"

PAGE_WIDTH_DXA = 9360

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 102, 117)
GRAY = RGBColor(242, 244, 247)
LIGHT_BLUE = RGBColor(232, 238, 245)
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)
GREEN = RGBColor(25, 169, 116)
GOLD = RGBColor(122, 90, 0)
RED = RGBColor(155, 28, 28)

PIL_NAVY = "#0B2545"
PIL_BLUE = "#2E74B5"
PIL_DARK = "#1F4D78"
PIL_MUTED = "#5A6675"
PIL_LIGHT = "#E8EEF5"
PIL_BG = "#F7F9FC"
PIL_GREEN = "#19A974"
PIL_GOLD = "#D49100"
PIL_RED = "#B42318"


def windows_font(name: str) -> str | None:
    for candidate in [Path("C:/Windows/Fonts") / name, Path("C:/Windows/Fonts") / name.lower()]:
        if candidate.exists():
            return str(candidate)
    return None


REGULAR_FONT = windows_font("arial.ttf")
BOLD_FONT = windows_font("arialbd.ttf")


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = BOLD_FONT if bold else REGULAR_FONT
    if path:
        return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    if table.rows:
        set_repeat_table_header(table.rows[0])


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    set_run_font(run, size=9, color=MUTED)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def add_toc_field(paragraph, levels="1-3"):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "{levels}" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Indice automatico: actualizar campos en Word para recalcular paginas."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def set_update_fields_on_open(docx_path: Path):
    tmp = docx_path.parent / "_toc_patch_tmp_execution"
    if tmp.exists():
        shutil.rmtree(tmp)
    tmp.mkdir()
    with zipfile.ZipFile(docx_path, "r") as zf:
        zf.extractall(tmp)
    settings = tmp / "word" / "settings.xml"
    parser = etree.XMLParser(remove_blank_text=False)
    if settings.exists():
        tree = etree.parse(str(settings), parser)
        root = tree.getroot()
    else:
        root = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}settings")
        tree = etree.ElementTree(root)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    uf = root.find("w:updateFields", namespaces=ns)
    if uf is None:
        uf = etree.Element(f"{{{ns['w']}}}updateFields")
        root.insert(0, uf)
    uf.set(f"{{{ns['w']}}}val", "true")
    tree.write(str(settings), xml_declaration=True, encoding="UTF-8", standalone="yes")
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in tmp.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(tmp).as_posix())
    shutil.rmtree(tmp)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)

    header = section.header.paragraphs[0]
    header.text = "ManeComb | Evidencias de fase de ejecucion"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_small_note(doc: Document, label: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [PAGE_WIDTH_DXA], 120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=10, bold=True, color=DARK_BLUE)
    run = p.add_run(text)
    set_run_font(run, size=10, color=BLACK)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=9, italic=True, color=MUTED)


def add_picture(paragraph, image_path: Path, width, alt_text: str):
    inline = paragraph.add_run().add_picture(str(image_path), width=width)
    inline._inline.docPr.set("title", alt_text)
    inline._inline.docPr.set("descr", alt_text)
    return inline


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], font_size=8.7, header_fill="F2F4F7"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths, 120)
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if len(str(value)) <= 14:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=BLACK)
    doc.add_paragraph()
    return table


def wrapped_lines(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines = []
    for raw in str(text).split("\n"):
        if not raw:
            lines.append("")
            continue
        words = raw.split()
        current = ""
        for word in words:
            candidate = word if not current else f"{current} {word}"
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines


def rounded_card(draw, box, fill="#FFFFFF", outline="#C9D3DF", width=3, radius=24):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def centered_text(draw, box, text, font, fill, line_gap=8):
    x1, y1, x2, y2 = box
    lines = wrapped_lines(draw, text, font, x2 - x1 - 36)
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, height in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += height + line_gap


def arrow(draw, start, end, fill=PIL_BLUE, width=5):
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    angle = math.atan2(ey - sy, ex - sx)
    length = 18
    p1 = (ex - length * math.cos(angle - math.pi / 7), ey - length * math.sin(angle - math.pi / 7))
    p2 = (ex - length * math.cos(angle + math.pi / 7), ey - length * math.sin(angle + math.pi / 7))
    draw.polygon([end, p1, p2], fill=fill)


def save_cover(path: Path):
    img = Image.new("RGB", (1800, 900), PIL_BG)
    d = ImageDraw.Draw(img)
    title = pil_font(62, bold=True)
    sub = pil_font(34)
    small = pil_font(24)
    rounded_card(d, (85, 85, 1715, 815), fill="#FFFFFF", outline="#DDE6F0", width=4, radius=50)
    d.rounded_rectangle((160, 180, 780, 610), radius=42, fill=PIL_NAVY)
    d.rounded_rectangle((230, 250, 710, 390), radius=18, fill="#D8E9FF")
    d.rectangle((245, 430, 695, 500), fill="#FFFFFF")
    d.ellipse((290, 520, 390, 620), fill="#142640")
    d.ellipse((555, 520, 655, 620), fill="#142640")
    d.ellipse((325, 555, 355, 585), fill="#FFFFFF")
    d.ellipse((590, 555, 620, 585), fill="#FFFFFF")
    d.line((910, 590, 1600, 265), fill="#B8D5EE", width=16)
    d.line((910, 590, 1140, 470, 1330, 515, 1600, 265), fill=PIL_BLUE, width=8)
    for x, y in [(910, 590), (1140, 470), (1330, 515), (1600, 265)]:
        d.ellipse((x - 22, y - 22, x + 22, y + 22), fill="#FFFFFF", outline=PIL_BLUE, width=6)
    d.text((860, 300), "ManeComb", font=title, fill=PIL_NAVY)
    d.text((865, 382), "Fase de ejecucion", font=sub, fill=PIL_DARK)
    d.text((867, 435), "Portafolio inicial de evidencias", font=small, fill=PIL_MUTED)
    img.save(path)


def save_architecture(path: Path):
    img = Image.new("RGB", (2200, 1300), PIL_BG)
    d = ImageDraw.Draw(img)
    title = pil_font(48, bold=True)
    head = pil_font(29, bold=True)
    body = pil_font(24)
    d.text((90, 55), "Arquitectura y dependencias del proyecto", font=title, fill=PIL_NAVY)
    boxes = {
        "mobile": (90, 190, 560, 430, "App movil\nReact Native CLI\nTypeScript"),
        "web": (90, 520, 560, 760, "Portal web\nReact + Vite\nCloudflare Pages"),
        "api": (845, 300, 1375, 610, "Backend API\nNode.js + Express\nJWT, REST, Socket.IO"),
        "db": (870, 825, 1350, 1050, "MongoDB Atlas\nUsuarios, rutas,\nviajes y pagos"),
        "services": (1620, 180, 2110, 470, "Integraciones\nMercado Pago\nGoogle Maps\nFirebase"),
        "infra": (1595, 735, 2120, 1045, "Infraestructura\nRender, Docker,\nPM2, Nginx"),
    }
    for key, (x1, y1, x2, y2, text) in boxes.items():
        fill = "#FFFFFF" if key != "api" else "#E8EEF5"
        rounded_card(d, (x1, y1, x2, y2), fill=fill, outline="#B7C9DA", width=4)
        lines = text.split("\n")
        d.text((x1 + 34, y1 + 34), lines[0], font=head, fill=PIL_NAVY)
        y = y1 + 92
        for line in lines[1:]:
            d.text((x1 + 34, y), line, font=body, fill=PIL_MUTED)
            y += 40
    arrow(d, (560, 315), (845, 415))
    arrow(d, (560, 640), (845, 495))
    arrow(d, (1110, 610), (1110, 825))
    arrow(d, (1375, 420), (1620, 325))
    arrow(d, (1375, 520), (1595, 885))
    d.text((620, 330), "HTTPS / JSON", font=pil_font(22), fill=PIL_DARK)
    d.text((1135, 705), "Persistencia", font=pil_font(22), fill=PIL_DARK)
    d.text((1430, 305), "Pagos, mapas, push", font=pil_font(22), fill=PIL_DARK)
    img.save(path)


def save_sprint_cycle(path: Path):
    img = Image.new("RGB", (2200, 1100), PIL_BG)
    d = ImageDraw.Draw(img)
    title = pil_font(48, bold=True)
    head = pil_font(28, bold=True)
    small = pil_font(22)
    d.text((90, 55), "Ciclo de sprint quincenal", font=title, fill=PIL_NAVY)
    items = [
        ("Planificacion", "Backlog, alcance y responsables"),
        ("Ejecucion", "Desarrollo backend, web y movil"),
        ("Revision tecnica", "PR, calidad y configuracion"),
        ("QA funcional", "Casos, defectos y evidencias"),
        ("Cierre", "Reporte, minuta y retrospectiva"),
    ]
    centers = [(360, 520), (760, 260), (1220, 260), (1660, 520), (1040, 820)]
    for (title_text, detail), (cx, cy) in zip(items, centers):
        rounded_card(d, (cx - 210, cy - 95, cx + 210, cy + 95), fill="#FFFFFF", outline="#ADC5DA", width=4)
        centered_text(d, (cx - 190, cy - 75, cx + 190, cy - 15), title_text, head, PIL_NAVY)
        centered_text(d, (cx - 190, cy - 10, cx + 190, cy + 70), detail, small, PIL_MUTED)
    for idx in range(len(centers)):
        start = centers[idx]
        end = centers[(idx + 1) % len(centers)]
        arrow(d, (start[0] + 205 if end[0] > start[0] else start[0], start[1]), (end[0] - 205 if end[0] > start[0] else end[0], end[1]), fill=PIL_BLUE, width=5)
    d.text((850, 510), "Duracion base: 2 semanas", font=pil_font(30, bold=True), fill=PIL_DARK)
    d.text((780, 560), "Cada sprint debe cerrar con reporte de avance y evidencias.", font=small, fill=PIL_MUTED)
    img.save(path)


def save_evidence_tree(path: Path):
    img = Image.new("RGB", (1900, 1300), PIL_BG)
    d = ImageDraw.Draw(img)
    title = pil_font(46, bold=True)
    body = pil_font(26)
    mono = pil_font(25)
    d.text((85, 55), "Estructura de portafolio de evidencias", font=title, fill=PIL_NAVY)
    rounded_card(d, (90, 145, 1810, 1210), fill="#FFFFFF", outline="#DDE6F0", width=4, radius=32)
    lines = [
        "evidencias-fase-ejecucion/",
        "  01-minutas/",
        "  02-reportes-avance/",
        "  03-capturas-pantalla/",
        "  04-diagramas/",
        "  05-casos-prueba/",
        "  06-bitacoras/",
        "  07-control-cambios/",
        "  08-riesgos/",
        "  09-entregables/",
        "  10-validacion-rubrica/",
    ]
    y = 225
    for idx, line in enumerate(lines):
        color = PIL_NAVY if idx == 0 else PIL_DARK
        d.text((150, y), line, font=mono, fill=color)
        y += 78
    d.rounded_rectangle((1120, 280, 1700, 610), radius=26, fill="#E8EEF5", outline="#ADC5DA", width=3)
    d.text((1160, 330), "Regla de control", font=pil_font(31, bold=True), fill=PIL_NAVY)
    for i, line in enumerate(["Nombre consistente", "Fecha y responsable", "Version trazable", "Evidencia verificable"]):
        d.text((1160, 405 + i * 45), f"- {line}", font=body, fill=PIL_MUTED)
    img.save(path)


def create_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets = {
        "cover": ASSET_DIR / "cover-ejecucion.png",
        "architecture": ASSET_DIR / "arquitectura-ejecucion.png",
        "sprint": ASSET_DIR / "ciclo-sprint.png",
        "evidence_tree": ASSET_DIR / "portafolio-evidencias.png",
    }
    save_cover(assets["cover"])
    save_architecture(assets["architecture"])
    save_sprint_cycle(assets["sprint"])
    save_evidence_tree(assets["evidence_tree"])
    return assets


def add_cover(doc: Document, assets: dict[str, Path]):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(p, assets["cover"], Inches(6.4), "Portada ManeComb fase de ejecucion")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Preparacion inicial de evidencias de la fase de ejecucion")
    set_run_font(r, size=20, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Administracion de Proyectos de TI")
    set_run_font(r, size=13, color=DARK_BLUE)
    rows = [
        ["Proyecto", "ManeComb"],
        ["Tipo", "Documento base para seguimiento, control y evidencias"],
        ["Fecha", date.today().isoformat()],
        ["Stack base", "React + Vite, React Native CLI, Node.js + Express, MongoDB Atlas, Socket.IO"],
        ["Infraestructura", "Render, Cloudflare, Docker, PM2 y Nginx"],
    ]
    add_table(doc, ["Campo", "Valor"], rows, [2100, 7260], font_size=9.5, header_fill="E8EEF5")
    add_small_note(doc, "Alcance", "Este documento organiza la fase de ejecucion del proyecto ManeComb. No agrega codigo de aplicacion ni cambia el alcance tecnico; define evidencias, plantillas y mecanismos de control.")
    doc.add_page_break()


def add_index(doc: Document):
    add_heading(doc, "Indice", 1)
    p = doc.add_paragraph()
    add_toc_field(p, "1-3")
    doc.add_page_break()


def add_context(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "Contexto del proyecto", 1)
    add_body(doc, "ManeComb es una plataforma SaaS para la administracion operativa de flotillas de transporte tipo combi. El sistema concentra monitoreo GPS, gestion de usuarios, rutas, viajes, incidencias, comunicacion interna y control de planes de suscripcion.")
    add_body(doc, "El alcance vigente contempla un portal web React + Vite, una aplicacion movil React Native CLI con TypeScript, un backend Node.js + Express, MongoDB Atlas como base de datos, Socket.IO para comunicacion en tiempo real, autenticacion JWT e integraciones con Mercado Pago, Google Maps y Firebase.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(p, assets["architecture"], Inches(6.45), "Arquitectura ManeComb con web movil backend datos integraciones e infraestructura")
    add_caption(doc, "Figura 1. Arquitectura contextual usada para planear la fase de ejecucion.")
    rows = [
        ["Frontend web", "React + Vite", "Portal de ventas y vistas web asociadas al flujo comercial."],
        ["Aplicacion movil", "React Native CLI + TypeScript", "Operacion movil, acceso, mapa, incidencias y comunicacion."],
        ["Backend", "Node.js + Express", "API REST, reglas de negocio, autenticacion JWT y webhooks."],
        ["Datos", "MongoDB Atlas", "Persistencia de usuarios, flota, viajes, documentos y pagos."],
        ["Tiempo real", "Socket.IO", "Ubicacion viva, chat y eventos operativos."],
        ["Infraestructura", "Render, Cloudflare, Docker, PM2, Nginx", "Despliegue, publicacion y operacion."],
    ]
    add_table(doc, ["Capa", "Tecnologia", "Uso en ManeComb"], rows, [1700, 2600, 5060], font_size=8.8)
    doc.add_page_break()


def section_1(doc: Document):
    add_heading(doc, "1. Organizacion del equipo", 1)
    add_body(doc, "La organizacion se plantea por roles funcionales, ya que los nombres de integrantes deben asignarse por el equipo. Cada rol se vincula con responsabilidades y evidencias verificables.")
    rows = [
        ["Lider de proyecto", "Coordinar avance, cronograma, reuniones, acuerdos, cambios, riesgos y cumplimiento de rubrica.", "Cronograma, minutas, control de cambios, matriz de riesgos."],
        ["Desarrollador backend", "Mantener API Express, JWT, reglas de negocio, MongoDB Atlas, Socket.IO e integraciones.", "PRs backend, pruebas API, documentacion de endpoints."],
        ["Desarrollador frontend web", "Mantener portal React + Vite, rutas web, formularios, estados y despliegue en Cloudflare.", "Capturas web, build, pruebas funcionales web."],
        ["Desarrollador movil", "Mantener app React Native CLI, navegacion, sesion, mapas, permisos y pruebas en Android.", "APK/debug, capturas moviles, pruebas en dispositivo."],
        ["QA y documentacion", "Definir casos de prueba, validar funcionalidad, registrar defectos y consolidar evidencias.", "Casos de prueba, reportes QA, bitacoras, evidencias finales."],
    ]
    add_table(doc, ["Rol", "Responsabilidades", "Evidencias"], rows, [1800, 4300, 3260], font_size=8.5)
    raci = [
        ["Definir alcance de ejecucion", "A/R", "C", "C", "C", "C"],
        ["Mantener cronograma", "A/R", "I", "I", "I", "C"],
        ["Preparar minutas", "A/R", "I", "I", "I", "R"],
        ["Registrar bitacora", "A", "R", "R", "R", "R"],
        ["Validar API backend", "C", "A/R", "C", "C", "C"],
        ["Validar portal web", "C", "C", "A/R", "I", "C"],
        ["Validar app movil", "C", "C", "I", "A/R", "C"],
        ["Ejecutar QA funcional", "A", "C", "C", "C", "R"],
        ["Gestionar riesgos/cambios", "A/R", "C", "C", "C", "C"],
        ["Preparar evidencias finales", "A", "C", "C", "C", "R"],
    ]
    add_table(doc, ["Actividad", "Lider", "Backend", "Web", "Movil", "QA/Doc"], raci, [2750, 1250, 1250, 1250, 1250, 1610], font_size=8)


def section_2(doc: Document):
    add_heading(doc, "2. Plan de seguimiento", 1)
    add_body(doc, "El seguimiento se organiza con reuniones semanales y cierre de sprint cada dos semanas. Las minutas y bitacoras deben actualizarse de forma constante para conservar trazabilidad.")
    meetings = [
        ["Planeacion de sprint", "Cada 2 semanas, dia 1", "Definir alcance, backlog priorizado, responsables y criterios de aceptacion.", "Lider de proyecto"],
        ["Seguimiento semanal", "Semanal", "Revisar avance, bloqueos, riesgos y evidencias pendientes.", "Lider de proyecto"],
        ["Revision tecnica", "Semanal o por PR", "Validar cambios tecnicos, calidad y compatibilidad de integraciones.", "Backend/Web/Movil"],
        ["QA y cierre", "Ultimos 2 dias del sprint", "Ejecutar pruebas, consolidar hallazgos y preparar reporte.", "QA y documentacion"],
    ]
    add_table(doc, ["Reunion", "Frecuencia", "Objetivo", "Responsable"], meetings, [1900, 1800, 4260, 1400], font_size=8.5)
    add_heading(doc, "2.1 Formato de minuta", 2)
    minuta = [
        ["Fecha", "AAAA-MM-DD"],
        ["Asistentes", "Nombre, rol y participacion."],
        ["Objetivo", "Tema central de la reunion."],
        ["Avances", "Actividades completadas y evidencias generadas."],
        ["Bloqueos", "Impedimentos tecnicos, de gestion o de dependencias externas."],
        ["Acuerdos", "Decision, responsable y fecha compromiso."],
        ["Riesgos nuevos", "Riesgo, probabilidad, impacto y mitigacion inicial."],
        ["Proxima revision", "Fecha y temas a validar."],
    ]
    add_table(doc, ["Campo", "Contenido requerido"], minuta, [2100, 7260], font_size=9)
    add_heading(doc, "2.2 Formato de bitacora", 2)
    bitacora = [
        ["Fecha", "Responsable", "Actividad", "Modulo", "Evidencia", "Estado"],
        ["AAAA-MM-DD", "Rol/persona", "Descripcion breve", "Backend/Web/Movil/QA", "Commit, captura, caso o reporte", "Pendiente/En proceso/Cerrado"],
    ]
    add_table(doc, bitacora[0], bitacora[1:], [1200, 1500, 2600, 1450, 1810, 800], font_size=8.5)


def section_3(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "3. Gestion del avance", 1)
    add_body(doc, "La metodologia recomendada es agil con sprints quincenales. El avance debe controlarse por backlog, tablero Kanban y reporte de sprint.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(p, assets["sprint"], Inches(6.35), "Ciclo de sprint quincenal para ManeComb")
    add_caption(doc, "Figura 2. Flujo de seguimiento para sprints quincenales.")
    rows = [
        ["Herramienta recomendada", "GitHub Projects o tablero equivalente", "Permite vincular issues, PRs, responsables y estados."],
        ["Flujo de trabajo", "Backlog -> Por hacer -> En proceso -> Revision -> QA -> Cerrado", "Cada tarea debe tener criterio de aceptacion y evidencia."],
        ["Estados de tareas", "Backlog, To Do, In Progress, Code Review, QA, Blocked, Done", "Blocked requiere causa y plan de desbloqueo."],
        ["Indicadores", "% avance, tareas cerradas, defectos abiertos, cobertura de pruebas, cumplimiento de sprint", "Se reportan al cierre de cada sprint."],
    ]
    add_table(doc, ["Elemento", "Definicion", "Aplicacion"], rows, [2300, 3100, 3960], font_size=8.5)
    add_heading(doc, "3.1 Indicadores de avance", 2)
    indicators = [
        ["Cumplimiento de sprint", "Tareas cerradas / tareas comprometidas", " >= 80%"],
        ["Defectos abiertos", "Errores activos por severidad", "Tendencia descendente"],
        ["Build y pruebas", "Ejecuciones exitosas de lint, typecheck y test", "Sin fallos bloqueantes"],
        ["Evidencias completas", "Evidencias generadas / evidencias requeridas", "100% al cierre"],
    ]
    add_table(doc, ["Indicador", "Medicion", "Meta inicial"], indicators, [3000, 4200, 2160], font_size=8.8)


def section_4(doc: Document):
    add_heading(doc, "4. Gestion de calidad", 1)
    add_body(doc, "La calidad se controla mediante listas de verificacion, revision tecnica y pruebas antes del cierre de cada sprint.")
    checklist = [
        ["Cumplimiento de requisitos", "La funcionalidad corresponde al alcance aprobado y no introduce cambios no solicitados.", "Lider / QA"],
        ["Revision de codigo", "PR revisado, cambios acotados, sin secretos y sin degradar configuracion productiva.", "Equipo tecnico"],
        ["Pruebas funcionales", "Flujos principales ejecutados con resultado esperado y evidencia.", "QA"],
        ["Validacion de interfaces", "Pantallas sin bloqueos, rutas correctas, mensajes claros y estados recuperables.", "Frontend/Movil/QA"],
        ["Revision de documentacion", "Minutas, bitacoras, riesgos, cambios y resultados de pruebas actualizados.", "QA y documentacion"],
    ]
    add_table(doc, ["Punto de control", "Criterio", "Responsable"], checklist, [2400, 5260, 1700], font_size=8.8)
    add_small_note(doc, "Politica", "Ningun incremento debe cerrarse sin evidencia minima: tarea, responsable, prueba o captura, resultado y fecha.")


def section_5(doc: Document):
    add_heading(doc, "5. Gestion de riesgos", 1)
    add_body(doc, "La matriz de riesgos inicial considera amenazas tecnicas y de gestion asociadas al despliegue, integraciones, autenticacion, servicios externos y calendario.")
    risks = [
        ["R-01", "Fallos de despliegue en Render o Cloudflare", "Media", "Alto", "Validar build, variables de entorno, health checks y rollback documentado.", "Backend/Web"],
        ["R-02", "Problemas de integracion Mercado Pago, Google Maps o Firebase", "Media", "Alto", "Usar ambientes de prueba, logs seguros y contratos de API documentados.", "Backend"],
        ["R-03", "Errores de autenticacion JWT o sesion movil", "Media", "Alto", "Pruebas de login, refresh, expiracion, roles y acceso operativo.", "Backend/Movil"],
        ["R-04", "Dependencia de servicios externos", "Media", "Medio", "Definir fallback, monitoreo y criterios de degradacion controlada.", "Lider/Backend"],
        ["R-05", "Retrasos de desarrollo", "Media", "Medio", "Priorizar backlog, limitar alcance por sprint y registrar bloqueos.", "Lider"],
        ["R-06", "Datos incompletos en MongoDB Atlas", "Baja", "Alto", "Validar migraciones, respaldos y scripts de diagnostico no destructivos.", "Backend/QA"],
        ["R-07", "Diferencias entre entorno local y produccion", "Media", "Alto", "Usar variables productivas documentadas y pruebas contra Render.", "Equipo tecnico"],
    ]
    add_table(doc, ["ID", "Riesgo", "Prob.", "Impacto", "Mitigacion", "Resp."], risks, [700, 2600, 850, 900, 3210, 1100], font_size=7.9)


def section_6(doc: Document):
    add_heading(doc, "6. Gestion de cambios", 1)
    add_body(doc, "Todo cambio de alcance, tecnologia, integracion, calendario o comportamiento productivo debe registrarse antes de implementarse.")
    fields = [
        ["ID del cambio", "CHG-001"],
        ["Descripcion", "Resumen del cambio solicitado."],
        ["Motivo", "Problema, oportunidad o ajuste de alcance que justifica la solicitud."],
        ["Impacto", "Modulo afectado, riesgos, esfuerzo, pruebas y documentacion requerida."],
        ["Prioridad", "Alta, media o baja."],
        ["Responsable", "Rol/persona encargada de coordinarlo."],
        ["Estado", "Propuesto, aprobado, rechazado, en ejecucion, validado o cerrado."],
    ]
    add_table(doc, ["Campo", "Formato"], fields, [2200, 7160], font_size=9)
    add_heading(doc, "6.1 Flujo de control", 2)
    for step in [
        "Registrar solicitud de cambio con ID unico.",
        "Evaluar impacto tecnico, calendario, calidad, riesgos y evidencias.",
        "Aprobar o rechazar el cambio en reunion o revision formal.",
        "Implementar solo si fue aprobado y cuenta con responsable.",
        "Validar con pruebas y actualizar documentacion afectada.",
    ]:
        add_numbered(doc, step)


def section_7(doc: Document):
    add_heading(doc, "7. Estrategia de pruebas", 1)
    add_body(doc, "La estrategia cubre pruebas unitarias, integracion, funcionales y aceptacion. Cada prueba debe producir resultado, evidencia y responsable.")
    rows = [
        ["Unitarias", "Funciones y utilidades aisladas.", "Jest, node:test, pruebas TS cuando aplique.", "Desarrollador del modulo"],
        ["Integracion", "API, base de datos, sockets, pagos y rutas entre componentes.", "npm test backend, pruebas de endpoints, mocks controlados.", "Backend/QA"],
        ["Funcionales", "Flujos reales: login, mapa, incidencias, pagos, portal y mobile.", "Checklist QA, capturas, pruebas en Android.", "QA"],
        ["Aceptacion", "Validacion contra criterios de rubrica y casos principales de ManeComb.", "Matriz de aceptacion, evidencia firmada o aprobada.", "Lider/QA"],
    ]
    add_table(doc, ["Tipo", "Objetivo", "Herramientas sugeridas", "Responsable"], rows, [1600, 2700, 3360, 1700], font_size=8.5)
    add_heading(doc, "7.1 Evidencia minima por prueba", 2)
    for item in [
        "Identificador del caso de prueba.",
        "Modulo o requisito cubierto.",
        "Datos de entrada y pasos ejecutados.",
        "Resultado esperado y resultado obtenido.",
        "Estado: aprobado, fallido, bloqueado o no aplica.",
        "Captura, log o enlace a commit cuando corresponda.",
    ]:
        add_bullet(doc, item)


def section_8(doc: Document):
    add_heading(doc, "8. Gestion de configuracion", 1)
    add_body(doc, "La configuracion controla codigo, versiones, ramas, respaldos y despliegues para mantener trazabilidad durante la ejecucion.")
    rows = [
        ["Ramas GitHub", "main como rama estable; ramas feature/fix/docs para trabajo puntual; evitar cambios directos sin revision."],
        ["Convencion commits", "Usar prefijos tipo feat:, fix:, docs:, test:, chore: con descripcion clara y alcance."],
        ["Pull requests", "Incluir objetivo, archivos afectados, pruebas ejecutadas, evidencia y riesgos."],
        ["Respaldos", "GitHub para codigo; MongoDB Atlas con respaldos/snapshots segun plan; conservar artefactos de evidencias por sprint."],
        ["Versionado", "Version semantica o etiquetas por hito: v0.1-ejecucion, v0.2-qa, v1.0-entrega."],
    ]
    add_table(doc, ["Elemento", "Politica inicial"], rows, [2200, 7160], font_size=8.8)
    add_heading(doc, "8.1 Checklist de PR", 2)
    for item in [
        "El PR describe que cambia y por que.",
        "No incluye secretos, .env, APKs no solicitados ni credenciales.",
        "Se ejecutaron pruebas relevantes o se justifica por que no aplican.",
        "Se actualizo documentacion cuando el cambio afecta operacion o despliegue.",
        "El cambio esta acotado al objetivo aprobado.",
    ]:
        add_bullet(doc, item)


def section_9(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "9. Portafolio de evidencias", 1)
    add_body(doc, "El portafolio organiza los archivos que demuestran la ejecucion del proyecto. La estructura debe mantenerse constante para facilitar revision por rubrica.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(p, assets["evidence_tree"], Inches(5.9), "Estructura de carpetas para portafolio de evidencias")
    add_caption(doc, "Figura 3. Estructura recomendada del portafolio de evidencias.")
    rows = [
        ["01-minutas", "Minutas semanales y acuerdos."],
        ["02-reportes-avance", "Reportes por sprint e indicadores."],
        ["03-capturas-pantalla", "Evidencias visuales web, movil y backend cuando aplique."],
        ["04-diagramas", "Arquitectura, procesos, riesgos y flujos."],
        ["05-casos-prueba", "Casos, resultados, defectos y evidencias QA."],
        ["06-bitacoras", "Registro de actividades por rol y fecha."],
        ["07-control-cambios", "Solicitudes de cambio y estado."],
        ["08-riesgos", "Matriz historica de riesgos y mitigaciones."],
        ["09-entregables", "Documentos finales, builds y reportes aprobados."],
        ["10-validacion-rubrica", "Matriz de criterios y evidencias asociadas."],
    ]
    add_table(doc, ["Carpeta", "Contenido"], rows, [2500, 6860], font_size=8.8)


def section_10(doc: Document):
    add_heading(doc, "10. Validacion con la rubrica", 1)
    add_body(doc, "La siguiente matriz propone 30 criterios evaluables alineados con los entregables solicitados para la fase de ejecucion. Si la rubrica oficial usa una redaccion distinta, estos criterios deben mapearse al enunciado institucional correspondiente sin perder la evidencia asociada.")
    criteria = [
        ["C01", "Contexto del proyecto documentado", "Descripcion de ManeComb y objetivo de ejecucion", "Documento base"],
        ["C02", "Objetivos del proyecto identificados", "Objetivo general y objetivos especificos", "Documento base"],
        ["C03", "Alcance operativo definido", "Alcance vigente y restricciones", "Documento base"],
        ["C04", "Arquitectura descrita", "Diagrama y tabla de capas tecnologicas", "Diagramas"],
        ["C05", "Tecnologias justificadas", "Tabla de stack y uso", "Documento base"],
        ["C06", "Roles definidos", "Tabla de roles y responsabilidades", "Organizacion"],
        ["C07", "Responsabilidades trazables", "Matriz RACI", "Organizacion"],
        ["C08", "Calendario de reuniones", "Plan semanal y cierre de sprint", "Seguimiento"],
        ["C09", "Formato de minuta", "Plantilla de minuta", "Minutas"],
        ["C10", "Bitacora de actividades", "Formato de bitacora", "Bitacoras"],
        ["C11", "Metodologia agil definida", "Sprints quincenales", "Gestion avance"],
        ["C12", "Flujo de trabajo definido", "Estados del tablero", "Gestion avance"],
        ["C13", "Estados de tarea documentados", "Backlog, To Do, In Progress, QA, Done", "Gestion avance"],
        ["C14", "Indicadores de avance", "Tabla de metricas y metas", "Reportes"],
        ["C15", "Cierre de sprint previsto", "Reporte, QA y retrospectiva", "Reportes"],
        ["C16", "Checklist de requisitos", "Control de cumplimiento funcional", "Calidad"],
        ["C17", "Revision de codigo", "Checklist de PR y revision tecnica", "Calidad"],
        ["C18", "Pruebas funcionales", "Casos y evidencias QA", "Pruebas"],
        ["C19", "Validacion de interfaces", "Capturas web/movil y checklist", "Capturas"],
        ["C20", "Revision documental", "Lista de documentos vivos", "Documentacion"],
        ["C21", "Matriz de riesgos", "Tabla de riesgos iniciales", "Riesgos"],
        ["C22", "Riesgo de despliegue considerado", "Mitigacion Render/Cloudflare", "Riesgos"],
        ["C23", "Riesgo de integracion considerado", "Mitigacion Mercado Pago/Maps/Firebase", "Riesgos"],
        ["C24", "Riesgo de autenticacion considerado", "Pruebas JWT y sesiones", "Riesgos"],
        ["C25", "Dependencias externas evaluadas", "Plan de fallback y monitoreo", "Riesgos"],
        ["C26", "Control de cambios definido", "Formato CHG y flujo", "Cambios"],
        ["C27", "Estrategia de pruebas", "Unitarias, integracion, funcionales, aceptacion", "Pruebas"],
        ["C28", "Gestion de configuracion", "Ramas, commits, PR y versionado", "Configuracion"],
        ["C29", "Portafolio de evidencias", "Estructura de carpetas", "Portafolio"],
        ["C30", "Actualizacion continua", "Documentos vivos y responsables", "Seguimiento"],
    ]
    add_table(doc, ["ID", "Criterio", "Evidencia necesaria", "Ubicacion"], criteria, [700, 2850, 4210, 1600], font_size=7.7)
    add_small_note(doc, "Uso", "Esta matriz debe actualizarse en cada sprint con enlace o referencia al archivo real que demuestra cada criterio.")


def add_appendix(doc: Document):
    add_heading(doc, "Anexo A. Documentos vivos de la fase de ejecucion", 1)
    docs = [
        ["Cronograma de ejecucion", "Lider de proyecto", "Semanal o ante cambio de fechas"],
        ["Minutas de reunion", "Lider / QA", "Despues de cada reunion"],
        ["Bitacora de actividades", "Todo el equipo", "Cada actividad relevante"],
        ["Matriz de riesgos", "Lider", "Semanal o ante riesgo nuevo"],
        ["Control de cambios", "Lider", "Cada solicitud"],
        ["Casos de prueba", "QA", "Cada sprint"],
        ["Reporte de avance", "Lider", "Cierre de sprint"],
        ["Matriz de rubrica", "QA", "Cierre de sprint y entrega final"],
    ]
    add_table(doc, ["Documento", "Responsable", "Frecuencia"], docs, [2900, 2500, 3960], font_size=8.8)
    add_heading(doc, "Anexo B. Convencion de nombres sugerida", 1)
    for item in [
        "minuta-AAAA-MM-DD.md",
        "bitacora-sprint-XX.md",
        "reporte-avance-sprint-XX.pdf",
        "casos-prueba-sprint-XX.xlsx",
        "riesgos-sprint-XX.md",
        "control-cambios-CHG-XXX.md",
        "captura-modulo-fecha-descripcion.png",
    ]:
        add_bullet(doc, item)


def build_doc() -> Path:
    assets = create_assets()
    doc = Document()
    configure_document(doc)
    add_cover(doc, assets)
    add_index(doc)
    add_context(doc, assets)
    section_1(doc)
    doc.add_page_break()
    section_2(doc)
    doc.add_page_break()
    section_3(doc, assets)
    doc.add_page_break()
    section_4(doc)
    section_5(doc)
    doc.add_page_break()
    section_6(doc)
    section_7(doc)
    doc.add_page_break()
    section_8(doc)
    section_9(doc, assets)
    doc.add_page_break()
    section_10(doc)
    doc.add_page_break()
    add_appendix(doc)

    doc.core_properties.title = "ManeComb - Evidencias de la fase de ejecucion"
    doc.core_properties.subject = "Organizacion, seguimiento, calidad, riesgos, cambios, pruebas, configuracion y rubrica"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Documento generado para Administracion de Proyectos de TI."
    doc.save(OUT_DOCX)
    set_update_fields_on_open(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_doc())
