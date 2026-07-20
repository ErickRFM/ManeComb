from __future__ import annotations

import math
import shutil
import textwrap
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "generated-doc-assets"
OUT_DOCX = DOCS_DIR / "documentacion-tecnica-sistema-inteligente-gestion-combis-ampliada.docx"

PAGE_WIDTH_DXA = 9360

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 102, 117)
GRAY = RGBColor(242, 244, 247)
LIGHT_BLUE = RGBColor(232, 238, 245)
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)

PIL_NAVY = "#0B2545"
PIL_BLUE = "#2E74B5"
PIL_DARK = "#1F4D78"
PIL_MUTED = "#5A6675"
PIL_LIGHT = "#E8EEF5"
PIL_BG = "#F7F9FC"
PIL_GREEN = "#19A974"
PIL_GOLD = "#D49100"
PIL_RED = "#B42318"


def windows_font(name: str) -> str | None:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts") / name.lower(),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


REGULAR_FONT = windows_font("arial.ttf")
BOLD_FONT = windows_font("arialbd.ttf")
ITALIC_FONT = windows_font("ariali.ttf")


def pil_font(size: int, bold: bool = False, italic: bool = False) -> ImageFont.FreeTypeFont:
    path = BOLD_FONT if bold else ITALIC_FONT if italic else REGULAR_FONT
    if path:
        return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, options in edges.items():
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        for key, value in options.items():
            node.set(qn(f"w:{key}"), str(value))


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(i, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_paragraph_border_bottom(paragraph, color="D7DBE2", size="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, size=9, color=MUTED)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def add_toc_field(paragraph, levels="1-3"):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "{levels}" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Índice automático: actualizar campos en Word para recalcular páginas."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def set_update_fields_on_open(docx_path: Path):
    tmp = docx_path.parent / "_toc_patch_tmp"
    if tmp.exists():
        shutil.rmtree(tmp)
    tmp.mkdir()
    with zipfile.ZipFile(docx_path, "r") as zf:
        zf.extractall(tmp)
    settings = tmp / "word" / "settings.xml"
    parser = etree.XMLParser(remove_blank_text=False)
    if settings.exists():
        tree = etree.parse(str(settings), parser)
        root = tree.getroot()
    else:
        root = etree.Element(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}settings")
        tree = etree.ElementTree(root)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    uf = root.find("w:updateFields", namespaces=ns)
    if uf is None:
        uf = etree.Element(f"{{{ns['w']}}}updateFields")
        root.insert(0, uf)
    uf.set(f"{{{ns['w']}}}val", "true")
    tree.write(str(settings), xml_declaration=True, encoding="UTF-8", standalone="yes")
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in tmp.rglob("*"):
            if p.is_file():
                zf.write(p, p.relative_to(tmp).as_posix())
    shutil.rmtree(tmp)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)

    header_p = section.header.paragraphs[0]
    header_p.text = "Sistema Inteligente de Gestión de Combis | Documentación técnica"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_p.runs:
        set_run_font(run, size=8.5, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, style=None, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_small_note(doc: Document, label: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [PAGE_WIDTH_DXA], 120)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10, color=BLACK)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=9, italic=True, color=MUTED)


def add_picture(paragraph, image_path: Path, width, alt_text: str):
    picture = paragraph.add_run().add_picture(str(image_path), width=width)
    picture._inline.docPr.set("title", alt_text)
    picture._inline.docPr.set("descr", alt_text)
    return picture


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], header_fill="0B2545", font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths, 120)
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    keep_row_together(table.rows[0])
    table.rows[0].height = Pt(24)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for i, header in enumerate(headers):
        cell = header_cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        header_text_color = WHITE if header_fill in {"0B2545", "1F4D78", "2E74B5"} else NAVY
        set_run_font(r, size=font_size, bold=True, color=header_text_color)
        set_cell_border(cell, bottom={"val": "single", "sz": "10", "color": "2E74B5"})
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        keep_row_together(table.rows[-1])
        for i, value in enumerate(row):
            if row_index % 2:
                set_cell_shading(cells[i], "F7F9FC")
            set_cell_border(
                cells[i],
                bottom={"val": "single", "sz": "4", "color": "D7E2EC"},
                start={"val": "nil"},
                end={"val": "nil"},
                top={"val": "nil"},
            )
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, bold=(i == 0), color=NAVY if i == 0 else BLACK)
    doc.add_paragraph()
    return table


def wrapped_lines(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines = []
    for raw in str(text).split("\n"):
        if not raw:
            lines.append("")
            continue
        words = raw.split()
        current = ""
        for word in words:
            candidate = word if not current else f"{current} {word}"
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines


def centered_text(draw, box, text, font, fill, line_gap=8):
    x1, y1, x2, y2 = box
    lines = wrapped_lines(draw, text, font, x2 - x1 - 36)
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, h in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += h + line_gap


def arrow(draw, start, end, fill=PIL_BLUE, width=5):
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    angle = math.atan2(ey - sy, ex - sx)
    length = 18
    p1 = (ex - length * math.cos(angle - math.pi / 7), ey - length * math.sin(angle - math.pi / 7))
    p2 = (ex - length * math.cos(angle + math.pi / 7), ey - length * math.sin(angle + math.pi / 7))
    draw.polygon([end, p1, p2], fill=fill)


def rounded_card(draw, box, fill="#FFFFFF", outline="#C9D3DF", width=3, radius=24):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle((x1 + 10, y1 + 12, x2 + 10, y2 + 12), radius=radius, fill="#DCE5EF")
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def diagram_title(draw, title: str, subtitle: str | None = None):
    draw.rounded_rectangle((70, 38, 86, 126), radius=8, fill=PIL_BLUE)
    draw.text((112, 48), title, font=pil_font(46, bold=True), fill=PIL_NAVY)
    if subtitle:
        draw.text((114, 105), subtitle, font=pil_font(23), fill=PIL_MUTED)


def diagram_legend(draw, items: list[tuple[str, str]], y: int, x: int = 90):
    cursor = x
    font = pil_font(21)
    for color, label in items:
        draw.rounded_rectangle((cursor, y, cursor + 28, y + 28), radius=7, fill=color)
        draw.text((cursor + 40, y - 1), label, font=font, fill=PIL_MUTED)
        cursor += 70 + draw.textbbox((0, 0), label, font=font)[2]


def save_cover_icon(path: Path):
    img = Image.new("RGB", (1800, 900), PIL_BG)
    d = ImageDraw.Draw(img)
    title = pil_font(54, bold=True)
    small = pil_font(30)
    d.rounded_rectangle((95, 80, 1705, 820), radius=48, fill="#FFFFFF", outline="#DDE6F0", width=4)
    for i, color in enumerate(["#E8EEF5", "#DCEAF7", "#F7E8C2"]):
        d.ellipse((1020 + i * 95, 110 + i * 20, 1580 + i * 95, 670 + i * 20), fill=color)
    d.line((220, 610, 1480, 260), fill="#89B5E0", width=20)
    d.line((220, 610, 720, 410, 1030, 450, 1480, 260), fill=PIL_BLUE, width=10)
    for x, y in [(220, 610), (720, 410), (1030, 450), (1480, 260)]:
        d.ellipse((x - 24, y - 24, x + 24, y + 24), fill="#FFFFFF", outline=PIL_BLUE, width=8)
    d.rounded_rectangle((245, 300, 810, 565), radius=36, fill=PIL_NAVY)
    d.rounded_rectangle((305, 345, 750, 445), radius=16, fill="#D8E9FF")
    d.rectangle((320, 460, 735, 510), fill="#FFFFFF")
    d.ellipse((350, 520, 435, 605), fill="#182B44")
    d.ellipse((620, 520, 705, 605), fill="#182B44")
    d.ellipse((375, 545, 410, 580), fill="#FFFFFF")
    d.ellipse((645, 545, 680, 580), fill="#FFFFFF")
    d.text((910, 355), "Transporte público", font=title, fill=PIL_NAVY)
    d.text((912, 425), "GPS + API REST + tiempo real", font=small, fill=PIL_MUTED)
    img.save(path)


def save_architecture_diagram(path: Path):
    img = Image.new("RGB", (2200, 1350), PIL_BG)
    d = ImageDraw.Draw(img)
    body = pil_font(28)
    small = pil_font(22)
    diagram_title(d, "Arquitectura lógica del sistema", "Capas, protocolos e integraciones principales")
    boxes = {
        "mobile": (90, 180, 520, 460, "Flutter / App móvil\nExpo Router en implementación actual\nZustand, mapas, secure storage"),
        "api": (820, 170, 1380, 500, "Backend API REST\nNode.js + Express\nJWT, validación, rate limit"),
        "socket": (1540, 180, 2080, 460, "Socket.IO\nUbicación viva, chat,\nradio y RTC"),
        "db": (780, 685, 1220, 970, "MongoDB / MySQL\nUsuarios, rutas, viajes,\ndocumentos, pagos"),
        "maps": (140, 720, 560, 960, "Google Maps API\nGeocodificación,\nrutas y ETA"),
        "pay": (1440, 720, 1860, 960, "Servicios externos\nPagos, correo,\nSentry, Cloudinary"),
    }
    # Draw connectors first so node cards mask line endings and protect labels.
    arrow(d, (520, 320), (820, 320))
    arrow(d, (1380, 320), (1540, 320))
    arrow(d, (1040, 500), (1010, 685))
    arrow(d, (440, 720), (820, 470))
    arrow(d, (1380, 470), (1550, 720))
    for key, (x1, y1, x2, y2, text) in boxes.items():
        fill = "#FFFFFF" if key != "api" else "#E8EEF5"
        rounded_card(d, (x1, y1, x2, y2), fill=fill, outline="#B7C9DA", width=4)
        lines = text.split("\n")
        d.text((x1 + 34, y1 + 34), lines[0], font=pil_font(31, bold=True), fill=PIL_NAVY)
        yy = y1 + 88
        body_lines = wrapped_lines(d, "\n".join(lines[1:]), body, x2 - x1 - 68)
        for line in body_lines:
            d.text((x1 + 34, yy), line, font=body, fill=PIL_MUTED)
            yy += 42
    d.text((610, 275), "HTTPS / JSON", font=small, fill=PIL_DARK)
    d.text((1410, 275), "Eventos", font=small, fill=PIL_DARK)
    d.text((1020, 595), "Persistencia", font=small, fill=PIL_DARK)
    diagram_legend(d, [(PIL_NAVY, "Núcleo"), (PIL_LIGHT, "Servicio"), (PIL_BLUE, "Flujo de datos")], 1245)
    img.save(path)


def save_mer_diagram(path: Path):
    img = Image.new("RGB", (2400, 1600), PIL_BG)
    d = ImageDraw.Draw(img)
    head = pil_font(28, bold=True)
    txt = pil_font(21)
    diagram_title(d, "Modelo Entidad-Relación (MER)", "Entidades, atributos clave y cardinalidades")
    entities = {
        "Usuarios": (95, 180, 520, 420, ["PK id_usuario", "correo UNIQUE", "rol", "estado"]),
        "Conductores": (95, 690, 520, 930, ["PK id_conductor", "licencia", "telefono", "estado"]),
        "Combis": (775, 430, 1200, 690, ["PK id_combi", "placa UNIQUE", "FK id_conductor", "estado"]),
        "Rutas": (1550, 180, 1975, 420, ["PK id_ruta", "origen", "destino", "distancia"]),
        "Viajes": (1550, 700, 1975, 960, ["PK id_viaje", "FK id_combi", "FK id_ruta", "estado"]),
        "UbicacionesGPS": (780, 1015, 1210, 1275, ["PK id_gps", "FK id_combi", "latitud", "fecha_hora"]),
        "Pagos": (95, 1130, 520, 1390, ["PK id_pago", "FK id_usuario", "monto", "estado"]),
        "Participantes": (1545, 1130, 1985, 1390, ["PK id_participante", "FK id_usuario", "FK id_conversacion"]),
    }
    relation_labels = []
    def conn(a, b, label):
        ax = (a[0] + a[2]) // 2
        ay = (a[1] + a[3]) // 2
        bx = (b[0] + b[2]) // 2
        by = (b[1] + b[3]) // 2
        d.line((ax, ay, bx, by), fill=PIL_BLUE, width=5)
        mx = (ax + bx) // 2
        my = (ay + by) // 2
        relation_labels.append((mx, my, label))
    e = entities
    # Relationships form the back layer; entity cards are painted over them.
    conn(e["Conductores"][:4], e["Combis"][:4], "1:N")
    conn(e["Combis"][:4], e["Rutas"][:4], "N:1")
    conn(e["Combis"][:4], e["Viajes"][:4], "1:N")
    conn(e["Rutas"][:4], e["Viajes"][:4], "1:N")
    conn(e["Combis"][:4], e["UbicacionesGPS"][:4], "1:N")
    conn(e["Usuarios"][:4], e["Pagos"][:4], "1:N")
    conn(e["Usuarios"][:4], e["Participantes"][:4], "N:M")
    # Cardinality labels belong to the relationship layer too; cards mask any
    # label whose midpoint would otherwise fall over entity content.
    for mx, my, label in relation_labels:
        d.rounded_rectangle((mx - 64, my - 22, mx + 64, my + 22), radius=12, fill="#FFFFFF", outline="#D7E2EC", width=2)
        tw = d.textbbox((0, 0), label, font=pil_font(20, bold=True))
        d.text((mx - (tw[2] - tw[0]) / 2, my - 13), label, font=pil_font(20, bold=True), fill=PIL_DARK)
    for name, (x1, y1, x2, y2, fields) in entities.items():
        rounded_card(d, (x1, y1, x2, y2), fill="#FFFFFF", outline="#ADC5DA", width=4)
        d.rounded_rectangle((x1, y1, x2, y1 + 58), radius=18, fill=PIL_LIGHT, outline="#ADC5DA", width=0)
        d.text((x1 + 24, y1 + 16), name, font=head, fill=PIL_NAVY)
        yy = y1 + 82
        for f in fields:
            d.text((x1 + 30, yy), f, font=txt, fill=PIL_MUTED)
            yy += 38
    diagram_legend(d, [(PIL_LIGHT, "Entidad"), (PIL_BLUE, "Relación"), (PIL_NAVY, "Clave")], 1510)
    img.save(path)


def save_relational_diagram(path: Path):
    img = Image.new("RGB", (2200, 1300), PIL_BG)
    d = ImageDraw.Draw(img)
    diagram_title(d, "Modelo relacional normalizado", "Tablas, claves y dependencias de referencia")
    tables = [
        ("usuarios", "PK id_usuario\ncorreo UNIQUE\nrol, estado"),
        ("conductores", "PK id_conductor\nFK id_usuario opcional\nlicencia UNIQUE"),
        ("combis", "PK id_combi\nFK id_conductor\nFK id_ruta"),
        ("rutas", "PK id_ruta\nnombre_ruta\norigen, destino"),
        ("viajes", "PK id_viaje\nFK id_combi\nFK id_ruta"),
        ("ubicaciones_gps", "PK id_gps\nFK id_combi\nlatitud, longitud"),
        ("pagos", "PK id_pago\nFK id_usuario\nmonto, metodo_pago"),
        ("conversaciones_usuarios", "PK compuesta\nFK id_usuario\nFK id_conversacion"),
    ]
    positions = []
    for i, (name, fields) in enumerate(tables):
        col = i % 4
        row = i // 4
        x1 = 95 + col * 520
        y1 = 190 + row * 470
        x2 = x1 + 430
        y2 = y1 + 300
        positions.append((x1, y1, x2, y2))
        rounded_card(d, (x1, y1, x2, y2), fill="#FFFFFF", outline="#B9CADB", width=4)
        d.rectangle((x1, y1, x2, y1 + 60), fill=PIL_NAVY)
        d.text((x1 + 24, y1 + 16), name, font=pil_font(28, bold=True), fill="#FFFFFF")
        yy = y1 + 88
        for line in fields.split("\n"):
            d.text((x1 + 28, yy), line, font=pil_font(23), fill=PIL_MUTED)
            yy += 40
    # Relationship spine
    arrow(d, (525, 340), (615, 340))
    arrow(d, (1045, 340), (1135, 340))
    arrow(d, (1565, 340), (1655, 340))
    arrow(d, (1315, 490), (1315, 660))
    arrow(d, (1810, 490), (1810, 660))
    d.text((90, 1120), "Integridad referencial: las FK se restringen o actualizan bajo reglas de negocio; los catálogos y estados se validan en backend.", font=pil_font(26), fill=PIL_MUTED)
    img.save(path)


def save_modules_diagram(path: Path):
    img = Image.new("RGB", (2200, 1250), PIL_BG)
    d = ImageDraw.Draw(img)
    diagram_title(d, "Módulos funcionales", "Mapa de capacidades alrededor del núcleo operativo")
    center = (860, 455, 1340, 795)
    modules = [
        ("Login y autenticación", 120, 220),
        ("Usuarios y roles", 650, 190),
        ("Conductores", 1440, 220),
        ("Combis", 1600, 520),
        ("Rutas y viajes", 1440, 840),
        ("GPS en tiempo real", 650, 900),
        ("Pagos y ventas", 120, 840),
        ("Administración", 120, 520),
    ]
    # Paint spokes before cards so no connector crosses node text.
    center_point = ((center[0] + center[2]) // 2, (center[1] + center[3]) // 2)
    for _, x, y in modules:
        box = (x, y, x + 430, y + 150)
        module_point = ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2)
        arrow(d, module_point, center_point, fill="#8AB8E6", width=4)
    for label, x, y in modules:
        box = (x, y, x + 430, y + 150)
        rounded_card(d, box, fill="#FFFFFF", outline="#B9CADB", width=4)
        centered_text(d, box, label, pil_font(27, bold=True), PIL_NAVY)
    rounded_card(d, center, fill=PIL_NAVY, outline=PIL_NAVY, width=4)
    centered_text(d, center, "API REST + Socket.IO\nCentro operativo", pil_font(35, bold=True), "#FFFFFF")
    diagram_legend(d, [(PIL_NAVY, "Núcleo"), ("#FFFFFF", "Módulo"), ("#8AB8E6", "Integración")], 1170)
    img.save(path)


def save_ecosystem_diagram(path: Path):
    img = Image.new("RGB", (2200, 1250), PIL_BG)
    d = ImageDraw.Draw(img)
    diagram_title(d, "Interconexión de aplicaciones", "App móvil, plataforma central y portal de ventas")
    cards = {
        "mobile": (90, 250, 650, 800, "App móvil", ["Operación en ruta", "GPS y navegación", "Chat, radio y llamadas", "Documentos e incidencias"]),
        "backend": (820, 200, 1380, 850, "Backend central", ["API REST + JWT", "Socket.IO y señalización RTC", "Reglas multiempresa", "Persistencia y auditoría"]),
        "sales": (1550, 250, 2110, 800, "Portal de ventas", ["Planes y checkout", "Onboarding y activación", "Unidades, rutas y usuarios", "Pagos, sesiones y documentos"]),
    }
    # Connectors are the back layer; cards keep all text unobstructed.
    arrow(d, (650, 455), (820, 455), fill=PIL_BLUE, width=6)
    arrow(d, (820, 590), (650, 590), fill=PIL_GREEN, width=6)
    arrow(d, (1380, 455), (1550, 455), fill=PIL_BLUE, width=6)
    arrow(d, (1550, 590), (1380, 590), fill=PIL_GOLD, width=6)
    for key, (x1, y1, x2, y2, title, items) in cards.items():
        fill = PIL_LIGHT if key == "backend" else "#FFFFFF"
        rounded_card(d, (x1, y1, x2, y2), fill=fill, outline="#ADC5DA", width=4)
        d.rounded_rectangle((x1, y1, x2, y1 + 82), radius=22, fill=PIL_NAVY)
        d.text((x1 + 34, y1 + 24), title, font=pil_font(32, bold=True), fill="#FFFFFF")
        yy = y1 + 125
        for item in items:
            d.ellipse((x1 + 36, yy + 7, x1 + 54, yy + 25), fill=PIL_BLUE)
            d.text((x1 + 74, yy), item, font=pil_font(27), fill=PIL_NAVY)
            yy += 78
    d.text((680, 408), "REST", font=pil_font(22, bold=True), fill=PIL_DARK)
    d.text((680, 615), "Eventos", font=pil_font(22, bold=True), fill=PIL_GREEN)
    d.text((1420, 408), "REST", font=pil_font(22, bold=True), fill=PIL_DARK)
    d.text((1410, 615), "Socket", font=pil_font(22, bold=True), fill=PIL_GOLD)
    diagram_legend(d, [(PIL_BLUE, "Solicitud"), (PIL_GREEN, "Actualización operativa"), (PIL_GOLD, "Evento comercial")], 1120)
    img.save(path)


def save_flow_diagram(path: Path, title: str, steps: list[str], accent=PIL_BLUE):
    img = Image.new("RGB", (2200, 900), PIL_BG)
    d = ImageDraw.Draw(img)
    diagram_title(d, title, "Secuencia operativa y puntos de responsabilidad")
    n = len(steps)
    cols = min(4, n)
    box_w = 450
    box_h = 150
    gap_x = 60
    gap_y = 90
    start_x = 120
    start_y = 220
    positions = []
    for i, step in enumerate(steps):
        row = i // cols
        col = i % cols
        x1 = start_x + col * (box_w + gap_x)
        y1 = start_y + row * (box_h + gap_y)
        box = (x1, y1, x1 + box_w, y1 + box_h)
        positions.append(box)
        rounded_card(d, box, fill="#FFFFFF", outline="#B9CADB", width=4)
        d.ellipse((x1 + 18, y1 + 48, x1 + 72, y1 + 102), fill=accent)
        centered_text(d, (x1 + 18, y1 + 48, x1 + 72, y1 + 102), str(i + 1), pil_font(24, bold=True), "#FFFFFF")
        centered_text(d, (x1 + 88, y1 + 18, x1 + box_w - 20, y1 + box_h - 18), step, pil_font(23, bold=True), PIL_NAVY)
    for i in range(len(positions) - 1):
        a = positions[i]
        b = positions[i + 1]
        if (i + 1) % cols == 0:
            ax = (a[0] + a[2]) // 2
            bx = (b[0] + b[2]) // 2
            mid_y = a[3] + gap_y // 2
            d.line(((ax, a[3]), (ax, mid_y), (bx, mid_y)), fill=accent, width=5)
            arrow(d, (bx, mid_y), (bx, b[1] - 12), fill=accent, width=5)
        else:
            arrow(d, (a[2] + 10, (a[1] + a[3]) // 2), (b[0] - 12, (b[1] + b[3]) // 2), fill=accent, width=5)
    diagram_legend(d, [(accent, "Paso"), (PIL_NAVY, "Secuencia")], 820)
    img.save(path)


def create_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "cover": ASSET_DIR / "cover-transporte-tech.png",
        "modules": ASSET_DIR / "modulos-funcionales.png",
        "mer": ASSET_DIR / "modelo-entidad-relacion.png",
        "relational": ASSET_DIR / "modelo-relacional.png",
        "architecture": ASSET_DIR / "arquitectura-sistema.png",
        "ecosystem": ASSET_DIR / "interconexion-aplicaciones.png",
    }
    save_cover_icon(paths["cover"])
    save_modules_diagram(paths["modules"])
    save_mer_diagram(paths["mer"])
    save_relational_diagram(paths["relational"])
    save_architecture_diagram(paths["architecture"])
    save_ecosystem_diagram(paths["ecosystem"])

    flows = {
        "login": ("Proceso: inicio de sesión", ["Captura credenciales", "Valida formato", "Autentica en backend", "Genera JWT y refresh token", "Carga dashboard", "Abre canal Socket.IO"], PIL_BLUE),
        "registro": ("Proceso: registro de usuarios", ["Captura datos", "Valida contraseña", "Verifica correo único", "Crea usuario", "Registra auditoría", "Inicia sesión"], PIL_GREEN),
        "asignacion": ("Proceso: asignación de rutas", ["Selecciona combi", "Define origen y destino", "Calcula ruta", "Valida rol administrador", "Asigna recorrido", "Notifica mapa operativo"], PIL_GOLD),
        "gps": ("Proceso: monitoreo GPS", ["Obtiene ubicación", "Envía coordenadas", "Actualiza unidad", "Emite location:updated", "Refresca mapa", "Genera alertas"], PIL_BLUE),
        "viajes": ("Proceso: registro de viajes", ["Inicia recorrido", "Captura salida", "Registra llegada", "Calcula duración", "Persiste bitácora", "Consulta historial"], PIL_GREEN),
        "pagos": ("Proceso: pago del usuario", ["Selecciona plan", "Captura facturación", "Crea orden", "Procesa pago", "Confirma estado", "Activa onboarding"], PIL_GOLD),
        "admin": ("Proceso: administración del sistema", ["Consulta tablero", "Gestiona usuarios", "Revisa documentos", "Atiende incidencias", "Audita eventos", "Cierra acciones"], PIL_RED),
    }
    for key, (title, steps, color) in flows.items():
        paths[key] = ASSET_DIR / f"flujo-{key}.png"
        save_flow_diagram(paths[key], title, steps, color)
    return paths


DATA_DICTIONARY = {
    "Usuarios": {
        "description": "Entidad base para personas que acceden al sistema. En la implementación backend corresponde principalmente a la colección users; los conductores también pueden representarse como usuarios con rol driver.",
        "rows": [
            ["id_usuario", "VARCHAR / ObjectId", "36", "PK", "Identificador único del usuario.", "Obligatorio; no nulo."],
            ["nombre", "VARCHAR", "80", "", "Nombre del usuario o nombre completo operativo.", "Obligatorio; trim; caracteres válidos."],
            ["apellido", "VARCHAR", "80", "", "Apellido del usuario cuando el registro se separa en nombre y apellido.", "Opcional si se usa nombre completo."],
            ["correo", "VARCHAR", "120", "UNIQUE", "Correo usado para login y notificaciones.", "Obligatorio; formato email; único."],
            ["contraseña", "VARCHAR", "255", "", "Se almacena como hash, nunca en texto plano.", "Mínimo 8; letras, números y símbolo; bcrypt."],
            ["telefono", "VARCHAR", "20", "", "Teléfono de contacto operativo.", "Opcional; formato nacional/internacional."],
            ["rol", "ENUM", "20", "", "Perfil de acceso: admin, supervisor, driver, owner, dispatcher, viewer.", "Valor permitido por catálogo."],
            ["fecha_registro", "DATETIME", "-", "", "Fecha de creación de la cuenta.", "Default NOW()."],
            ["estado", "ENUM", "20", "", "Estado de cuenta: active, pending, suspended.", "Default active."],
        ],
    },
    "Conductores": {
        "description": "Especialización operativa del usuario que conduce una unidad. Puede modelarse como tabla propia o como vista de Usuarios filtrada por rol driver con campos extendidos.",
        "rows": [
            ["id_conductor", "VARCHAR / ObjectId", "36", "PK", "Identificador único del conductor.", "Obligatorio; no nulo."],
            ["nombre", "VARCHAR", "120", "", "Nombre completo del conductor.", "Obligatorio."],
            ["licencia", "VARCHAR", "40", "UNIQUE", "Número de licencia o permiso de conducción.", "Obligatorio; único; vigente."],
            ["telefono", "VARCHAR", "20", "", "Teléfono del conductor.", "Opcional."],
            ["direccion", "VARCHAR", "180", "", "Domicilio o referencia administrativa.", "Opcional; protegido como dato personal."],
            ["estado", "ENUM", "20", "", "Disponible, en ruta, suspendido, baja.", "Valor de catálogo."],
            ["fecha_ingreso", "DATE", "-", "", "Fecha de alta laboral u operativa.", "No mayor a fecha actual."],
        ],
    },
    "Combis": {
        "description": "Unidad física de transporte. En el backend actual corresponde a vehicles, con código operativo, placa, capacidad, ubicación y ruta asignada.",
        "rows": [
            ["id_combi", "VARCHAR / ObjectId", "36", "PK", "Identificador único de la unidad.", "Obligatorio."],
            ["placa", "VARCHAR", "20", "UNIQUE", "Placa vehicular oficial.", "Obligatorio; único."],
            ["modelo", "VARCHAR", "60", "", "Modelo, marca o año de la unidad.", "Opcional."],
            ["capacidad", "INT", "-", "", "Capacidad máxima de pasajeros.", "Mayor a 0."],
            ["estado", "ENUM", "25", "", "Estado operativo: en ruta, mantenimiento, disponible, fuera de servicio.", "Valor de catálogo."],
            ["id_conductor", "VARCHAR / ObjectId", "36", "FK", "Conductor asignado a la unidad.", "Referencia a Conductores.id_conductor; puede ser nulo."],
        ],
    },
    "Rutas": {
        "description": "Trayecto autorizado o planeado con origen, destino y características de recorrido. El repo usa routes y assignedRoute para rutas fijas y rutas calculadas.",
        "rows": [
            ["id_ruta", "VARCHAR / ObjectId", "36", "PK", "Identificador único de la ruta.", "Obligatorio."],
            ["nombre_ruta", "VARCHAR", "100", "", "Nombre comercial u operativo de la ruta.", "Obligatorio."],
            ["origen", "VARCHAR", "120", "", "Punto inicial o base de salida.", "Obligatorio."],
            ["destino", "VARCHAR", "120", "", "Punto final del trayecto.", "Obligatorio."],
            ["distancia", "DECIMAL", "10,2", "", "Distancia estimada en kilómetros.", "Mayor o igual a 0."],
            ["tiempo_estimado", "INT", "-", "", "Duración estimada en minutos.", "Mayor a 0."],
        ],
    },
    "Viajes": {
        "description": "Bitácora de recorridos realizados por una combi en una ruta. En el backend actual se relaciona con trip_logs.",
        "rows": [
            ["id_viaje", "VARCHAR / ObjectId", "36", "PK", "Identificador único del viaje.", "Obligatorio."],
            ["id_combi", "VARCHAR / ObjectId", "36", "FK", "Unidad que realizó el viaje.", "Referencia a Combis.id_combi."],
            ["id_ruta", "VARCHAR / ObjectId", "36", "FK", "Ruta recorrida o planeada.", "Referencia a Rutas.id_ruta."],
            ["hora_salida", "DATETIME", "-", "", "Fecha y hora de inicio.", "Obligatorio; menor que hora_llegada."],
            ["hora_llegada", "DATETIME", "-", "", "Fecha y hora de cierre.", "Obligatorio al finalizar."],
            ["estado", "ENUM", "20", "", "Programado, activo, completado, retrasado, cancelado.", "Valor de catálogo."],
        ],
    },
    "UbicacionesGPS": {
        "description": "Registro histórico o instantáneo de coordenadas enviadas por las unidades. Sirve para mapa vivo, auditoría y análisis de puntualidad.",
        "rows": [
            ["id_gps", "VARCHAR / ObjectId", "36", "PK", "Identificador del registro GPS.", "Obligatorio."],
            ["id_combi", "VARCHAR / ObjectId", "36", "FK", "Unidad que reporta la ubicación.", "Referencia a Combis.id_combi."],
            ["latitud", "DECIMAL", "10,7", "", "Latitud geográfica.", "Rango -90 a 90."],
            ["longitud", "DECIMAL", "10,7", "", "Longitud geográfica.", "Rango -180 a 180."],
            ["velocidad", "DECIMAL", "6,2", "", "Velocidad reportada en km/h.", "Mayor o igual a 0."],
            ["fecha_hora", "DATETIME", "-", "", "Momento exacto de captura.", "Default NOW(); indexado."],
        ],
    },
    "Pagos": {
        "description": "Registro económico de compras, mensualidades o pagos de usuario/empresa. En el repo actual se refleja en commercial_leads/orders y perfiles de pago.",
        "rows": [
            ["id_pago", "VARCHAR / ObjectId", "36", "PK", "Identificador único del pago.", "Obligatorio."],
            ["id_usuario", "VARCHAR / ObjectId", "36", "FK", "Usuario o propietario que realiza el pago.", "Referencia a Usuarios.id_usuario."],
            ["monto", "DECIMAL", "12,2", "", "Importe del pago.", "Mayor a 0."],
            ["metodo_pago", "ENUM", "20", "", "Tarjeta, SPEI, transferencia u otro proveedor.", "Valor permitido."],
            ["fecha_pago", "DATETIME", "-", "", "Fecha de creación o aprobación del pago.", "Obligatorio."],
            ["estado", "ENUM", "25", "", "Pendiente, aprobado, rechazado, cancelado, reembolsado.", "Valor de catálogo."],
        ],
    },
}


RELATIONAL_ROWS = [
    ["Usuarios", "id_usuario", "-", "correo UNIQUE; rol y estado por catálogo; password_hash no expone contraseña."],
    ["Conductores", "id_conductor", "id_usuario -> Usuarios.id_usuario", "La licencia debe ser única y vigente."],
    ["Combis", "id_combi", "id_conductor -> Conductores.id_conductor; id_ruta -> Rutas.id_ruta", "Una combi puede estar sin conductor temporalmente."],
    ["Rutas", "id_ruta", "-", "Código/nombre de ruta recomendado como índice secundario."],
    ["Viajes", "id_viaje", "id_combi -> Combis.id_combi; id_ruta -> Rutas.id_ruta", "hora_llegada debe ser mayor a hora_salida."],
    ["UbicacionesGPS", "id_gps", "id_combi -> Combis.id_combi", "Índice por id_combi y fecha_hora descendente."],
    ["Pagos", "id_pago", "id_usuario -> Usuarios.id_usuario", "monto positivo; estado controlado por proveedor/pasarela."],
    ["ConversacionesUsuarios", "id_conversacion + id_usuario", "id_usuario -> Usuarios.id_usuario", "Resuelve relación N:M de chat/radio."],
]


PROCESS_EXPLANATIONS = [
    ("Inicio de sesión", "El usuario captura correo y contraseña. El backend valida existencia, compara hash con bcrypt, crea sesión, firma JWT con expiración y devuelve dashboard inicial. La app guarda el token de forma segura y usa refresh token para recuperación controlada."),
    ("Registro de usuarios", "El registro valida campos mínimos, fuerza política de contraseña, revisa correo único y crea el perfil con rol inicial. Las cuentas de empresa pueden iniciar flujo comercial y las cuentas operativas quedan listas para asignación por administración."),
    ("Asignación de rutas", "Un administrador selecciona unidad, origen y destino. El backend calcula una ruta, guarda assignedRoute en la combi y emite un evento en tiempo real para que mapa y dashboard reflejen la nueva instrucción."),
    ("Monitoreo GPS", "La app del conductor o unidad envía coordenadas con velocidad. La API valida token y pertenencia de la unidad, actualiza la ubicación, emite location:updated y alimenta alertas por retraso, inactividad o desvío."),
    ("Registro de viajes", "Cada recorrido registra salida, llegada, duración, distancia y usuario responsable. La bitácora permite consultar historial por unidad y fecha de servicio, además de calcular puntualidad y vueltas por turno."),
    ("Pago del usuario", "El comprador selecciona plan, envía datos de empresa y facturación, se crea una orden comercial y se confirma el pago por proveedor o transferencia. Al aprobarse, se habilita onboarding y configuración inicial de flotilla."),
    ("Administración del sistema", "Administradores y supervisores consultan métricas, gestionan usuarios, atienden incidencias, revisan documentos, observan eventos técnicos y ajustan operación de acuerdo con permisos y auditoría."),
]


REFERENCES = [
    "Express.js. (2026). Express routing guide. https://expressjs.com/en/guide/routing.html",
    "Flutter. (2026). Flutter documentation. https://docs.flutter.dev/",
    "Google for Developers. (2026). Routes API: Compute routes. https://developers.google.com/maps/documentation/routes/compute_route_directions",
    "Internet Engineering Task Force. (2015). RFC 7519: JSON Web Token (JWT). https://datatracker.ietf.org/doc/rfc7519/",
    "MongoDB. (2026). Data modeling in MongoDB. https://www.mongodb.com/docs/manual/data-modeling/",
    "OWASP Foundation. (2026). OWASP API Security Project. https://owasp.org/API-Security/",
    "Socket.IO. (2026). Socket.IO documentation. https://socket.io/",
]


def add_cover(doc: Document, assets: dict[str, Path]):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("Sistema Inteligente de Gestión de Combis")
    set_run_font(run, size=28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("Documentación técnica: aplicación móvil, backend API REST y base de datos")
    set_run_font(run, size=14, color=MUTED)

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(pic_p, assets["cover"], Inches(5.95), "Icono tecnológico de combi conectada con ruta GPS")

    doc.add_paragraph()
    meta = [
        ("Materia", "Ingeniería de Software"),
        ("Alumno", "[Nombre del alumno]"),
        ("Universidad", "[Nombre de la universidad]"),
        ("Fecha", "26 de mayo de 2026"),
        ("Tipo de proyecto", "Aplicación móvil + Backend API REST"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [2300, 5100], 900)
    set_repeat_table_header(table.rows[0])
    for i, (label, value) in enumerate(meta):
        set_cell_shading(table.cell(i, 0), "E8EEF5")
        p = table.cell(i, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=10, bold=True, color=NAVY)
        p = table.cell(i, 1).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=10, color=BLACK)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Versión académica profesional | Preparada para entrega y revisión técnica")
    set_run_font(run, size=10, italic=True, color=MUTED)
    doc.add_page_break()


def add_toc(doc: Document):
    add_heading(doc, "Índice automático", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    add_toc_field(p)
    add_small_note(
        doc,
        "Uso en Word",
        "El índice se insertó como campo TOC. Al abrir el archivo en Microsoft Word, usar Ctrl+A y F9 para actualizar numeración de páginas si el visor no lo hace automáticamente.",
    )
    doc.add_page_break()


def add_intro(doc: Document):
    add_heading(doc, "1. Introducción", 1)
    add_body(doc, "El Sistema Inteligente de Gestión de Combis es una plataforma orientada a digitalizar y supervisar la operación diaria de rutas de transporte público. La solución integra una aplicación móvil para usuarios operativos y conductores, una API REST para la lógica central, una base de datos persistente y canales de comunicación en tiempo real para mapa, chat, radio, incidencias y seguimiento de viajes.")
    add_body(doc, "La problemática principal que resuelve es la fragmentación operativa: llamadas, hojas de cálculo, mensajes aislados y documentación física dificultan conocer el estado real de unidades, conductores, rutas, pagos y eventos de seguridad. El sistema concentra datos críticos en un flujo digital auditable y accesible según rol.")
    add_body(doc, "El objetivo general es permitir que administradores, supervisores, conductores y propietarios de empresas de transporte consulten información confiable, tomen decisiones rápidas y mantengan trazabilidad sobre viajes, ubicación, documentos, pagos, alertas y comunicaciones.")
    add_heading(doc, "1.1 Beneficios esperados", 2)
    for item in [
        "Mayor visibilidad de unidades y rutas mediante ubicación GPS y actualizaciones en vivo.",
        "Reducción de tiempos de respuesta ante incidencias, retrasos, mantenimiento y documentos vencidos.",
        "Mejor control de usuarios, roles, permisos, sesiones y acceso a módulos sensibles.",
        "Registro histórico para análisis de viajes, puntualidad, uso de flotilla y operación por turno.",
        "Base escalable para comercialización, pagos, onboarding de clientes y administración multiempresa.",
    ]:
        add_bullet(doc, item)
    add_small_note(doc, "Nota técnica", "El prompt académico solicita Flutter como frontend móvil. El repositorio revisado contiene una implementación Expo/React Native con módulos equivalentes; el documento conserva Flutter como tecnología objetivo y anota la equivalencia cuando corresponde.")
    doc.add_page_break()


def add_general_description(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "2. Descripción general del sistema", 1)
    add_body(doc, "La plataforma se organiza por módulos funcionales conectados a una API central. Cada módulo expone operaciones REST y, cuando el caso lo requiere, eventos Socket.IO para actualización inmediata en clientes conectados.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(p, assets["modules"], Inches(6.25), "Diagrama de módulos funcionales del sistema de combis")
    add_caption(doc, "Figura 1. Vista general de módulos funcionales.")
    module_rows = [
        ["Login y autenticación", "Inicio de sesión, registro, refresh token, cierre de sesión y recuperación de sesión.", "JWT, bcrypt, sesiones, rate limit."],
        ["Usuarios", "Alta, edición, suspensión, roles, perfil y datos de empresa.", "Administración y portal de cliente."],
        ["Conductores", "Operadores asignables a unidades, documentos, estado de turno y contacto.", "Rol driver o entidad especializada."],
        ["Combis", "Unidades físicas con placa, capacidad, estado, ubicación, aforo y ruta asignada.", "Colección vehicles."],
        ["Rutas", "Rutas fijas y rutas calculadas por origen/destino.", "Google Maps API o proveedor interno."],
        ["GPS", "Actualización y visualización en vivo de coordenadas, velocidad y estado.", "REST + Socket.IO."],
        ["Pagos", "Planes comerciales, checkout, confirmación y activación.", "Mercado Pago/manual; órdenes comerciales."],
        ["Administración", "Usuarios, documentos, observabilidad, incidencias, RTC y ventas.", "Permisos por rol y auditoría."],
        ["Historial de viajes", "Bitácora por unidad, fecha de servicio, duración y distancia.", "trip_logs / Viajes."],
    ]
    add_table(doc, ["Módulo", "Responsabilidad", "Implementación"], module_rows, [2100, 4400, 2860], font_size=8.7)
    add_heading(doc, "2.1 Flujo operativo general", 2)
    for step in [
        "La app valida si existe una sesión local y, si hay token, consulta el perfil y el dashboard.",
        "El backend autentica con JWT, recupera datos del store y normaliza respuestas por rol.",
        "La app consume endpoints REST para dashboard, mapa, usuarios, documentos, pagos, navegación e incidencias.",
        "Socket.IO mantiene presencia y emite eventos de ubicación, chat, radio, RTC e incidencias.",
        "La base de datos conserva usuarios, unidades, rutas, viajes, documentos, sesiones, eventos y órdenes comerciales.",
    ]:
        add_numbered(doc, step)
    doc.add_page_break()


def add_data_dictionary(doc: Document):
    add_heading(doc, "3. Diccionario de datos", 1)
    add_body(doc, "El diccionario de datos describe las entidades mínimas solicitadas para el sistema. Cuando el repositorio actual usa nombres técnicos distintos, se indica la equivalencia conceptual para mantener trazabilidad entre el modelo académico y la implementación.")
    add_table(
        doc,
        ["Convención", "Significado"],
        [
            ["PK", "Clave primaria que identifica de forma única cada registro."],
            ["FK", "Clave foránea que enlaza el registro con otra tabla."],
            ["UNIQUE", "Restricción de unicidad para evitar duplicidad lógica."],
            ["ENUM", "Campo restringido a valores permitidos por catálogo o validación de backend."],
        ],
        [2100, 7260],
        font_size=9,
    )
    dictionary_names = list(DATA_DICTIONARY)
    for index, name in enumerate(dictionary_names):
        data = DATA_DICTIONARY[name]
        add_heading(doc, f"3.{list(DATA_DICTIONARY).index(name) + 1} Tabla: {name}", 2)
        add_body(doc, data["description"])
        add_table(
            doc,
            ["Campo", "Tipo", "Long.", "PK/FK", "Descripción", "Restricciones"],
            data["rows"],
            [1450, 1250, 700, 700, 3650, 1610],
            header_fill="E8EEF5",
            font_size=8,
        )
        if index < len(dictionary_names) - 1:
            doc.add_page_break()
    doc.add_page_break()


def add_mer(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "4. Modelo Entidad-Relación (MER)", 1)
    add_body(doc, "El MER identifica las entidades principales del dominio de transporte, sus atributos esenciales y las cardinalidades que gobiernan la operación. El modelo está pensado para una base relacional o document-oriented con referencias, manteniendo integridad lógica desde la capa de servicio.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(p, assets["mer"], Inches(6.45), "Modelo entidad relación con usuarios conductores combis rutas viajes GPS pagos y participantes")
    add_caption(doc, "Figura 2. Modelo Entidad-Relación con entidades y cardinalidades.")
    add_heading(doc, "4.1 Relaciones principales", 2)
    relationship_rows = [
        ["Conductores - Combis", "1:N", "Un conductor puede operar varias unidades en distintos turnos históricos; una combi activa debe tener máximo un conductor asignado."],
        ["Rutas - Combis", "1:N", "Una ruta puede tener muchas combis; cada combi trabaja una ruta principal o una ruta asignada temporal."],
        ["Combis - Viajes", "1:N", "Cada combi registra múltiples viajes por fecha de servicio."],
        ["Rutas - Viajes", "1:N", "Una ruta acumula viajes realizados por diferentes combis."],
        ["Combis - UbicacionesGPS", "1:N", "Cada unidad reporta muchas ubicaciones durante su operación."],
        ["Usuarios - Pagos", "1:N", "Un usuario o propietario puede generar múltiples pagos u órdenes comerciales."],
        ["Usuarios - Conversaciones", "N:M", "Los usuarios participan en múltiples conversaciones; una conversación tiene múltiples usuarios."],
    ]
    add_table(doc, ["Relación", "Cardinalidad", "Justificación"], relationship_rows, [2200, 1300, 5860], font_size=8.8)
    add_body(doc, "Las relaciones N:M se resuelven mediante tablas intermedias. En chat/radio, la relación entre usuarios y conversaciones se representa con participantes. Para rutas asignadas dinámicamente, el backend puede guardar la asignación activa dentro de la unidad y conservar bitácora en Viajes.")
    doc.add_page_break()


def add_relational_model(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "5. Modelo relacional", 1)
    add_body(doc, "La transformación del MER al modelo relacional define tablas, claves primarias, claves foráneas y restricciones mínimas. Aunque MongoDB permite documentos embebidos, el modelo relacional facilita validación académica, diseño de índices y migración futura a MySQL si se requiere.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(p, assets["relational"], Inches(6.25), "Modelo relacional normalizado con tablas y claves principales")
    add_caption(doc, "Figura 3. Transformación del MER a tablas relacionales.")
    add_table(doc, ["Tabla", "PK", "FK", "Reglas de integridad"], RELATIONAL_ROWS, [1900, 1400, 3000, 3060], font_size=8.5)
    add_heading(doc, "5.1 Esquema relacional propuesto", 2)
    schema_lines = [
        "Usuarios(id_usuario PK, nombre, apellido, correo UNIQUE, password_hash, telefono, rol, fecha_registro, estado)",
        "Conductores(id_conductor PK, id_usuario FK NULL, nombre, licencia UNIQUE, telefono, direccion, estado, fecha_ingreso)",
        "Combis(id_combi PK, placa UNIQUE, modelo, capacidad, estado, id_conductor FK, id_ruta FK)",
        "Rutas(id_ruta PK, nombre_ruta, origen, destino, distancia, tiempo_estimado)",
        "Viajes(id_viaje PK, id_combi FK, id_ruta FK, hora_salida, hora_llegada, estado)",
        "UbicacionesGPS(id_gps PK, id_combi FK, latitud, longitud, velocidad, fecha_hora)",
        "Pagos(id_pago PK, id_usuario FK, monto, metodo_pago, fecha_pago, estado)",
    ]
    for line in schema_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line)
        set_run_font(r, size=9.2, name="Consolas", color=BLACK)
    doc.add_page_break()


def add_process_model(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "6. Modelo de procesos", 1)
    add_body(doc, "El modelo de procesos representa los flujos principales del sistema desde una perspectiva BPMN simple: evento inicial, tareas, validaciones, persistencia y notificación. Los diagramas son intencionalmente compactos para documentar responsabilidades sin entrar en detalles de implementación interna.")
    flow_keys = ["login", "registro", "asignacion", "gps", "viajes", "pagos", "admin"]
    for idx, ((title, explanation), key) in enumerate(zip(PROCESS_EXPLANATIONS, flow_keys), start=1):
        add_heading(doc, f"6.{idx} {title}", 2)
        add_body(doc, explanation)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture(p, assets[key], Inches(6.2), f"Diagrama de flujo para {title.lower()}")
        add_caption(doc, f"Figura {idx + 3}. Diagrama de flujo: {title.lower()}.")
        if idx < len(flow_keys):
            doc.add_page_break()
    doc.add_page_break()


def add_architecture(doc: Document, assets: dict[str, Path]):
    add_heading(doc, "7. Arquitectura del sistema", 1)
    add_body(doc, "La arquitectura sigue un patrón cliente-servidor con API REST y canal de eventos en tiempo real. La aplicación móvil concentra la experiencia del usuario, el backend centraliza reglas de negocio, la base de datos conserva el estado operativo y las integraciones externas enriquecen rutas, pagos, archivos y observabilidad.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(p, assets["architecture"], Inches(6.45), "Arquitectura lógica de frontend backend base de datos sockets mapas y servicios externos")
    add_caption(doc, "Figura 11. Arquitectura lógica de frontend, backend, persistencia e integraciones.")
    rows = [
        ["Frontend móvil", "Flutter como tecnología objetivo; Expo/React Native en implementación revisada.", "UI, navegación, estado local, mapas, cámara/documentos, almacenamiento seguro."],
        ["Backend API", "Node.js + Express.", "Autenticación, reglas de negocio, validación, módulos REST, errores, trazabilidad."],
        ["Persistencia", "MongoDB y diseño relacional compatible con MySQL.", "Usuarios, vehículos, rutas, viajes, documentos, sesiones, eventos y pagos."],
        ["Tiempo real", "Socket.IO.", "Ubicación viva, chat, radio, RTC, incidencias y presencia."],
        ["Mapas", "Google Maps API / proveedor interno.", "Búsqueda de lugares, cálculo de rutas, polilíneas, distancia y ETA."],
        ["Servicios externos", "Mercado Pago, Cloudinary, Sentry, correo/WhatsApp.", "Checkout, almacenamiento, monitoreo y comunicación comercial."],
    ]
    add_table(doc, ["Capa", "Tecnología", "Responsabilidad"], rows, [1800, 2800, 4760], font_size=8.8)
    add_heading(doc, "7.1 Endpoints principales", 2)
    endpoint_rows = [
        ["Auth", "POST /api/auth/login; POST /api/auth/register; GET /api/auth/session", "Sesión, registro y perfil autenticado."],
        ["Dashboard", "GET /api/dashboard/overview", "Resumen operativo por rol."],
        ["Ubicación", "GET /api/locations/live; POST /api/locations/update", "Mapa vivo y actualización de coordenadas."],
        ["Navegación", "GET /api/navigation/search; POST /api/navigation/plan; POST /api/navigation/assign; GET/POST /api/navigation/trips", "Búsqueda, planeación, asignación e historial."],
        ["Incidencias", "GET/POST /api/incidents; PATCH /api/incidents/:id/status", "Alta y seguimiento de eventos operativos."],
        ["Chat/documentos", "GET/POST /api/chat; GET/POST/PATCH /api/documents", "Comunicación y control documental."],
        ["Comercial", "GET /api/commercial/plans; POST /api/commercial/checkout; POST /api/commercial/confirm", "Planes, compras y confirmación de pagos."],
    ]
    add_table(doc, ["Módulo", "Endpoint", "Uso"], endpoint_rows, [1600, 4500, 3260], font_size=8.3)
    doc.add_page_break()

    add_heading(doc, "7.2 Aplicación móvil operativa", 2)
    add_body(doc, "La aplicación móvil es la herramienta de trabajo del personal en campo. Consume la misma identidad y los mismos datos organizacionales que el portal, pero prioriza continuidad operativa, navegación, ubicación y comunicación. El cliente centraliza el acceso HTTP, adjunta el token de sesión, intenta renovar credenciales cuando corresponde y mantiene separadas las acciones del usuario de la representación visual de cada pantalla.")
    for item in [
        "Autenticación y sesión: inicio de sesión, registro, renovación de token, recuperación de contraseña y almacenamiento seguro de credenciales.",
        "Operación de unidades: consulta de la proyección operacional, conductor, ruta, estado, coordenadas, velocidad, ETA e incidencias asociadas.",
        "Navegación: búsqueda y geocodificación, cálculo de ruta, paradas, asignación a unidad, inicio o pausa de sesión y bitácora de viajes.",
        "Comunicación: conversaciones directas o generales, mensajes de texto, audio y medios, radio operativa, notificaciones y sesiones RTC.",
        "Trabajo en segundo plano: envío de ubicación y protección de llamadas activas mediante un servicio foreground nativo en Android.",
        "Seguridad: control de acceso por rol, aislamiento por organización y soporte de sobres cifrados para conversaciones directas.",
    ]:
        add_bullet(doc, item)
    add_small_note(doc, "Fuente de verdad", "La proyección /operational-units concentra estado, GPS, ruta, conductor y ETA. Las pantallas no deben reconstruir esos datos a partir de fuentes paralelas.")

    add_heading(doc, "7.3 Portal web de ventas y administración", 2)
    add_body(doc, "El portal de ventas no es únicamente una página de cobro. Funciona como consola de cuenta para propietarios, administradores y responsables de facturación. El frontend web comparte el backend con la app móvil y organiza la experiencia en pantallas de operaciones, unidades, rutas, usuarios, documentos, incidencias, onboarding, plan, pagos, facturación, sesiones y perfil.")
    sales_rows = [
        ["Operaciones", "Resumen de cuenta, métricas, accesos rápidos y estado general.", "GET /portal/overview"],
        ["Onboarding", "Pasos de activación, claves, usuarios y unidades necesarias para arrancar.", "GET /portal/onboarding"],
        ["Comercial", "Planes, checkout, confirmación, webhook y activación de suscripción.", "/commercial/*"],
        ["Cuenta", "Plan contratado, cambio o cancelación, facturas y sesiones abiertas.", "/account/*"],
        ["Operación", "Unidades, rutas guardadas, sesiones de ruta, posiciones, métricas e incidencias.", "/vehicles, /navigation, /incidents"],
        ["Control", "Usuarios, claves de activación, documentos y revisión administrativa.", "/users, /admin, /documents"],
    ]
    add_table(doc, ["Área", "Responsabilidad", "Integración principal"], sales_rows, [1700, 4760, 2900], font_size=8.5)
    add_body(doc, "El estado del portal se conserva en un store compartido. Las cargas relacionadas se agrupan, se evita duplicar solicitudes concurrentes y se aplica una ventana corta de caché. Después de una acción —por ejemplo cambiar el plan, revocar una sesión o revisar un documento— se actualiza el estado local y se solicita únicamente la información necesaria.")
    doc.add_page_break()

    add_heading(doc, "7.4 Interconexión entre app, ventas y backend", 2)
    add_body(doc, "Las dos aplicaciones cliente no se conectan directamente entre sí. El backend actúa como punto de coordinación: autentica, aplica permisos, filtra por organizationId, persiste cambios y publica eventos a las salas correctas. De esta manera una acción comercial puede habilitar capacidades operativas sin compartir bases locales ni duplicar reglas de negocio.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(p, assets["ecosystem"], Inches(6.35), "Interconexión entre aplicación móvil, backend central y portal web de ventas")
    add_caption(doc, "Figura 12. Interconexión funcional de la app móvil y el portal de ventas mediante el backend.")
    for step in [
        "El cliente inicia sesión y recibe identidad, rol, organización, access token y mecanismo de renovación.",
        "Las solicitudes REST llegan a módulos especializados; los middlewares verifican autenticación, permiso y acceso al portal u operación.",
        "El backend consulta o modifica la persistencia y devuelve una respuesta normalizada al cliente que originó la acción.",
        "Cuando el cambio afecta a otros usuarios, Socket.IO emite el evento a salas por usuario, rol u organización.",
        "El portal aplica eventos comerciales como payment:confirmed, plan:active, subscription:updated u onboarding:updated y refresca la vista correspondiente.",
        "La app recibe eventos operativos de ubicación, chat, radio, notificaciones o señalización RTC y actualiza la experiencia en tiempo real.",
    ]:
        add_numbered(doc, step)

    add_heading(doc, "7.5 Sistema de llamadas y comunicación RTC", 2)
    add_body(doc, "La base de llamadas ya está implementada sobre WebRTC. El backend entrega configuración ICE autenticada; siempre incluye STUN y añade TURN cuando existen credenciales dinámicas o estáticas. Socket.IO se utiliza como canal de señalización para coordinar participantes y registrar sesiones, mientras que el audio o video viaja por la conexión RTC entre clientes o a través de TURN cuando la red no permite conexión directa.")
    call_rows = [
        ["Configuración", "GET /rtc/config entrega iceServers y modo de credenciales.", "Implementado"],
        ["Señalización", "Eventos Socket.IO coordinan unión, oferta, respuesta, candidatos y cierre.", "Implementado"],
        ["Cliente", "Capa WebRTC detecta disponibilidad y crea descripciones y candidatos ICE.", "Implementado"],
        ["Android", "Servicio foreground conserva micrófono/cámara durante una llamada en segundo plano.", "Implementado"],
        ["Auditoría", "Administradores pueden consultar sesiones RTC registradas.", "Implementado"],
        ["Robustez futura", "TURN administrado, métricas de calidad, reconexión y pruebas multioperador.", "Planificado"],
    ]
    add_table(doc, ["Componente", "Detalle", "Estado"], call_rows, [1900, 5660, 1800], font_size=8.5)
    add_small_note(doc, "Limitación actual", "Sin TURN configurado el sistema opera en modo STUN-only. Puede funcionar en redes sencillas, pero no garantiza llamadas estables detrás de NAT restrictivo, redes corporativas o ciertos operadores móviles.")
    doc.add_page_break()

    add_heading(doc, "7.6 Sistema de correo y notificaciones", 2)
    add_body(doc, "El correo comercial se integra mediante el módulo de comunicación y el proveedor Resend. El servicio selecciona una plantilla según el evento y el estado de pago, incorpora datos de la orden y registra proveedor, plantilla, resultado, error y fecha de contacto. Si faltan RESEND_API_KEY o RESEND_FROM_EMAIL, el envío se omite de forma explícita y el estado queda marcado como no configurado.")
    email_rows = [
        ["order-created", "Confirmación de orden, referencia, plan, monto y siguiente paso."],
        ["payment-pending", "Pago pendiente o en espera de configuración o confirmación manual."],
        ["payment-approved", "Pago aprobado y continuidad de activación."],
        ["payment-rejected", "Pago fallido, cancelado o rechazado con instrucciones de recuperación."],
        ["subscription-activated", "Cuenta y suscripción activadas; acceso al dashboard."],
        ["subscription-cancelled", "Confirmación de cancelación y estado final de la cuenta."],
    ]
    add_table(doc, ["Plantilla", "Propósito"], email_rows, [2800, 6560], font_size=8.8)
    add_body(doc, "Las notificaciones operativas siguen otro canal complementario: se guardan en la cuenta, se emiten por Socket.IO y, cuando existen suscripciones válidas, se envían como push. El cliente Android puede mostrar categorías de chat, radio, SOS, incidencias y notificaciones, con deep links hacia la pantalla correspondiente. El registro de token push remoto permanece como trabajo de integración hasta disponer de credenciales FCM o un proveedor equivalente.")

    add_heading(doc, "7.7 Actualizaciones futuras recomendadas", 2)
    roadmap_rows = [
        ["Corto plazo", "Correo", "Configurar dominio, SPF, DKIM y DMARC; activar Resend en producción; verificar rebotes y quejas."],
        ["Corto plazo", "Llamadas", "Desplegar TURN con credenciales dinámicas, TLS y monitoreo de disponibilidad."],
        ["Corto plazo", "Push", "Integrar FCM/Notifee, alta y rotación de tokens, apertura fiable por deep link."],
        ["Mediano plazo", "Correo", "Cola persistente, reintentos con backoff, idempotencia, métricas y panel de entregabilidad."],
        ["Mediano plazo", "Llamadas", "Timbrado entrante, reconexión, selección de dispositivo, mute, altavoz y transferencia."],
        ["Mediano plazo", "Ventas", "Automatizar recordatorios de pago, renovaciones, vencimientos y secuencias de onboarding."],
        ["Largo plazo", "RTC", "Indicadores de jitter, pérdida, latencia y duración; alertas y análisis de calidad por operador."],
        ["Largo plazo", "Omnicanal", "Unificar historial de correo, WhatsApp, push, llamadas e incidencias por cuenta."],
    ]
    add_table(doc, ["Horizonte", "Área", "Actualización"], roadmap_rows, [1600, 1700, 6060], font_size=8.4)
    doc.add_page_break()


def add_security(doc: Document):
    add_heading(doc, "8. Seguridad", 1)
    add_body(doc, "El sistema aplica controles de autenticación, autorización, protección de credenciales, validación de entrada y observabilidad. Estos controles reducen riesgos comunes en APIs móviles y servicios expuestos a internet.")
    rows = [
        ["JWT", "El backend firma tokens con subject de usuario, rol, correo, organización y sesión.", "Expiración corta, refresh token rotado y validación en middleware authenticate."],
        ["Contraseñas", "Las contraseñas se guardan como hash bcrypt.", "Nunca devolver passwordHash; política mínima de 8 caracteres con letra, número y símbolo."],
        ["Roles y permisos", "Roles owner, admin, dispatcher, supervisor, billing_manager, support, viewer y driver.", "Permisos por módulo; tenant access por organizationId."],
        ["Protección de APIs", "Helmet, CORS configurable, rate limit, trazabilidad x-trace-id y manejo centralizado de errores.", "Reducir superficie de ataque y facilitar auditoría."],
        ["Validaciones", "Campos obligatorios, catálogos, coordenadas, fechas, correos únicos y montos positivos.", "Validar en cliente y backend, priorizando backend."],
        ["Tiempo real", "Socket.IO exige token y valida pertenencia de unidad para location:update.", "Evitar spoofing de ubicación y acceso a salas no permitidas."],
        ["Datos sensibles", "Cifrado E2EE opcional en chat directo y almacenamiento seguro de tokens en cliente.", "Evitar exposición de respaldos, tokens push y llaves privadas."],
    ]
    add_table(doc, ["Control", "Descripción", "Aplicación en el sistema"], rows, [1700, 3550, 4110], font_size=8.5)
    add_heading(doc, "8.1 Recomendaciones de endurecimiento", 2)
    for item in [
        "Configurar JWT_SECRET fuerte y diferente por ambiente; no usar valores por defecto en producción.",
        "Exigir HTTPS para API, sockets y descargas de documentos.",
        "Activar MongoDB con REQUIRE_MONGO=true, respaldos y usuarios de base de datos con mínimos privilegios.",
        "Registrar auditoría para cambios de usuarios, documentos, rutas, pagos e incidencias críticas.",
        "Aplicar pruebas de autorización por rol y por organización para evitar acceso cruzado entre empresas.",
        "Mantener límites de carga para archivos y escanear documentos si se reciben adjuntos de usuarios externos.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def add_conclusions(doc: Document):
    add_heading(doc, "9. Conclusiones", 1)
    add_body(doc, "El Sistema Inteligente de Gestión de Combis representa una solución integral para operaciones de transporte público que requieren visibilidad, control documental, comunicación inmediata, seguimiento GPS y trazabilidad de viajes. La combinación de app móvil, backend REST, base de datos persistente y eventos en tiempo real permite convertir procesos manuales en flujos digitales auditables.")
    add_body(doc, "Desde el punto de vista de ingeniería de software, el diseño separa responsabilidades entre frontend, backend, persistencia e integraciones. Esta separación facilita mantenimiento, escalabilidad, pruebas y evolución hacia modelos multiempresa, paneles analíticos, optimización de rutas y mantenimiento predictivo.")
    add_heading(doc, "9.1 Mejoras futuras", 2)
    for item in [
        "Seguimiento GPS continuo en segundo plano con reglas de batería y privacidad.",
        "Permisos granulares por ruta, patio, organización y tipo de acción.",
        "Reportes exportables en PDF/CSV para viajes, puntualidad, documentos, pagos e incidencias.",
        "Optimización de rutas con demanda histórica, tráfico y horarios pico.",
        "Módulo de mantenimiento preventivo por kilometraje, evidencia y vencimientos.",
        "Portal público para pasajeros con tiempos estimados, avisos e incidencias relevantes.",
    ]:
        add_bullet(doc, item)
    add_body(doc, "El impacto tecnológico esperado es una operación más coordinada, con menos información dispersa, mejores tiempos de respuesta y una base de datos capaz de sostener análisis posteriores para mejorar seguridad, puntualidad y calidad del servicio.")
    doc.add_page_break()


def add_references_and_appendix(doc: Document):
    add_heading(doc, "Referencias bibliográficas", 1)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(ref)
        set_run_font(r, size=10, color=BLACK)

    add_heading(doc, "Anexo A. Matriz de trazabilidad técnica", 1)
    rows = [
        ["Autenticación", "backend/src/modules/auth/routes.js", "Login, registro, refresh token, logout y sesión actual."],
        ["JWT", "backend/src/utils/jwt.js", "Firma y verificación de tokens de acceso."],
        ["Permisos", "backend/src/middlewares/access-control.js", "Roles, permisos y aislamiento por organización."],
        ["Modelos", "backend/src/data/models.js", "Schemas MongoDB para usuarios, vehículos, rutas, viajes y eventos."],
        ["Sockets", "backend/src/sockets/index.js", "Presencia, chat, ubicación y RTC en tiempo real."],
        ["Navegación", "backend/src/modules/navigation/routes.js", "Búsqueda, cálculo, asignación y bitácora de viajes."],
        ["Cliente API", "mobile/src/api/client.ts", "Consumo REST, reintentos, refresh y trazabilidad x-trace-id."],
    ]
    add_table(doc, ["Área", "Archivo", "Responsabilidad"], rows, [1800, 3600, 3960], font_size=8.6)

    add_heading(doc, "Anexo B. Glosario", 1)
    glossary = [
        ["API REST", "Interfaz de comunicación basada en recursos HTTP y mensajes JSON."],
        ["JWT", "Token firmado que transporta identidad y claims de autorización."],
        ["Socket.IO", "Canal bidireccional usado para eventos de baja latencia."],
        ["MER", "Modelo Entidad-Relación que describe entidades, atributos y cardinalidades."],
        ["GPS", "Sistema de posicionamiento usado para ubicar combis en tiempo real."],
        ["GridFS", "Mecanismo de MongoDB para almacenar archivos grandes en la base de datos."],
        ["RTC", "Comunicación en tiempo real para sesiones de audio/video o radio operativo."],
    ]
    add_table(doc, ["Término", "Definición"], glossary, [1900, 7460], font_size=9)


def build_doc():
    assets = create_assets()
    doc = Document()
    configure_document(doc)
    add_cover(doc, assets)
    add_toc(doc)
    add_intro(doc)
    add_general_description(doc, assets)
    add_data_dictionary(doc)
    add_mer(doc, assets)
    add_relational_model(doc, assets)
    add_process_model(doc, assets)
    add_architecture(doc, assets)
    add_security(doc)
    add_conclusions(doc)
    add_references_and_appendix(doc)

    doc.core_properties.title = "Documentación técnica del Sistema Inteligente de Gestión de Combis"
    doc.core_properties.subject = "Diccionario de datos, MER, modelo relacional, procesos, arquitectura y seguridad"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Documento generado para entrega universitaria."
    doc.save(OUT_DOCX)
    set_update_fields_on_open(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    output = build_doc()
    print(output)
