from __future__ import annotations

import shutil
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "actividad-3-2-documentacion-asistencia-tecnica-manecomb.docx"
CONTENT_WIDTH = 9360
BLACK = RGBColor(0, 0, 0)
NAVY = RGBColor(31, 55, 78)
BLUE = RGBColor(47, 84, 150)
GRAY = RGBColor(89, 89, 89)
LIGHT = "EAF0F6"


def font(run, size=12, bold=False, italic=False, color=BLACK):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    node = pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        pr.append(node)
    node.set(qn("w:fill"), fill)


def cell_width(cell, width):
    pr = cell._tc.get_or_add_tcPr()
    node = pr.find(qn("w:tcW"))
    if node is None:
        node = OxmlElement("w:tcW")
        pr.append(node)
    node.set(qn("w:w"), str(width))
    node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    pr = table._tbl.tblPr
    for tag, attrs in [
        ("tblW", {"w": str(sum(widths)), "type": "dxa"}),
        ("tblInd", {"w": "120", "type": "dxa"}),
        ("tblLayout", {"type": "fixed"}),
    ]:
        node = pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            pr.append(node)
        for key, value in attrs.items():
            node.set(qn(f"w:{key}"), value)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for i, cell in enumerate(row.cells):
            cell_width(cell, widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    pr.append(node)


def add_table(doc, headers, rows, widths, size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table_geometry(table, widths)
    repeat_header(table.rows[0])
    for i, value in enumerate(headers):
        shade(table.cell(0, i), "1F374E")
        p = table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        font(r, size=size, bold=True, color=RGBColor(255, 255, 255))
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if index % 2:
                shade(cells[i], "F5F7FA")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value))
            font(r, size=size, bold=(i == 0))
    doc.add_paragraph()
    return table


def paragraph(doc, text, bold_lead=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        font(r)
    else:
        r = p.add_run(text)
        font(r)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        font(r)
    if not p.runs:
        font(p.add_run(text))
    else:
        p.runs[0].text = text


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(text))


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def callout(doc, label, text, fill="EEF3F8"):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [CONTENT_WIDTH])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(f"{label}. ")
    font(r, bold=True, color=NAVY)
    font(p.add_run(text))
    doc.add_paragraph()


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])
    font(run, size=12)


def add_toc(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "Índice automático. Actualice los campos en Word para mostrar la numeración definitiva."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, cached, end])
    font(run)


def update_fields(path):
    tmp = path.parent / "_apa_fields"
    if tmp.exists():
        shutil.rmtree(tmp)
    tmp.mkdir()
    with zipfile.ZipFile(path) as zf:
        zf.extractall(tmp)
    settings = tmp / "word" / "settings.xml"
    tree = etree.parse(str(settings))
    root = tree.getroot()
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    node = root.find(f"{{{ns}}}updateFields")
    if node is None:
        node = etree.Element(f"{{{ns}}}updateFields")
        root.insert(0, node)
    node.set(f"{{{ns}}}val", "true")
    tree.write(str(settings), xml_declaration=True, encoding="UTF-8", standalone="yes")
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for item in tmp.rglob("*"):
            if item.is_file():
                zf.write(item, item.relative_to(tmp).as_posix())
    shutil.rmtree(tmp)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(8)
    for name, size, before, after in [("Heading 1", 14, 12, 8), ("Heading 2", 12, 10, 6), ("Heading 3", 12, 8, 4)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
    add_page_number(section.footer.paragraphs[0])


def cover(doc):
    today = date.today()
    months = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("MANECOMB"), size=20, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Documentación y Asistencia Técnica Centrada en el Usuario"), size=16, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Actividad 3.2"), size=14, bold=True, color=BLUE)
    doc.add_paragraph()
    rows = [
        ["Materia", "Usabilidad - Diseño de Interfaces"],
        ["Proyecto", "ManeComb"],
        ["Equipo", "Equipo XX"],
        ["Integrantes", "____________________________________________\n____________________________________________"],
        ["Docente", "____________________________________________"],
        ["Fecha", f"{today.day} de {months[today.month - 1]} de {today.year}"],
    ]
    add_table(doc, ["Dato", "Información"], rows, [2300, 7060], size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Documento académico en formato APA 7"), italic=True, color=GRAY)
    doc.add_page_break()


def introduction(doc):
    heading(doc, "Introducción", 1)
    paragraph(doc, "La documentación constituye un componente esencial de cualquier sistema interactivo porque traduce las funciones técnicas en acciones comprensibles para las personas. Una aplicación puede disponer de numerosos módulos y, aun así, resultar difícil de utilizar cuando los usuarios no comprenden qué deben hacer, qué información requiere el sistema o cómo recuperarse de un error. La documentación adecuada reduce esta distancia al ofrecer instrucciones, explicaciones, ejemplos y criterios de decisión alineados con las tareas reales.")
    paragraph(doc, "La asistencia técnica amplía esta función al proporcionar apoyo en el momento en que aparece una dificultad. No se limita a un manual externo: también incluye mensajes dentro de la interfaz, ayudas contextuales, advertencias preventivas, preguntas frecuentes, tutoriales y mecanismos de recuperación. Desde la perspectiva del diseño centrado en el usuario, estos recursos deben construirse a partir de necesidades observables, lenguaje familiar y condiciones de uso concretas, en lugar de reflejar únicamente la estructura interna del software.")
    paragraph(doc, "La usabilidad depende de que una persona pueda alcanzar sus objetivos con eficacia, eficiencia y satisfacción. Cuando un sistema explica con claridad sus estados, confirma acciones importantes y ofrece orientación antes de que se cometa un error, disminuye la carga cognitiva y mejora el aprendizaje. Norman (2013) señala la relevancia de la visibilidad, la retroalimentación y las correspondencias comprensibles entre controles y resultados. De manera complementaria, Shneiderman (2013) destaca la consistencia, la prevención de errores y el apoyo al control del usuario.")
    paragraph(doc, "ManeComb requiere especialmente estos elementos porque combina tareas operativas sensibles: asignar rutas, cargar documentos, monitorear unidades, reportar incidencias y comunicarse mediante radio PTT. Un error puede afectar la continuidad de una ruta, el cumplimiento documental o la coordinación del personal. Por ello, este trabajo diagnostica puntos de fricción, propone documentación impresa y digital, selecciona modalidades de asistencia y formula un prototipo accesible que fortalece la experiencia general del sistema.")
    paragraph(doc, "Dentro de esa experiencia, el centro de alertas y el chat constituyen una capa transversal. Las alertas informan cambios que exigen atención —por ejemplo, pérdida de comunicación, desvíos o incidencias— y el chat permite coordinar una respuesta conservando participantes, tiempos y evidencia. Su valor aumenta cuando ambos se enlazan con la unidad, ruta o incidente correspondiente: el operador no solo recibe un aviso, sino que comprende su prioridad y puede conversar, asignar responsabilidad y dar seguimiento desde el mismo contexto.")
    doc.add_page_break()


def friction(doc, number, title, description, causes, impact, risk, current, solution):
    heading(doc, f"{number}. {title}", 2)
    paragraph(doc, f"Descripción. {description}", "Descripción. ")
    paragraph(doc, f"Por qué ocurre. {causes}", "Por qué ocurre. ")
    paragraph(doc, f"Impacto en el usuario. {impact}", "Impacto en el usuario. ")
    paragraph(doc, f"Riesgo. {risk}", "Riesgo. ")
    paragraph(doc, f"Estado actual. {current}", "Estado actual. ")
    paragraph(doc, f"Posible solución. {solution}", "Posible solución. ")


def communication_diagnosis(doc):
    heading(doc, "Diagnóstico complementario: alertas y chat", 2)
    paragraph(doc, "Aunque los tres puntos anteriores responden a la actividad académica, la experiencia completa de ManeComb también depende de cómo se comunica la operación. Las alertas deben convertir eventos técnicos en prioridades comprensibles, mientras que el chat debe conservar el contexto entre administración, conductores y personal de seguimiento. Ambos módulos conectan rutas, unidades, incidencias y decisiones humanas; por ello, una falla de comunicación puede ser tan crítica como un dato incorrecto.")
    add_table(doc, ["Módulo", "Fricción observable", "Consecuencia", "Mejora propuesta"], [["Centro de alertas", "Mensajes con igual peso visual, duplicados o sin indicar unidad, ruta, hora y gravedad.", "El operador tarda en distinguir una incidencia crítica de un aviso informativo y puede dejar eventos sin atender.", "Clasificar por severidad, agrupar duplicados, mostrar contexto y ofrecer Marcar atendida, Posponer y Ver incidencia."], ["Chat operativo", "No siempre es evidente si el mensaje está enviándose, fue entregado, leído o quedó pendiente por falta de conexión.", "Se repiten mensajes, se cambia innecesariamente a otro canal y se pierde continuidad en la coordinación.", "Mostrar estados textuales, reintento seguro, no leídos, hora, participantes y acceso directo desde alerta, ruta o unidad."], ["Interconexión", "La alerta y la conversación pueden quedar separadas del objeto que originó el problema.", "El usuario reconstruye manualmente qué unidad, ruta o incidencia se está discutiendo.", "Crear enlaces contextuales y una tarjeta compartida con identificador, prioridad, responsable y última actualización."]], [1500, 2820, 2520, 2520], size=9)
    heading(doc, "Criterios operativos para alertas", 3)
    for item in ["Distinguir niveles crítico, alto, medio e informativo mediante texto, icono y contraste, no solo color.", "Ordenar primero los eventos no atendidos y conservar un historial auditable de reconocimiento, responsable y fecha.", "Evitar notificaciones duplicadas mediante agrupación por unidad, tipo de evento y periodo.", "Abrir la ruta, unidad, incidencia o conversación relacionada sin obligar a buscarla nuevamente."]:
        bullet(doc, item)
    heading(doc, "Criterios operativos para chat", 3)
    for item in ["Conservar mensajes pendientes cuando la conexión sea inestable y explicar cuándo se reintentará el envío.", "Diferenciar enviado, entregado y leído con etiquetas accesibles y marcas de tiempo comprensibles.", "Permitir adjuntar evidencia y respuestas rápidas sin ocultar el contexto de la conversación.", "Escalar una conversación a incidencia o llamada cuando la urgencia exceda la comunicación asíncrona."]:
        bullet(doc, item)


def diagnosis(doc):
    heading(doc, "Fase 1. Diagnóstico", 1)
    heading(doc, "Diagnóstico de usabilidad", 2)
    paragraph(doc, "Un diagnóstico de usabilidad es un proceso sistemático para identificar obstáculos que dificultan la interacción entre una persona y un producto digital. Considera las metas del usuario, el contexto de uso, la claridad de la información, la consistencia de los controles, la prevención de errores, la retroalimentación y la capacidad de recuperación. Su propósito no consiste solamente en señalar defectos visuales, sino en comprender por qué se produce una dificultad y qué consecuencias tiene para la tarea.")
    paragraph(doc, "En ManeComb el diagnóstico se orienta a actividades que combinan información técnica y responsabilidad operativa. Se seleccionaron tres puntos de fricción representativos porque afectan a perfiles distintos y requieren modalidades diferentes de asistencia: la asignación de rutas, la carga documental y el uso de radio PTT.")
    friction(doc, "1.1", "Configuración y asignación de rutas", "El administrador debe elegir una ruta y vincularla con una unidad. La interfaz puede completar la acción sin que el usuario comprenda con certeza qué ruta permanece activa o cómo sustituirla.", "La tarea reúne selección de unidad, selección de ruta, confirmación y actualización del mapa. Si la jerarquía visual o la retroalimentación son insuficientes, el usuario interpreta el cambio a partir de suposiciones.", "Se incrementa el tiempo de verificación y aparece inseguridad después de guardar. Un usuario nuevo puede abandonar la pantalla, repetir la acción o solicitar apoyo externo.", "Una asignación equivocada puede dirigir una unidad a un recorrido incorrecto, afectar horarios, generar reportes inconsistentes y complicar el monitoreo.", "No existe ayuda contextual suficiente que explique la selección, la confirmación y el procedimiento para modificar una ruta ya asignada.", "Incorporar texto de ayuda junto al selector, mostrar origen y destino antes de confirmar, solicitar confirmación con nombre de unidad y ruta, presentar un mensaje de éxito persistente y ofrecer la acción «Cambiar ruta» con explicación breve.")
    doc.add_page_break()
    friction(doc, "1.2", "Carga de documentos del conductor", "El usuario debe seleccionar y enviar licencias, identificaciones u otras evidencias, pero no conoce con anticipación los formatos permitidos, el tamaño máximo ni el significado de los estados de revisión.", "La información técnica se encuentra implícita o aparece únicamente después de un intento fallido. Además, el proceso de validación administrativa ocurre en otro momento, por lo que el usuario no observa una relación inmediata entre la carga y la aprobación.", "La persona prueba archivos al azar, repite envíos, teme perder información y no sabe si debe esperar, corregir o contactar al administrador.", "La documentación incompleta o incorrecta puede impedir que un conductor sea habilitado, provocar vencimientos y aumentar la carga de soporte.", "El sistema permite cargar y revisar documentos, pero la asistencia previa y la explicación de estados pueden resultar insuficientes para usuarios sin experiencia.", "Mostrar formatos y tamaño antes de seleccionar, validar inmediatamente, visualizar progreso, distinguir «Cargado», «En revisión», «Aprobado» y «Rechazado», e incluir una razón de rechazo con acción directa para reemplazar el archivo.")
    doc.add_page_break()
    friction(doc, "1.3", "Uso del sistema Radio PTT", "La radio PTT exige mantener presionado para transmitir y liberar para finalizar. Los usuarios nuevos pueden tocar brevemente, hablar antes de que el canal esté disponible o confundir estados técnicos.", "El modelo de interacción no coincide con una llamada convencional. Además, los estados «Conectando», «Grabando», «Subiendo» y «Reproduciendo» representan etapas distintas que requieren retroalimentación visual, textual y, cuando sea pertinente, háptica.", "El usuario duda sobre si su mensaje fue capturado, repite transmisiones o interrumpe la comunicación de otros participantes.", "Los mensajes operativos pueden perderse, llegar tarde o duplicarse en situaciones que requieren coordinación inmediata.", "La función dispone de estados de radio, pero el aprendizaje depende demasiado de prueba y error y no siempre existe ayuda proactiva durante el primer uso.", "Agregar un tutorial inicial breve, texto «Mantén presionado para hablar», cambio de color y vibración al comenzar, indicador de subida, confirmación de envío y explicación accesible de cada estado.")
    add_table(doc, ["Punto", "Severidad", "Frecuencia probable", "Prioridad"], [["Asignación de rutas", "Alta", "Media", "Alta"], ["Carga documental", "Alta", "Alta", "Alta"], ["Radio PTT", "Media-alta", "Alta en usuarios nuevos", "Alta"]], [2700, 1800, 2460, 2400], size=10)
    communication_diagnosis(doc)
    doc.add_page_break()


def printed_manual(doc):
    heading(doc, "Fase 2. Documentación impresa", 1)
    heading(doc, "Cómo cargar correctamente un documento del conductor", 2)
    callout(doc, "Objetivo", "Orientar al usuario para seleccionar, validar y enviar un documento legible, reduciendo rechazos y repeticiones.")
    heading(doc, "Requisitos", 3)
    for item in ["Iniciar sesión con una cuenta autorizada.", "Tener el documento completo y vigente.", "Utilizar PDF o JPG.", "Verificar que el archivo no supere 10 MB.", "Asegurar que nombres, fechas y números sean legibles."]:
        bullet(doc, item)
    heading(doc, "Procedimiento paso a paso", 3)
    for item in ["Abra el módulo Documentos y seleccione al conductor correspondiente.", "Identifique el tipo de documento que desea cargar.", "Pulse Seleccionar archivo y localice el PDF o JPG.", "Revise la vista previa y confirme que no existan páginas cortadas, reflejos o texto borroso.", "Compruebe el nombre del archivo, formato y tamaño indicados por el sistema.", "Pulse Subir y espere hasta observar la confirmación de carga.", "Verifique el estado En revisión. El administrador notificará si el documento fue aprobado o rechazado."]:
        numbered(doc, item)
    heading(doc, "Notas importantes y recomendaciones", 3)
    for item in ["No cierre la pantalla mientras el archivo se está enviando.", "Evite fotografías inclinadas o con sombras.", "No cargue contraseñas, archivos comprimidos ni documentos ajenos.", "Conserve el archivo original hasta recibir aprobación.", "Si el documento es rechazado, lea la causa antes de sustituirlo."]:
        bullet(doc, item)
    heading(doc, "Errores comunes", 3)
    add_table(doc, ["Error", "Cómo prevenirlo"], [["Formato no permitido", "Convierta el documento a PDF o JPG."], ["Archivo mayor de 10 MB", "Comprima el archivo sin perder legibilidad."], ["Imagen borrosa", "Repita la captura con buena iluminación."], ["Documento vencido", "Compruebe la vigencia antes de enviarlo."], ["Carga interrumpida", "Revise la conexión y vuelva a intentarlo una sola vez."]], [3500, 5860], size=10)
    callout(doc, "Espacio para evidencia", "[Inserte aquí captura etiquetada del sistema]\n\nFigura 1. Selección del tipo de documento.\nFigura 2. Validación del archivo antes del envío.\nFigura 3. Confirmación y estado de revisión.", fill="F7F7F7")
    paragraph(doc, "El formato impreso es adecuado para esta tarea porque ofrece una referencia estable que puede consultarse durante capacitación, entregarse a operadores y utilizarse incluso cuando no existe conexión. De acuerdo con Lidwell, Holden y Butler (2010), la reducción de complejidad, la jerarquía visual y el reconocimiento favorecen el aprendizaje. La secuencia numerada, los errores comunes y las figuras previstas permiten localizar información sin memorizar todo el procedimiento.")
    doc.add_page_break()


def digital(doc):
    heading(doc, "Documentación digital", 1)
    callout(doc, "Base de conocimiento ManeComb", "🔎 Buscar en ayuda: [________________________________________]", fill="E7F0FA")
    paragraph(doc, "Categorías: Operación | Alertas | Chat | Conductores | Documentos | Rutas | Radio PTT | Cuenta", align=WD_ALIGN_PARAGRAPH.CENTER)
    heading(doc, "Cómo cargar correctamente un documento del conductor", 2)
    paragraph(doc, "Este artículo explica cómo preparar, seleccionar y enviar un documento para revisión administrativa. Tiempo estimado: 3 minutos.")
    heading(doc, "Antes de comenzar", 3)
    for item in ["Cuenta autorizada y conductor seleccionado.", "Archivo PDF o JPG de máximo 10 MB.", "Documento completo, vigente y legible."]:
        bullet(doc, item)
    heading(doc, "Pasos", 3)
    for item in ["Entre a Documentos.", "Elija el conductor y el tipo de documento.", "Pulse Seleccionar archivo.", "Revise la vista previa, formato y tamaño.", "Pulse Subir y espere la confirmación.", "Consulte el estado En revisión, Aprobado o Rechazado."]:
        numbered(doc, item)
    callout(doc, "Nota", "Los documentos son revisados por un administrador. La carga correcta no equivale a aprobación inmediata.")
    callout(doc, "Advertencia", "No abandone la pantalla durante el envío. Si aparece un error, compruebe la conexión antes de repetir.", fill="FFF2CC")
    heading(doc, "Preguntas frecuentes", 3)
    add_table(doc, ["Pregunta", "Respuesta"], [["¿Qué formatos se aceptan?", "PDF y JPG."], ["¿Cuál es el tamaño máximo?", "10 MB por archivo."], ["¿Cuánto tarda la validación?", "Depende de la revisión administrativa; el estado se actualizará en el sistema."], ["¿Qué hago si fue rechazado?", "Consulte la causa, corrija el archivo y utilice Reemplazar documento."], ["¿Puedo cerrar la aplicación?", "Después de recibir la confirmación de carga, sí."]], [3900, 5460], size=10)
    callout(doc, "Recurso multimedia", "▶ Video tutorial\nDuración: 3:20", fill="EEF3F8")
    paragraph(doc, "Ver también:\n• Cómo actualizar documentos\n• Validación de documentos\n• Recuperar documentos rechazados")
    heading(doc, "Artículos operativos relacionados", 2)
    add_table(doc, ["Artículo", "Contenido principal", "Acción inmediata"], [["Interpretar y atender una alerta", "Severidad, origen, unidad, ruta, hora, estado y responsable.", "Abrir el contexto, reconocer el evento o escalarlo a incidencia."], ["Usar el chat durante una operación", "Participantes, mensajes no leídos, entrega, lectura, adjuntos y conexión.", "Responder, reintentar, compartir evidencia o iniciar llamada."], ["Pasar de una alerta al chat", "Conversación vinculada al evento sin perder su identificador ni prioridad.", "Coordinar responsables y registrar la decisión tomada."]], [2700, 3840, 2820], size=9)
    callout(doc, "Preguntas frecuentes de comunicación", "¿Por qué recibí esta alerta? Abra Detalles para consultar el evento que la originó.\n¿Mi mensaje llegó? Revise la etiqueta Enviado, Entregado, Leído o Pendiente.\n¿La alerta ya está resuelta? El estado Atendida indica responsable y fecha; no elimina el historial.", fill="EEF3F8")
    paragraph(doc, "El formato digital conserva el contenido esencial del manual y añade búsqueda, navegación interna, actualización inmediata y recursos multimedia. Los principios de accesibilidad, revelación progresiva y reconocimiento descritos por Lidwell et al. (2010) justifican presentar primero la información indispensable y permitir que el usuario amplíe detalles mediante preguntas frecuentes y enlaces relacionados.")
    doc.add_page_break()


def assistance(doc):
    heading(doc, "Fase 3. Asistencia técnica", 1)
    paragraph(doc, "La asistencia debe corresponder con el momento, la complejidad y el riesgo de cada tarea. Una única modalidad no resuelve todas las dificultades: algunas requieren una aclaración junto al control, otras necesitan analizar el archivo y otras deben anticiparse durante el primer uso.")
    add_table(doc, ["Punto de fricción", "Tipo de asistencia", "Justificación", "Aplicación en ManeComb"], [["Configuración de rutas", "Ayuda contextual", "La duda aparece junto a selectores y confirmación.", "Textos breves, resumen de ruta, confirmación y acción Cambiar ruta."], ["Carga de documentos", "Asistencia inteligente", "El sistema puede detectar formato, tamaño, vigencia aparente y legibilidad antes del envío.", "Validación inmediata, mensajes específicos y recomendación según el error."], ["Radio PTT", "Ayuda activa / proactiva", "La interacción es poco familiar y el error afecta comunicación inmediata.", "Tutorial inicial, indicaciones durante primeros usos, color, vibración y estados explicados."]], [2100, 2100, 2580, 2580], size=9)
    heading(doc, "Justificación de las decisiones", 2)
    paragraph(doc, "La ayuda contextual es apropiada para la asignación de rutas porque mantiene la explicación cerca del control y evita interrumpir el flujo con un manual completo. El usuario necesita comprender la consecuencia de su selección en el instante en que la realiza. Un resumen con unidad, origen, destino y confirmación reduce la posibilidad de una asociación errónea.")
    paragraph(doc, "La asistencia inteligente resulta pertinente en documentos porque el sistema puede evaluar condiciones objetivas. En lugar de mostrar un mensaje genérico después del fallo, puede indicar que el archivo excede 10 MB, que el formato no es compatible o que la imagen tiene baja resolución. Esto convierte la asistencia en una herramienta de prevención y recuperación.")
    paragraph(doc, "La ayuda activa o proactiva se justifica en Radio PTT debido a que el gesto de mantener presionado no siempre es evidente. Durante los primeros usos, ManeComb puede explicar el control antes de la transmisión y reforzar cada etapa con texto, color y respuesta háptica. Después de que el usuario demuestra dominio, la ayuda puede reducirse para no obstaculizar la operación.")
    heading(doc, "Asistencia complementaria para comunicación operativa", 2)
    add_table(doc, ["Módulo", "Tipo de asistencia", "Aplicación"], [["Alertas", "Proactiva y basada en prioridad", "Notifica solo cuando corresponde, explica gravedad y causa, agrupa repeticiones y ofrece acciones para atender o escalar."], ["Chat", "Contextual e inteligente", "Explica estados de entrega, conserva borradores, detecta desconexión y propone reintento, llamada o creación de incidencia."], ["Alertas + chat", "Asistencia conectada", "Abre una conversación con el contexto del evento y mantiene visibles unidad, ruta, prioridad y responsable."]], [1800, 2500, 5060], size=9)
    paragraph(doc, "La asistencia no debe aumentar el ruido operativo. Una alerta crítica requiere persistencia y confirmación; un aviso informativo puede permanecer en el centro de notificaciones sin interrumpir. En chat, la ayuda debe aparecer cuando existe una condición real —por ejemplo, desconexión o fallo de entrega— y retirarse cuando el usuario recupera el flujo normal.")
    doc.add_page_break()


def prototype(doc):
    heading(doc, "Prototipo visual de ayuda contextual", 1)
    callout(doc, "Prototipo textual", "Subir documento\n\n[ Seleccionar archivo ]\n\nⓘ Solo PDF o JPG\nMáximo 10 MB\nLos documentos serán revisados por el administrador.\n\n[ Subir ]", fill="EAF0F6")
    heading(doc, "Funcionamiento propuesto", 2)
    paragraph(doc, "La información esencial se presenta antes de que el usuario seleccione el archivo. El icono informativo acompaña un texto visible, por lo que no es la única forma de comunicar la restricción. Después de elegir el archivo, el sistema muestra nombre, formato, tamaño y una vista previa. El botón Subir permanece deshabilitado hasta que las condiciones sean válidas.")
    paragraph(doc, "Durante el envío se sustituye la etiqueta por «Subiendo…» y se muestra progreso. Al finalizar, aparece una confirmación con el estado «En revisión». Si ocurre un error, el foco se mueve al mensaje, se explica la causa y se conserva el archivo seleccionado cuando sea seguro reintentar. Esta retroalimentación permite que el usuario entienda qué sucedió y qué acción está disponible.")
    heading(doc, "Prototipo del centro de alertas", 2)
    callout(doc, "Alerta operativa", "CRÍTICA · Unidad C-24 sin comunicación\nRuta: Centro–Terminal · Hace 2 min\nResponsable: Sin asignar\n\n[ Ver unidad ]  [ Abrir chat ]  [ Marcar atendida ]", fill="FFF2CC")
    paragraph(doc, "La tarjeta combina severidad escrita, contexto, antigüedad y responsable. Abrir chat crea o recupera una conversación vinculada a la alerta; Marcar atendida solicita responsable y conserva el evento en el historial. Las alertas repetidas de la misma unidad se agrupan y muestran su cantidad, evitando que el centro se convierta en una lista de duplicados.")
    heading(doc, "Prototipo de chat operativo", 2)
    callout(doc, "Chat · Unidad C-24", "Conductor: conexión inestable\nOperaciones · 10:42 · Leído\n«Confirma tu ubicación cuando recuperes señal.»\n\nConductor · 10:44 · Pendiente de envío\n«Estoy cerca de la Terminal.»\n\n[ Adjuntar ]  [ Escribir mensaje… ]  [ Enviar ]\nAlerta vinculada: C-24 sin comunicación", fill="E7F0FA")
    paragraph(doc, "El encabezado conserva la unidad y el estado de conexión. Cada mensaje comunica autor, hora y estado sin depender exclusivamente de iconos. Si falla la red, el texto permanece visible como Pendiente de envío, se habilita Reintentar y se evita duplicarlo. La alerta vinculada permite regresar al evento sin perder el hilo de coordinación.")
    heading(doc, "Cumplimiento con WCAG 2.2", 1)
    heading(doc, "Texto alternativo", 2)
    paragraph(doc, "La vista previa y los iconos informativos deben incluir nombres accesibles que expliquen su propósito. Una captura del documento puede anunciarse como «Vista previa del archivo seleccionado», mientras el icono debe asociarse al texto de formatos y tamaño. No se debe depender de imágenes para transmitir instrucciones indispensables.")
    heading(doc, "Operabilidad mediante teclado", 2)
    paragraph(doc, "Todos los controles deben recibir foco en un orden lógico. Seleccionar archivo, eliminar, reemplazar y subir deben activarse mediante teclado. El indicador de foco necesita ser visible y los mensajes dinámicos deben anunciarse sin obligar al usuario a recorrer nuevamente toda la página.")
    heading(doc, "Contraste adecuado", 2)
    paragraph(doc, "Texto, bordes, controles y estados deben alcanzar relaciones de contraste acordes con WCAG 2.2. El color puede reforzar éxito, advertencia o error, pero nunca debe ser el único indicador; se combina con texto e iconografía.")
    heading(doc, "Mensajes comprensibles", 2)
    paragraph(doc, "Los mensajes emplean lenguaje directo y explican cómo corregir el problema. «El archivo supera 10 MB; comprímalo o seleccione otro» resulta más útil que «Error de validación». Esta redacción beneficia a personas nuevas, reduce solicitudes de soporte y cumple el propósito de identificación de errores.")
    paragraph(doc, "En alertas y chat, los cambios dinámicos deben anunciarse mediante regiones de estado sin mover el foco inesperadamente. Una alerta crítica se anuncia con gravedad y contexto; un mensaje nuevo identifica conversación y remitente. La severidad incluye texto e icono, las marcas de entrega tienen nombres accesibles y las acciones poseen áreas de interacción suficientes. El usuario puede pausar avisos no críticos, recorrer mensajes con teclado y volver al elemento que originó la conversación.")
    doc.add_page_break()


def integration_conclusion(doc):
    heading(doc, "Fase 4. Integración", 1)
    paragraph(doc, "Como parte de la integración se plantea una revisión entre los equipos responsables de análisis, diseño, desarrollo y pruebas. La revisión conjunta permite comprobar que las instrucciones coincidan con el comportamiento real del sistema, que el lenguaje sea comprensible y que las recomendaciones de accesibilidad puedan implementarse sin contradicciones.")
    callout(doc, "Observación simulada", "El lenguaje técnico puede simplificarse para usuarios nuevos.", fill="FFF2CC")
    paragraph(doc, "La observación se atendió sustituyendo expresiones internas por acciones y resultados. Por ejemplo, «el archivo no cumple el MIME type permitido» se transformó en «seleccione un archivo PDF o JPG». De igual forma, «estado pending_review» se presentó como «En revisión». Se conservaron los términos técnicos únicamente cuando aportan valor y se acompañaron con una explicación.")
    paragraph(doc, "También se revisó la coherencia entre manual impreso, artículo digital y prototipo. Los tres recursos utilizan el mismo límite de 10 MB, los mismos formatos y los mismos estados. Esta consistencia evita que una persona reciba instrucciones contradictorias según el canal consultado.")
    paragraph(doc, "La revisión se amplía a alertas y chat con participación de operación, backend, frontend y pruebas. Backend verifica el origen, prioridad, entrega y trazabilidad de eventos; frontend comprueba jerarquía, estados y navegación contextual; pruebas valida desconexión, reintentos, mensajes duplicados, lectores de pantalla y teclado; operación confirma que el vocabulario corresponda con decisiones reales. Así, términos como «socket desconectado» se transforman en «Sin conexión; el mensaje se enviará al recuperar señal».")
    heading(doc, "Conclusiones", 1)
    paragraph(doc, "La documentación y la asistencia técnica son elementos de diseño, no complementos posteriores. En ManeComb permiten que tareas críticas como asignar rutas, cargar documentos, atender alertas, coordinarse por chat o utilizar Radio PTT sean comprensibles para usuarios con diferentes niveles de experiencia. El diagnóstico mostró que la falta de retroalimentación, restricciones visibles, contexto y explicación de estados puede convertirse en errores operativos y solicitudes repetidas de soporte.")
    paragraph(doc, "La combinación de manual impreso, base de conocimiento y ayuda dentro de la interfaz ofrece apoyo en distintos contextos. El manual facilita capacitación y consulta estable; el artículo digital permite búsqueda, actualización y multimedia; la ayuda contextual acompaña la acción sin separar al usuario de su objetivo. La accesibilidad amplía estos beneficios al asegurar que la información pueda percibirse y operarse mediante diferentes medios.")
    paragraph(doc, "Las propuestas desarrolladas mejoran la experiencia de ManeComb porque previenen errores, explican consecuencias, ofrecen recuperación y fortalecen la confianza. Una implementación coherente puede reducir cargas incorrectas, asignaciones equivocadas, alertas desatendidas, mensajes duplicados y transmisiones fallidas. La interconexión entre evento, alerta, conversación, incidencia, ruta y unidad proporciona continuidad: el usuario comprende qué ocurrió, a quién afecta, quién responde y cuál es la siguiente acción.")
    doc.add_page_break()


def references(doc):
    heading(doc, "Referencias", 1)
    refs = [
        "Lidwell, W., Holden, K., & Butler, J. (2010). Universal principles of design. Rockport Publishers.",
        "Norman, D. (2013). The design of everyday things. Basic Books.",
        "Shneiderman, B. (2013). Designing the user interface: Strategies for effective human-computer interaction (5th ed.). Pearson.",
        "World Wide Web Consortium. (2023). Web Content Accessibility Guidelines (WCAG) 2.2. https://www.w3.org/TR/WCAG22/",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(8)
        font(p.add_run(ref))
    heading(doc, "Espacios reservados para capturas", 2)
    paragraph(doc, "Antes de la entrega final pueden insertarse capturas reales de ManeComb en los espacios de las Figuras 1, 2 y 3. Cada imagen debe conservar un título descriptivo, texto alternativo y una referencia dentro del cuerpo del documento.")


def build():
    doc = Document()
    configure(doc)
    cover(doc)
    heading(doc, "Índice", 1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()
    introduction(doc)
    diagnosis(doc)
    printed_manual(doc)
    digital(doc)
    assistance(doc)
    prototype(doc)
    integration_conclusion(doc)
    references(doc)
    doc.core_properties.title = "Documentación y asistencia técnica centrada en el usuario - ManeComb"
    doc.core_properties.subject = "Actividad 3.2 de Usabilidad y Diseño de Interfaces"
    doc.core_properties.author = "Equipo XX"
    doc.core_properties.comments = "Documento académico listo para completar integrantes, docente y capturas."
    doc.save(OUT)
    update_fields(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
