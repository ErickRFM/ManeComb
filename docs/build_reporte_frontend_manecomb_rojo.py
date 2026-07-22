from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

import build_reporte_integral_desarrollo_50_paginas as report


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "reporte-frontend-manecomb-app-ventas-estilo-rojo.docx"
RED = RGBColor(192, 0, 0)
DARK_RED = RGBColor(139, 0, 0)
GRAY = RGBColor(89, 89, 89)
LIGHT_RED = "FDECEC"
ALT_RED = "FFF5F5"

report.NAVY = DARK_RED
report.BLUE = RED
report.base.NAVY = DARK_RED
report.base.BLUE = RED


def p(doc, text, lead=None):
    return report.p(doc, text, lead=lead)


def h(doc, text, level=1):
    return report.h(doc, text, level)


def bullets(doc, items):
    return report.bullets(doc, items)


def table(doc, headers, rows, widths, size=8):
    return report.table(doc, headers, rows, widths, size=size)


def callout(doc, label, text):
    return report.callout(doc, label, text, fill=LIGHT_RED)


def next_page(doc):
    report.next_page(doc)


TOPICS = [
    ("1. Arquitectura del frontend", "El frontend de ManeComb se divide en dos experiencias principales: la aplicación móvil React Native CLI y la web Ventas/Portal construida con React Native Web y Vite. Ambas consumen el mismo backend, pero responden a contextos distintos. Mobile prioriza operación en campo, conectividad variable, ubicación, comunicación y controles táctiles. Ventas cubre adquisición, compra y activación; Portal ofrece administración y supervisión desde escritorio.", ["Mobile concentra mapa, control, incidencias, chat, Radio PTT, llamadas y perfil.", "Ventas pública presenta planes, checkout, pago y acceso.", "Portal concentra Operaciones, Equipo, Unidades, Rutas, plan, pagos y cuenta.", "Stores Zustand mantienen estado de sesión, operación y experiencia comercial.", "Los clientes API normalizan contratos, errores y URLs por entorno."], "La separación evita forzar la misma pantalla a usuarios y condiciones incompatibles. La lógica autoritativa permanece en backend; el frontend interpreta estados, conserva contexto y ofrece recuperación."),
    ("2. Sistema visual y componentes compartidos", "El trabajo de interfaz pasó de estilos y controles locales hacia primitivas reutilizables. En Mobile existen AppCard, PrimaryButton, StatusPill, AppShell, banners y componentes de mapa. En Portal se introdujeron PortalButton, PortalDataList, PortalDataRow, cards, layout, tema y paleta. El objetivo fue reducir ajustes repetidos sin convertir casos operativos específicos en componentes universales.", ["Jerarquía tipográfica y colores por tema.", "Estados loading, disabled, error y vacío.", "Botones con variantes, tamaños, iconos y ancho completo.", "Filas composables con leading, body, meta y actions.", "Layouts responsive para escritorio, tableta y móvil web."], "Las extracciones se verificaron en runtime porque una migración anterior provocó un fallo de módulo aunque typecheck y build pasaban. Desde entonces, estilos dependientes se resuelven dentro del render."),
    ("3. Navegación móvil y estructura de pantallas", "La navegación móvil se centralizó en un router, registro de rutas y políticas de acceso. El arranque decide si debe mostrar autenticación, gate de cuenta o la aplicación operativa. Los deep links pueden dirigir a chat, radio o incidencias, pero se vuelven a validar contra sesión y permisos.", ["Registro explícito de rutas y pruebas de navegación.", "Tabs operativas en español.", "Redirecciones seguras durante bootstrap.", "Deep links desde notificaciones.", "Infraestructura para teclado e inputs sin superposición."], "La navegación conecta estados comerciales y operativos: una cuenta sin activación no entra directamente a funciones de flotilla, y una notificación nunca omite la verificación de autorización."),
    ("4. Acceso, registro y estados de cuenta", "Las pantallas de autenticación móvil y Ventas comparten principios visuales, pero sus objetivos difieren. Mobile autentica a usuarios operativos y dirige según rol y activación. Ventas registra o autentica compradores y los conduce al checkout o Portal. Los formularios incorporan validación, carga, recuperación y mensajes de error.", ["Campos con etiquetas y mensajes de ayuda.", "Indicadores de fortaleza de contraseña.", "Loading dentro del botón de envío.", "Recuperación de sesión mediante refresh.", "Branding consistente entre superficies."], "La UI presenta el estado que devuelve backend y evita diagnosticar credenciales por su cuenta. Los errores técnicos se transforman en mensajes accionables sin exponer detalles internos."),
    ("5. Mapa y seguimiento en la app", "El mapa móvil es una de las vistas principales. Se compone de MapCanvas, controles flotantes, HUD, selector de ruta, recuperación de datos y panel inferior. Los hooks separan ubicación, cámara, selección, sincronización y datos de tracking. La interfaz distingue una posición actual de una ubicación histórica o vencida.", ["Marcadores de unidades y polilíneas de ruta.", "Última actualización y frescura GPS.", "Selector de unidad y ruta.", "Panel inferior con jornada y recorrido.", "Fallback y recuperación cuando faltan datos."], "El frontend consume gpsFreshness calculado por backend. Esto corrigió etapas donde el rastreo podía parecer desconectado o mostrar datos antiguos sin suficiente claridad."),
    ("6. Control, jornada y rutas guardadas", "La pantalla de Control organiza las acciones del conductor durante la jornada. Permite seleccionar o crear rutas, iniciar, pausar, reanudar y finalizar sesiones, registrar checkpoints y consultar historial. Por su tamaño, fue una de las pantallas con mayor refinamiento de UI y lógica.", ["Estado visible de jornada.", "Acciones principales con protección contra doble envío.", "Mensajes de sincronización y cola offline.", "Editor de ruta en panel modal.", "Confirmación para eliminar rutas guardadas."], "Control se conecta con mapa y seguimiento. La acción elegida cambia lo que Operaciones ve en el Portal; por eso la interfaz conserva feedback incluso cuando la red es inestable."),
    ("7. Incidencias y confirmación operativa", "La pantalla de incidencias permite crear reportes, filtrar estados y severidad, visualizar evidencia y ejecutar SOS. La experiencia debe confirmar visualmente las acciones críticas y diferenciar carga, vacío y error.", ["Formulario de nueva incidencia.", "Severidad baja, media, alta y crítica.", "Estados abierta, en proceso y resuelta.", "Ubicación adjunta solo cuando es fresca.", "Actualización realtime de listas."], "La incidencia enlaza unidad, ruta, reportero y ubicación. Su confirmación no puede depender solo de vibración; debe indicar si fue registrada, quedó pendiente o necesita reintento."),
    ("8. Chat móvil y directorio", "El chat se refactorizó en una vista, encabezado, compositor, contenido multimedia y hooks de control. El directorio permite seleccionar conversaciones y contactos; el encabezado conserva participante, presencia y acciones de llamada. El compositor maneja texto, adjuntos y notas de voz sin que el teclado cubra los mensajes recientes.", ["Conversaciones generales y directas.", "Estados Pendiente, Enviado, Entregado, Leído y Fallido.", "Reintento sin duplicar mensajes.", "Adjuntos de imagen, video y audio.", "Acciones de llamada y radio."], "Las correcciones de teclado fueron importantes: el layout y scroll deben reaccionar a la apertura del teclado, mantener visible el compositor y preservar la posición de lectura."),
    ("9. Radio PTT", "Radio PTT utiliza una interfaz especializada con botón de transmisión, waveform, participantes, tarjetas de audio e indicadores de fase. El diseño enseña mantener presionado para hablar y comunica si el canal está disponible, ocupado, grabando, enviando o en error.", ["Respuesta háptica al comenzar transmisión.", "Temporizador y onda de audio.", "Reducer como fuente de estado de sesión.", "Mensajes para permisos y conexión.", "Historial y reproducción de transmisiones."], "Las fallas de radio observadas durante el desarrollo se relacionaron con conexión, arbitraje de canal, audio y estados duplicados. Las auditorías posteriores consolidaron lifecycle y SSOT."),
    ("10. Llamadas desde Chat", "Las llamadas de audio y video se presentan desde la conversación directa. La UI comunica iniciando, conectando, conectada, reconectando, finalizada o fallida, e incluye controles de micrófono, cámara y colgado.", ["Acciones estandarizadas en ChatHeader.", "Avisos busy, forbidden y timeout.", "Foreground service visible en Android.", "Controles accesibles de audio y video.", "Regreso a la conversación al finalizar."], "El frontend mantiene la máquina de estados; el backend valida ingreso y retransmite señalización. La prueba real entre dispositivos y redes es indispensable para certificar la experiencia."),
    ("11. Perfil, edición y preferencias", "Perfil muestra identidad, rol, organización, unidad y datos de contacto. La edición se realiza en una pantalla separada con validación y feedback. También se ofrecen acceso legal, seguridad y cierre de sesión.", ["Avatar y datos normalizados.", "Formulario de edición.", "Mensajes de guardado y error.", "Preferencias de cuenta.", "Acciones de sesión y seguridad."], "Los cambios se reflejan en chat, incidencias y operación sin mantener copias independientes. Cerrar sesión y descartar cambios requieren confirmaciones consistentes."),
    ("12. Estados offline y recuperación visual", "La interfaz móvil no oculta los problemas de conectividad. ConnectionBanner informa offline, reconexión y sincronización; las pantallas preservan datos útiles y distinguen información local de confirmada.", ["Banners no bloqueantes.", "Mensajes pendientes en chat.", "Última ubicación válida con hora.", "Reintento de acciones.", "Protección contra respuestas de sesiones anteriores."], "Cada módulo aplica una política distinta. El mapa puede conservar una posición histórica; chat puede guardar un pendiente; una incidencia crítica debe exigir confirmación del servidor."),
    ("13. Ventas pública y presentación de planes", "La landing de Ventas presenta el producto, beneficios y cinco planes obtenidos del backend. El diseño se refinó para responder en escritorio y móvil, mostrar carga real, diferenciar error de catálogo vacío y reducir movimiento cuando el sistema lo solicita.", ["Hero y propuesta de valor.", "Tarjetas de planes reales.", "Addon de Radio cuando corresponde.", "Estados error, reintento y vacío.", "Navegación hacia registro y checkout."], "El plan seleccionado conserva planId a través de todo el recorrido. La UI no duplica precios ni muestra planes ficticios cuando falla la API."),
    ("14. Checkout y confirmación de compra", "El checkout guía datos de empresa, contacto, facturación y selección. La interfaz valida antes de enviar, bloquea dobles clics y presenta progreso mientras backend crea la orden y Mercado Pago prepara la preferencia.", ["Resumen de plan y addons.", "Validaciones de campos.", "Estados de envío y error.", "Redirección segura al proveedor.", "Confirmación y acceso posterior."], "El importe visible es informativo; backend recalcula y valida. Los errores 500 o fallos de red deben permitir reintentar sin borrar datos ingresados."),
    ("15. Portal: layout y navegación", "PortalLayout organiza una aplicación administrativa completa dentro de Ventas. El sidebar separa Cuenta, Administración y Ayuda, y aplica permisos por rol. En pantallas pequeñas, la navegación se adapta sin perder acceso a acciones.", ["Protección de ruta y redirección a login.", "Sidebar y cabecera responsive.", "Control de permisos.", "Error boundary y toast global.", "Persistencia de contexto de cuenta."], "El layout es la capa común para todas las pantallas. Un error de render se aísla mediante boundary, aunque cada pantalla también debe manejar errores de datos sin convertirlos en estados vacíos falsos."),
    ("16. Portal: Operaciones y dashboard", "Operaciones reúne métricas, unidades, jornadas, eventos, historial y mapa. Es la superficie web más compleja y fue refinada repetidamente para mejorar densidad, jerarquía y acciones rápidas.", ["Tarjetas KPI.", "Filtros por unidad, conductor, ruta y estado.", "Detalle y replay de jornada.", "Lista operativa y acciones.", "Mapa conectado a posiciones reales."], "La pantalla consume endpoints de tracking e historial. No usa fixtures y debe mostrar etiquetas comprensibles en lugar de UUID. Las extracciones de botones y listas redujeron repetición."),
    ("17. Portal: mapa, polilíneas y rutas", "El mapa del Portal muestra unidades y geometría de recorridos. Las polilíneas atravesaron varias correcciones porque podían no cargarse, no coincidir con la selección o depender de configuración incompleta.", ["OperationsMap y thumbnail de geometría.", "Fallback cuando Mapbox no está disponible.", "Polilíneas de rutas y sesiones.", "Frescura GPS compartida.", "Filtros que no truncan resultados."], "La UI conserva la última ubicación real y evita mostrar variables técnicas. Cuando la geometría no existe o falla, explica el estado sin simular una ruta."),
    ("18. Portal: Equipo, Unidades y Rutas", "Estas pantallas permiten administrar la estructura de la operación. Las listas se migraron a primitivas compartidas, los botones a PortalButton y las acciones destructivas usan confirmación.", ["Usuarios y estados.", "Unidades, placa, conductor y condición.", "Asignación y liberación de rutas.", "Formularios y modales.", "Loading, vacío y error."], "Los datos alimentan Operaciones y Mobile. La interfaz debe impedir que editar presentación altere silenciosamente una asignación válida."),
    ("19. Portal: plan, pagos, facturación y perfil", "Las pantallas de cuenta muestran suscripción, importe, periodo, facturas, orden, seguridad y soporte. La UI fue ajustada para eliminar promesas de renovación o reactivación que no existían en backend.", ["Mi Plan y acciones reales.", "Pagos y reintento.", "Facturas y descarga.", "Empresa y sesiones activas.", "Soporte y seguridad."], "El error de carga no debe presentarse como ausencia de suscripción. La navegación mantiene al cliente dentro del Portal y abre checkout solo cuando existe un plan válido."),
    ("20. Responsive, accesibilidad y movimiento", "El frontend se afinó para escritorio, móvil web y dispositivos Android. La accesibilidad considera foco, teclado, lectores de pantalla, contraste, objetivos táctiles, anuncios de estado y reducción de movimiento.", ["Layouts adaptativos.", "Foco visible y orden lógico.", "Etiquetas accesibles para iconos.", "No depender solo de color.", "prefers-reduced-motion y reduceMotionChanged."], "La accesibilidad se aplica a componentes compartidos para evitar corregir pantalla por pantalla. Los mapas deben tener alternativa textual y los estados dinámicos no deben mover el foco inesperadamente."),
    ("21. Calidad frontend y verificación runtime", "La calidad del frontend combina typecheck, pruebas unitarias, build, smokes y validación manual. Los errores de runtime demostraron que compilar no es suficiente: una referencia de estilos o evaluación prematura puede derribar una pantalla completa.", ["TypeScript en Mobile y Ventas.", "Jest para navegación, mapa, radio y utilidades.", "Build Vite.", "Android assembleDebug.", "Carga manual de pantallas críticas."], "El criterio correcto es verificar contenido real, interacciones y estados de error. Portal requiere ampliar pruebas automatizadas de UI; Mobile requiere pruebas físicas para GPS, radio, teclado y llamadas."),
]


ERRORS = [
    ("Rastreo sin conexión o sin actualización", "La ubicación no se reflejaba correctamente entre Mobile y Portal, o se mostraba como si fuera actual aunque el dato estuviera vencido.", "Se unificó la frescura GPS, el orden temporal y la recuperación. El frontend ahora distingue actual, demorada y sin datos recientes."),
    ("Radio PTT con fallas de conexión o estado", "La transmisión podía quedar en una fase incorrecta, competir por el canal o no comunicar con claridad si el audio se envió.", "Se consolidó el reducer, lifecycle, arbitraje de canal y mensajes de error; la prueba en dispositivo sigue siendo obligatoria."),
    ("Chat cubierto por el teclado", "Al escribir, el teclado podía ocultar el compositor o los mensajes recientes, afectando respuesta y adjuntos.", "Se reforzó el layout seguro de teclado, el scroll y la separación de ChatComposer y ChatScreenView."),
    ("Polilíneas de Ventas/Portal no cargaban", "La geometría podía faltar, no corresponder a la unidad seleccionada o depender de Mapbox sin fallback suficiente.", "Se corrigió la selección, el consumo de geometría y el fallback con datos reales, sin inventar recorridos."),
    ("Error 500 en consultas o acciones", "Una solicitud del frontend podía recibir un error interno y dejar la pantalla vacía o mostrar un mensaje técnico.", "Se normalizaron errores, se agregaron estados de reintento y se evitó confundir fallo de API con ausencia real de datos."),
]


def configure(doc):
    report.configure(doc)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[name].font.color.rgb = DARK_RED if name == "Heading 1" else RED
    for section in doc.sections:
        for run in section.header.paragraphs[0].runs:
            run.font.color.rgb = DARK_RED


def set_fill(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def apply_red_theme(doc):
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            for run in paragraph.runs:
                run.font.color.rgb = DARK_RED if paragraph.style.name == "Heading 1" else RED
    for tbl in doc.tables:
        is_header = tbl.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None
        if is_header:
            for cell in tbl.rows[0].cells:
                set_fill(cell, "8B0000")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
            for index, row in enumerate(tbl.rows[1:], start=1):
                for cell in row.cells:
                    set_fill(cell, ALT_RED if index % 2 == 0 else "FFFFFF")
        else:
            for row in tbl.rows:
                for cell in row.cells:
                    set_fill(cell, LIGHT_RED)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.bold:
                                run.font.color.rgb = DARK_RED


def cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report.font(para.add_run("MANECOMB"), size=20, bold=True, color=DARK_RED)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report.font(para.add_run("Reporte de desarrollo frontend"), size=28, bold=True, color=DARK_RED)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report.font(para.add_run("Aplicación móvil, Ventas y Portal operativo"), size=16, italic=True, color=RED)
    for _ in range(7):
        doc.add_paragraph()
    table(doc, ["Campo", "Detalle"], [["Enfoque", "Arquitectura visual, pantallas, interacción, estados, responsive y accesibilidad"], ["Evidencia final", "Espacios reservados para capturas principales y errores representativos"], ["Estilo", "Edición roja"], ["Fecha", "21/07/2026"]], [2200, 7160], size=10)
    next_page(doc)


def toc(doc):
    h(doc, "Contenido", 1)
    report.base.add_toc(doc.add_paragraph())
    p(doc, "El índice es automático. En Word, utilice Actualizar tabla para recalcular los números de página después de insertar las capturas.")
    callout(doc, "Estructura", "La primera parte documenta el frontend. Las últimas páginas están reservadas para capturas de vistas principales y evidencias de cinco errores seleccionados.")
    next_page(doc)


def introduction(doc):
    h(doc, "Resumen ejecutivo frontend", 1)
    p(doc, "El frontend de ManeComb traduce contratos técnicos y datos operativos en dos experiencias coordinadas. La app móvil acompaña al conductor y al personal de campo; Ventas atrae, registra y conduce al pago; y el Portal permite administrar la organización y observar la operación. El desarrollo no se limitó a cambiar colores o tarjetas: incluyó navegación, manejo de sesión, estados offline, mapas, geometría, teclado, audio, responsive, accesibilidad y recuperación ante errores.")
    table(doc, ["Superficie", "Usuarios", "Funciones principales"], [["App móvil", "Conductores, supervisores y operación", "Mapa, jornada, incidencias, chat, Radio PTT, llamadas y perfil."], ["Ventas pública", "Compradores y empresas", "Planes, registro, checkout, pago y descarga."], ["Portal", "Administradores y propietarios", "Operaciones, Equipo, Unidades, Rutas, pagos, plan y cuenta."]], [1900, 2800, 4660], size=9)
    p(doc, "El reporte pone más énfasis en lógica de interfaz: cómo se decide qué renderizar, qué ocurre durante una carga, cómo se comunica una desconexión, cómo se conserva contexto al navegar y cómo se integran acciones con el backend. Al final se incluyen espacios amplios para capturas aportadas por el equipo.")
    next_page(doc)


def contents_map(doc):
    h(doc, "Mapa de vistas y relación entre superficies", 1)
    table(doc, ["Flujo", "App", "Ventas / Portal", "Conexión"], [["Acceso", "Login y gate operativo", "Login comprador y protección Portal", "Sesión, rol y suscripción."], ["Operación", "Mapa, Control e Incidencias", "Dashboard, mapa, unidades y rutas", "Tracking, jornadas y eventos."], ["Comunicación", "Chat, Radio PTT y llamadas", "Alertas y contexto administrativo", "Socket.IO, presencia y notificaciones."], ["Cuenta", "Perfil y edición", "Empresa, seguridad, plan y pagos", "Usuario, organización y permisos."], ["Distribución", "UpdateBanner", "App Center y descarga", "Versión configurada en backend."]], [1600, 2500, 2860, 2400], size=8)
    p(doc, "La continuidad entre superficies es un criterio central. Una unidad creada en Portal debe aparecer en operación; una ruta asignada debe reflejarse en Mobile; una incidencia enviada desde campo debe llegar al tablero; y una actualización publicada desde App Center debe mostrarse al usuario móvil.")
    callout(doc, "Principio", "Cada pantalla presenta una parte del mismo sistema. El frontend no debe crear datos paralelos solo para completar una maqueta.")
    next_page(doc)


def topic_page(doc, item, index):
    title, overview, points, integration = item
    h(doc, title, 1)
    p(doc, overview)
    h(doc, "Elementos desarrollados", 2)
    bullets(doc, points)
    h(doc, "Comportamiento e integración", 2)
    p(doc, integration)
    table(doc, ["Aspecto", "Criterio frontend"], [["Carga", "Mostrar progreso sin sustituir datos válidos por un vacío falso."], ["Error", "Explicar qué ocurrió y ofrecer reintento o recuperación."], ["Vacío", "Usarlo solo cuando la consulta fue correcta y realmente no existen registros."], ["Runtime", "Cargar la vista y verificar contenido e interacción, no solo DOM o compilación."], ["Responsive", "Conservar jerarquía y acciones sin desbordes ni controles ocultos."]], [2100, 7260], size=8)
    callout(doc, "Resultado", "La vista debe conservar contexto, comunicar estado y depender de contratos reales del sistema.")
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    report.font(para.add_run(f"Sección frontend {index} de {len(TOPICS)}"), size=9, italic=True, color=GRAY)
    next_page(doc)


def error_summary(doc):
    h(doc, "Errores representativos seleccionados", 1)
    p(doc, "El desarrollo tuvo múltiples ajustes, pero este reporte limita la evidencia de errores a cinco casos visuales representativos. La intención es mostrar evolución y corrección sin convertir el documento en un inventario de fallas.")
    table(doc, ["Caso", "Síntoma observado", "Respuesta aplicada"], [[name, symptom, fix] for name, symptom, fix in ERRORS], [2200, 3540, 3620], size=8)
    callout(doc, "Uso de capturas", "En las páginas finales coloque una imagen del problema y, cuando exista, otra del resultado corregido. Evite incluir tokens, correos, teléfonos, claves o datos productivos.")
    next_page(doc)


def placeholder(doc, title, caption, notes, full=False):
    h(doc, title, 1)
    p(doc, caption)
    box = doc.add_table(rows=1, cols=1)
    report.base.table_geometry(box, [9360])
    cell = box.cell(0, 0)
    set_fill(cell, "FFF9F9")
    first = cell.paragraphs[0]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first.paragraph_format.space_after = Pt(10)
    report.font(first.add_run("INSERTAR CAPTURA AQUÍ"), size=14, bold=True, color=RED)
    blank_count = 20 if full else 12
    for _ in range(blank_count):
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        report.font(paragraph.add_run(" "), size=11)
    footer = cell.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report.font(footer.add_run("Espacio reservado para imagen"), size=10, italic=True, color=GRAY)
    doc.add_paragraph()
    table(doc, ["Pie de figura", "Notas"], [[notes, "Fecha / versión: ____________________\nResultado observado: ______________________________"]], [3000, 6360], size=9)
    next_page(doc)


def build():
    doc = Document()
    configure(doc)
    cover(doc)
    toc(doc)
    introduction(doc)
    contents_map(doc)
    for index, item in enumerate(TOPICS, start=1):
        topic_page(doc, item, index)
    error_summary(doc)
    placeholder(doc, "Anexo A. Vistas principales de la app", "Inserte una composición con Inicio/Control, Mapa y Perfil.", "Figura A1. Vistas principales de la aplicación móvil ManeComb.", full=True)
    placeholder(doc, "Anexo B. Mapa y seguimiento móvil", "Inserte capturas del mapa, polilínea, unidad seleccionada y panel de seguimiento.", "Figura B1. Seguimiento de unidad y ruta en la app.", full=True)
    placeholder(doc, "Anexo C. Chat y Radio PTT", "Inserte las vistas principales de conversación, compositor, estados de mensaje y botón PTT.", "Figura C1. Comunicación mediante chat y Radio PTT.", full=True)
    placeholder(doc, "Anexo D. Ventas pública", "Inserte landing, tarjetas de planes y checkout.", "Figura D1. Experiencia comercial de ManeComb.", full=True)
    placeholder(doc, "Anexo E. Portal operativo", "Inserte dashboard, mapa, Equipo, Unidades o Rutas.", "Figura E1. Vistas principales del Portal web.", full=True)
    placeholder(doc, "Anexo F. Evidencia: rastreo y polilíneas", "Inserte una captura del rastreo sin conexión y/o polilínea que no cargaba; puede comparar antes y después.", "Figura F1. Falla y corrección de rastreo o geometría.", full=True)
    placeholder(doc, "Anexo G. Evidencia: Radio y Chat", "Inserte capturas de una falla de Radio PTT y del teclado cubriendo el chat.", "Figura G1. Errores representativos de comunicación móvil.", full=True)
    placeholder(doc, "Anexo H. Evidencia: error 500 en Ventas/Portal", "Inserte una captura del error 500 o de una pantalla que no cargaba, junto con el estado corregido si está disponible.", "Figura H1. Manejo de error de servidor en el frontend web.", full=True)
    h(doc, "Conclusión frontend", 1)
    p(doc, "El frontend de ManeComb evolucionó mediante integración real, refactor de componentes y correcciones basadas en comportamiento. La app móvil concentra interacción operativa sensible a red, ubicación, teclado y audio; Ventas y Portal convierten el producto en una experiencia comercial y administrativa coherente. Los errores seleccionados muestran que la calidad no depende únicamente de compilar: requiere probar mapas, geometría, sockets, teclado, audio y respuestas del servidor en ejecución.")
    p(doc, "Los anexos permiten completar el reporte con evidencia visual proporcionada por el equipo. Al insertar las imágenes, se recomienda mantener un tamaño consistente, recortar datos sensibles y actualizar el índice para conservar la numeración.")
    apply_red_theme(doc)
    doc.core_properties.title = "Reporte frontend ManeComb - App, Ventas y Portal"
    doc.core_properties.subject = "Desarrollo de interfaces y espacios para evidencia visual"
    doc.core_properties.author = "Proyecto ManeComb"
    doc.save(OUT)
    report.base.update_fields(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
