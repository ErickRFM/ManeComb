from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_actividad_usabilidad_manecomb as base


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "reporte-integral-desarrollo-manecomb-app-ventas-backend-frontend.docx"
FONT = "Calibri"
NAVY = RGBColor(31, 55, 78)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(89, 89, 89)


def font(run, size=11, bold=False, italic=False, color=RGBColor(0, 0, 0)):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


base.font = font


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)
    for name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, NAVY),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_after = Pt(4)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("ManeComb | Reporte integral de desarrollo"), size=9, color=GRAY)
    base.add_page_number(section.footer.paragraphs[0])


def p(doc, text, lead=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.line_spacing = 1.25
    para.paragraph_format.space_after = Pt(6)
    if lead and text.startswith(lead):
        font(para.add_run(lead), bold=True)
        font(para.add_run(text[len(lead):]))
    else:
        font(para.add_run(text))
    return para


def h(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    para.paragraph_format.keep_with_next = True
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.keep_together = True
        font(para.add_run(item))


def table(doc, headers, rows, widths, size=8):
    return base.add_table(doc, headers, rows, widths, size=size)


def callout(doc, label, text, fill="EEF3F8"):
    base.callout(doc, label, text, fill=fill)


def next_page(doc):
    doc.add_page_break()


TOPICS = [
    {
        "title": "1. Origen, propósito y alcance del producto",
        "overview": "ManeComb nació como una plataforma para concentrar la operación diaria de flotillas de transporte colectivo. El problema inicial combinaba información dispersa, seguimiento manual de unidades, comunicación por canales no integrados y poca trazabilidad sobre rutas, incidencias y decisiones. La solución evolucionó hacia un monorepo que reúne API, aplicación móvil, experiencia comercial, portal administrativo y servicios de comunicación.",
        "items": ["Centralización de usuarios, roles, unidades, rutas y operación.", "Seguimiento geográfico y temporal conectado con incidencias y comunicación.", "Canal comercial para compra, activación, facturación y acceso al Portal.", "Aplicación móvil para conductores y operación en campo.", "Backend unificado para contratos REST, Socket.IO, persistencia y seguridad."],
        "integration": "El alcance no se limita a una aplicación de mapa ni a una tienda de planes. Ventas inicia la relación con la organización; el backend convierte la compra en orden, suscripción y activación; el Portal permite administrar la cuenta y la flotilla; y la app móvil ejecuta la operación diaria. Esta cadena explica por qué los cambios comerciales, de autorización y de seguimiento deben mantenerse sincronizados.",
        "evidence": "docs/alcance-sistema-combis.md; docs/project-master.md; historial Git desde el baseline del 31/05/2026",
        "status": "Producto integrado con frentes comercial, administrativo y operativo activos.",
    },
    {
        "title": "2. Consolidación del monorepo y etapas de evolución",
        "overview": "El repositorio se estableció formalmente como monorepo y después atravesó ciclos de estabilización, migración móvil, rediseño comercial, endurecimiento productivo y certificaciones por dominio. El historial visible contiene 131 revisiones hasta el corte analizado, con una concentración significativa de trabajo durante junio y julio de 2026.",
        "items": ["Baseline y eliminación inicial de código confirmado como muerto.", "Migración del cliente móvil hacia React Native CLI.", "Correcciones de CORS, variables de entorno y despliegue web.", "Rediseño de Ventas y compra de planes SaaS.", "Ciclos RC para comunicación, seguimiento, portal, autorización y release."],
        "integration": "La evolución fue incremental: primero se construyó la base funcional; después se estabilizaron accesos y producción; luego se profundizó en consistencia operacional, comunicación y experiencia del Portal. Los documentos RC registran hallazgos intermedios que no deben confundirse con el estado final: por ejemplo, una auditoría detectó pantallas faltantes y certificaciones posteriores documentaron su integración.",
        "evidence": "git log; RC-RELEASE-CANDIDATE-01.md; RC-PORTAL-FINAL-01.md; múltiples RC de junio-julio",
        "status": "Evolución trazable por commits y reportes de certificación; quedan validaciones manuales puntuales.",
    },
    {
        "title": "3. Arquitectura general y separación de responsabilidades",
        "overview": "La arquitectura actual separa interfaces por contexto sin duplicar la lógica central. Mobile atiende operación en campo; Ventas contiene landing, checkout y Portal; Backend expone la API principal y el socket general; communication-service concentra correo y mensajería saliente tolerante a fallos. MongoDB actúa como persistencia productiva y existe un store embebido para pruebas y desarrollo controlado.",
        "items": ["REST para operaciones transaccionales y consultas.", "Socket.IO para ubicación, presencia, chat, radio, llamadas e incidencias.", "Zustand para estado de cliente en Mobile y Portal.", "Adaptadores comerciales para separar UI, reglas y proveedores.", "Servicios compartidos de seguridad, telemetría, archivos y notificaciones."],
        "integration": "La cadena principal es UI -> store/hook -> cliente API -> ruta backend -> store -> respuesta normalizada -> render. En eventos dinámicos, el backend persiste o valida y después emite a salas segmentadas. Esta estructura reduce fuentes paralelas: Portal y Mobile consumen los mismos datos operativos, y el mapa no mantiene un seguimiento independiente del backend.",
        "evidence": "backend/src/app.js; backend/src/server.js; ventas/features; mobile/src; communication-service/docs/ARCHITECTURE.md",
        "status": "Arquitectura modular con backend unificado y servicios auxiliares acotados.",
    },
    {
        "title": "4. Fundación del backend Express",
        "overview": "El backend está construido sobre Node.js y Express, con montaje modular de rutas, middlewares transversales y un servidor HTTP compartido con Socket.IO. La base incluye compresión, CORS configurable, Helmet, limitación de solicitudes, trazabilidad por identificador y manejo público de errores.",
        "items": ["Inicialización de configuración y comprobaciones de preparación.", "Montaje de módulos bajo /api con contratos normalizados.", "Autenticación y autorización mediante middleware.", "Manejo centralizado de 404 y errores no controlados.", "Integración del store embebido o Mongo según disponibilidad y política."],
        "integration": "La decisión de mantener un backend principal simplifica autenticación, permisos y tiempo real. Chat, GPS, radio y señalización RTC comparten la sesión del socket; las funciones comerciales reutilizan el mismo store y telemetría. La política REQUIRE_MONGO permite impedir degradaciones silenciosas en producción.",
        "evidence": "backend/src/app.js; backend/src/server.js; backend/src/config; backend/src/middlewares",
        "status": "Base productiva endurecida; la disponibilidad final depende de variables y servicios externos.",
    },
    {
        "title": "5. Autenticación, sesión y recuperación de acceso",
        "overview": "El módulo de autenticación cubre inicio de sesión, registro permitido por flujo, consulta de sesión, refresh y recuperación de contraseña. El cliente conserva sesión de forma segura y ejecuta inicialización, renovación ante 401 y limpieza coordinada cuando la sesión deja de ser válida.",
        "items": ["JWT para identidad y autorización de solicitudes.", "Hash de contraseñas con bcrypt y política de fortaleza.", "Promesa singleton de refresh para evitar carreras concurrentes.", "Ruteo por estado de cuenta, rol y suscripción.", "Endurecimiento específico del flujo de recuperación de contraseña."],
        "integration": "La autenticación conecta el resultado comercial con la operación: no basta con credenciales válidas; el acceso móvil también considera activación y suscripción. En Portal, el layout protege rutas y aplica permisos antes de renderizar consumidores. El respaldo de claves E2EE se vincula al perfil autenticado sin exponerse en respuestas generales.",
        "evidence": "backend/src/modules/auth; backend/src/utils/jwt.js; mobile/src/store; RC-AUTHORIZATION-FINAL-01.md; commit 8563da2",
        "status": "Flujo funcional con refresh, políticas y recuperación endurecida.",
    },
    {
        "title": "6. Persistencia, modelos y estrategia de store",
        "overview": "La capa de datos abstrae dos implementaciones: store embebido para desarrollo y pruebas, y MongoDB/Mongoose para persistencia productiva. Los módulos consumen métodos del store mediante req.app.locals.store, evitando que las rutas acoplen directamente su lógica a una base específica.",
        "items": ["Modelos para usuarios, vehículos, rutas, incidencias y conversaciones.", "Persistencia de órdenes, suscripciones, sesiones RTC y eventos.", "GridFS para binarios documentales cuando Mongo está disponible.", "Semillas controladas para entorno local.", "Reglas de integridad y limpieza de asociaciones al eliminar entidades."],
        "integration": "La misma interfaz de store permite ejecutar pruebas sin infraestructura externa y desplegar con Mongo sin reescribir módulos. No obstante, los datos productivos requieren auditorías específicas: el informe de tracking dejó pendiente contar rutas huérfanas porque el entorno no resolvió DNS, y se evitó inferir resultados inexistentes.",
        "evidence": "backend/src/data/models.js; store.js; mongo-store.js; seedData.js; RC-DATA-INTEGRITY-01.md",
        "status": "Persistencia dual consolidada; producción debe exigir Mongo y auditar datos heredados.",
    },
    {
        "title": "7. Usuarios, roles y autorización RBAC",
        "overview": "La gestión de usuarios evolucionó desde cuentas administrativas generales hasta reglas más precisas por rol, organización y permiso. Administradores operan el Portal; supervisores consultan operación; conductores acceden a funciones móviles; compradores gestionan cuenta, plan y activación.",
        "items": ["Listado, actualización y eliminación con validaciones de integridad.", "Normalización de roles y sanitización de datos sensibles.", "Permisos de Portal para users, vehicles, routes y billing.", "Segmentación por organización en consultas y sockets.", "Alta de conductores mediante activation keys, no creación directa desde Portal."],
        "integration": "RBAC se aplica en backend y se refleja en navegación. La interfaz no sustituye la autorización del servidor: ocultar una opción mejora la experiencia, pero cada endpoint conserva validación. La eliminación de usuarios limpia o protege asociaciones relevantes para evitar conversaciones, unidades o documentos incongruentes.",
        "evidence": "backend/src/modules/users; require-admin.js; ventas/features/portal/utils/access.ts; RC-RBAC-CONTROL-01.md",
        "status": "Autorización integrada por rol y organización, con controles tanto en API como en Portal.",
    },
    {
        "title": "8. Unidades, conductores y consistencia operativa",
        "overview": "El dominio de unidades reúne código, placa, conductor, ruta, estado, capacidad, kilometraje y datos operativos. La evolución reciente se concentró en evitar duplicados, conservar asignaciones y alinear la representación de una misma unidad entre Mobile, Portal y backend.",
        "items": ["CRUD y consulta de vehículos mediante API.", "Asignación de conductor y ruta con reglas de integridad.", "Prevención de creación duplicada y normalización de identificadores.", "Presentación de código, conductor y ruta en lugar de UUID internos.", "Consistencia del vehículo activo en la aplicación móvil."],
        "integration": "Las unidades actúan como nodo común: reciben ubicación, participan en jornadas, originan alertas, se asocian a incidencias y aparecen en filtros del mapa. Un hallazgo RC detectó que editar una unidad assigned podía convertirla en available; los ciclos posteriores de consistencia y certificación se orientaron a impedir pérdidas silenciosas de asignación.",
        "evidence": "backend/src/modules/vehicles; portal-units-screen.tsx; RC-UNITS-DUPLICATE-CREATE-01.md; RC-MOBILE-VEHICLE-DATA-CONSISTENCY-01.md",
        "status": "Gestión consolidada con atención especial a integridad de asignaciones.",
    },
    {
        "title": "9. Rutas, asignación y geometría",
        "overview": "El sistema administra rutas con nombre, código, color, origen, destino y geometría. La asignación vincula rutas con unidades y alimenta tanto la navegación móvil como la vista operacional del Portal.",
        "items": ["Creación y edición de rutas con coordenadas validadas.", "Asignación, liberación y duplicación controlada.", "Polilíneas en mapas y miniaturas geométricas.", "Historial de viajes y sesiones por unidad.", "Aislamiento de rutas por organización."],
        "integration": "La ruta no es solo una figura en el mapa: define contexto para ubicación, jornada, incidentes, checkpoints y métricas. Las mejoras de UX agregaron confirmaciones y evitaron sobrescrituras poco visibles. En backend, crear rutas exige organización para impedir que registros sin tenant se filtren como globales.",
        "evidence": "backend/src/modules/navigation; mobile/src/screens/map; portal-routes-screen.tsx; RC-RUTAS-UX-03.md",
        "status": "Flujo funcional de definición, asignación y visualización con geometría compartida.",
    },
    {
        "title": "10. Seguimiento GPS e integridad temporal",
        "overview": "El seguimiento fue uno de los frentes con mayor número de iteraciones. La implementación unificó ubicación enviada por HTTP y Socket, orden temporal, frescura, recuperación y presentación coherente en Mobile y Portal.",
        "items": ["Normalización de reloj cliente con tolerancia configurable.", "Timestamps de origen, recepción y procesamiento para auditoría.", "Prevención de que paquetes antiguos sobrescriban ubicaciones nuevas.", "gpsFreshness autoritativo compartido por clientes.", "Segmentación de emisiones por organización, rol y usuario."],
        "integration": "El backend calcula freshUntil y los clientes lo interpretan sin inventar umbrales paralelos. Una ubicación vencida puede conservarse como evidencia histórica, pero no debe mostrarse como posición actual ni adjuntarse a una incidencia como dato fresco. El Portal y Mobile consumen la misma decisión para evitar contradicciones.",
        "evidence": "backend/src/services/tracking-time.js; locations/routes.js; mobile/src/screens/map/utils/tracking.ts; RC-TRACKING-EXECUTION-01.md",
        "status": "Fases técnicas completadas; persisten conteo productivo y matriz manual en dispositivo.",
    },
    {
        "title": "11. Mapa móvil y experiencia de seguimiento",
        "overview": "La pantalla de mapa móvil se descompuso en canvas, controles flotantes, HUD, selector de ruta, panel inferior y hooks especializados. Esta separación redujo la concentración de lógica y facilitó pruebas sobre selección, cámara, ubicación y estados operativos.",
        "items": ["MapCanvas y variantes nativa, web y de tipos compartidos.", "Motor de ubicación y sincronización con reducers y servicios.", "BottomTrackingPanel con información de unidad y jornada.", "Recuperación cuando el mapa o datos no están disponibles.", "Preferencia reduceMotion para omitir animaciones de layout."],
        "integration": "El mapa combina ruta asignada, posición, jornada, checkpoints e incidentes. La selección debe sobrevivir actualizaciones y cambios de cámara. Las polilíneas se mejoraron en commits recientes, y la vista evita representar como movimiento un dato que no ha sido actualizado.",
        "evidence": "mobile/src/screens/map; app-map.native.tsx; commits de polilínea y seguimiento; RC-GEOLOCATION-FORENSIC-AUDIT-01.md",
        "status": "Arquitectura modular y pruebas focalizadas; certificación física sigue siendo indispensable.",
    },
    {
        "title": "12. Jornadas, control, checklist e historial",
        "overview": "La pantalla de control/checklist concentra inicio, pausa, reanudación y cierre de jornada, selección de rutas, puntos de control e historial. El módulo reutiliza contratos de navegación y sesiones, no un backend de checklist completamente separado.",
        "items": ["Acciones de sesión encapsuladas en route-session-actions.", "Cola o recuperación para operaciones con conectividad limitada.", "Creación y eliminación de rutas guardadas.", "Historial y checkpoints conectados con seguimiento.", "Mensajes y confirmaciones para acciones operativas."],
        "integration": "Una auditoría RC señaló que no existe una entidad backend checklist autónoma. El diseño actual debe describirse con honestidad: la pantalla funciona sobre navegación, rutas y sesiones. Esto reduce duplicación, pero crea dependencia del modelo operacional y exige pruebas de regresión cuando cambian sus contratos.",
        "evidence": "mobile/src/screens/checklist-screen.tsx; mobile/src/services/route-session-actions.ts; RC-CONTROL-CERTIFICATION-01.md",
        "status": "Flujo operativo amplio, con deuda de definición sobre si checklist debe convertirse en dominio propio.",
    },
    {
        "title": "13. Incidencias y flujo SOS",
        "overview": "El módulo de incidencias permite registrar problemas, asociarlos con ruta y unidad, definir severidad y evolucionar el estado de abierta a en proceso o resuelta. Los eventos relevantes se distribuyen en tiempo real y alimentan dashboard y alertas.",
        "items": ["Creación y listado con permisos por rol y organización.", "Estados y severidades normalizados.", "Evidencias y ubicación condicionada por frescura GPS.", "Notificación de eventos críticos y SOS.", "Pantallas Mobile y Portal para alta, consulta y seguimiento."],
        "integration": "El flujo conecta ubicación, usuario reportero, unidad y ruta. Una auditoría detectó que el patrón /^sos/i podía generar falsos positivos con palabras no relacionadas; el reporte histórico conserva ese hallazgo como evidencia de por qué las reglas de emergencia deben ser exactas y probadas.",
        "evidence": "backend/src/modules/incidents; mobile/src/screens/incidents-screen.tsx; portal-incidents-screen.tsx; RC-RELEASE-CANDIDATE-01.md",
        "status": "Dominio conectado a tiempo real; reglas críticas deben conservar pruebas contra falsos positivos.",
    },
    {
        "title": "14. Alertas, notificaciones y retroalimentación",
        "overview": "ManeComb combina notificaciones persistidas, suscripciones push, banners de conexión, mensajes inline y confirmaciones. La auditoría de alertas mostró una base funcional, pero también inconsistencias entre Mobile y Ventas, componentes duplicados y acciones que necesitaban confirmación más clara.",
        "items": ["Notificaciones dirigidas por usuario o rol con nivel y categoría.", "Deep links hacia chat, radio e incidencias.", "ConnectionBanner para offline, reconexión y sincronización.", "ConfirmModal y Toast en Portal.", "Revisión de lenguaje técnico, acciones destructivas y estados de éxito."],
        "integration": "Las alertas no deben existir aisladas: deben abrir el objeto que las originó y conservar contexto. La auditoría RC-ALERTS-01 no aprobó la experiencia en su corte por confirmaciones faltantes y mensajes inconsistentes; los trabajos posteriores de UI y operaciones deben leerse como respuesta progresiva, no como negación del hallazgo.",
        "evidence": "backend/src/modules/notifications; mobile/src/components/connection-banner.tsx; RC-ALERTS-01.md",
        "status": "Infraestructura implementada; la coherencia de feedback requiere revisión continua por flujo.",
    },
    {
        "title": "15. Chat: conversaciones, mensajes y estados",
        "overview": "El chat evolucionó desde una pantalla concentrada hacia componentes, hooks y utilidades especializadas. Soporta directorio, conversaciones generales y directas, mensajes de texto, adjuntos, voz y estados locales de entrega.",
        "items": ["ChatHeader, ChatComposer, ChatScreenView y MessageMedia.", "Controlador de conversación y directorio desacoplados.", "Estados pending, sent, delivered, read y failed en la experiencia.", "Reintento de mensajes y continuidad ante desconexión.", "Acciones de llamada y radio integradas en el encabezado."],
        "integration": "Los mensajes se persisten mediante REST y se retransmiten por Socket.IO. El cliente usa identificadores locales para evitar duplicados y mantiene estados visibles. La refactorización también estandarizó acciones del encabezado y corrigió problemas de teclado, scroll y similitud de render.",
        "evidence": "backend/src/modules/chat; mobile/src/screens/chat; RC-CHAT-FINAL-02.md; RC-CHAT-MESSAGE-STATUS-UX-01.md",
        "status": "Chat modular, multimedia y conectado a presencia, llamadas y cifrado directo.",
    },
    {
        "title": "16. Cifrado E2EE y protección de conversaciones directas",
        "overview": "Las conversaciones directas incorporan cifrado de extremo a extremo con TweetNaCl. Las claves y respaldos se asocian a la sesión del usuario, mientras el backend conserva contenido cifrado y metadatos necesarios para entrega.",
        "items": ["Generación y manejo de claves en cliente.", "Cifrado y descifrado para conversaciones directas.", "Respaldo E2EE protegido por endpoints de autenticación.", "Pruebas unitarias de utilidades criptográficas.", "Sanitización para evitar exponer backups sensibles."],
        "integration": "E2EE convive con estados, multimedia y persistencia, pero limita qué puede inspeccionar el servidor. La documentación distingue privacidad de transporte y autorización: cifrar contenido no sustituye validar participantes, organización, acceso a conversación ni seguridad del dispositivo.",
        "evidence": "mobile/src/utils/chat-e2ee.ts; backend/src/utils/chat-crypto.js; commit 2830ac4; pruebas chat-e2ee",
        "status": "Cifrado directo conectado; requiere mantener gestión de claves y recuperación bajo pruebas de regresión.",
    },
    {
        "title": "17. Multimedia, notas de voz y archivos de chat",
        "overview": "El chat admite audio, imagen y video mediante servicios de carga, reproducción y presentación específicos. La aplicación maneja selección de archivos, previsualización, errores de reproducción y visualización a pantalla completa.",
        "items": ["Endpoints de media y audio por conversación.", "Abstracciones nativas de selector, imagen, video y audio.", "MessageMedia para renderizar contenido según tipo.", "Transcripción opcional cuando existe proveedor configurado.", "Controles de tamaño, acceso y almacenamiento."],
        "integration": "La multimedia depende de storage y autorización de conversación. Los fallos de proveedor no deben derribar el flujo de texto. En Radio, los audios también se conservan como mensajes, pero una auditoría identificó una búsqueda O(N*M) que requiere paginación o acceso directo por identificador para escalar.",
        "evidence": "backend/src/services/chat-media.js; audio-transcription.js; mobile/src/screens/chat/components/message-media.tsx",
        "status": "Soporte multimedia implementado; rendimiento de recuperación de audios debe vigilarse.",
    },
    {
        "title": "18. Radio PTT y máquina de estado",
        "overview": "Radio PTT es un subsistema operativo separado de las llamadas. La experiencia móvil utiliza reducer de sesión, servicios de audio y tiempo real, waveform, tarjetas de transmisión y estados explícitos para captura, envío, reproducción y error.",
        "items": ["Mantener presionado para transmitir con retroalimentación háptica.", "Arbitraje de canal activo mediante Redis en la evolución reciente.", "Servicios independientes para audio y eventos realtime.", "Persistencia de transmisiones y reproducción desde historial.", "Auditorías específicas de SSOT, render, carreras y lifecycle."],
        "integration": "Radio comparte autenticación, conversación y socket, pero no comparte ciclo de vida ni foreground service con llamadas. Esta separación evita que un cambio de WebRTC afecte la reproducción PTT. Los reportes de radio documentan la consolidación de una sola fuente de verdad y la eliminación de estados duplicados.",
        "evidence": "mobile/src/screens/radio; backend/src/modules/radio; docs/radio-ssot-audit; commit 73275cc",
        "status": "Subsistema maduro y auditado, con dependencia operacional de audio, red y Redis según despliegue.",
    },
    {
        "title": "19. Presencia y estado de conexión",
        "overview": "La presencia informa si usuarios y dispositivos están conectados y evita inferencias inconsistentes en chat, radio y operación. Un ciclo específico consolidó connection state como fuente única de verdad.",
        "items": ["Eventos de conexión y desconexión por Socket.IO.", "Salas por usuario, rol y organización.", "Indicadores de presencia en UI.", "Manejo de reconexión y epoch de sesión.", "Pruebas de utilidades de presencia y realtime-state."],
        "integration": "La presencia no equivale a entrega de mensaje ni a ubicación fresca. El reporte separa estos conceptos: un usuario puede estar conectado sin GPS reciente, y un mensaje enviado al servidor puede no estar entregado al dispositivo. Mantener estados independientes evita promesas falsas en la UI.",
        "evidence": "mobile/src/components/presence-indicator.tsx; mobile/src/utils/presence.ts; commit 6188d6e; RC-PRESENCE-FINAL-01.md",
        "status": "Estado de conexión consolidado y reutilizado por comunicación.",
    },
    {
        "title": "20. Llamadas WebRTC y señalización",
        "overview": "Las llamadas 1 a 1 usan el socket general para señalización y WebRTC P2P para medios. STUN/TURN resuelve conectividad, mientras el backend valida acceso a la sala y conserva sesiones RTC como historial.",
        "items": ["Eventos join, leave, offer, answer, ICE y hangup.", "Ack de ingreso con busy y forbidden.", "Foreground service Android para micrófono y cámara.", "Timer de gracia de 15 segundos ante desconexión.", "Configuración TURN estática o coturn REST."],
        "integration": "La arquitectura rechazó un microservicio y socket adicional porque la escala actual no los justificaba. Chat es dueño de la conexión y llamadas la consume. El CDR reutiliza el store existente; las funciones grupales, grabación y SFU permanecen fuera del alcance actual.",
        "evidence": "docs/CALLING-ARCHITECTURE.md; backend/src/sockets/index.js; mobile/src/native/webrtc.ts; RC-WEBRTC-CERTIFICATION-01.md",
        "status": "Diseño proporcional y funcional; TURN y foreground requieren prueba entre dispositivos y redes reales.",
    },
    {
        "title": "21. Servicio de correo y comunicaciones salientes",
        "overview": "communication-service organiza correos transaccionales y otros canales salientes mediante proveedores intercambiables, plantillas, prioridades, colas, métricas e historial. El endurecimiento se enfocó en impedir que DNS, Redis o un proveedor externo derriben Node.js.",
        "items": ["Providers para Resend, SMTP, SES, Mailgun, Postmark y SendGrid.", "Resultado estructurado success/error en lugar de excepciones no capturadas.", "BullMQ cuando Redis está disponible y cola local como fallback.", "Validación de configuración antes de habilitar proveedor.", "Workers con reintentos y registro de fallo."],
        "integration": "Los flujos comerciales y de cuenta pueden solicitar correo sin asumir que el proveedor está disponible. Si la cola falla, el servicio intenta envío directo; si el proveedor no está configurado, devuelve error controlado. Esta resiliencia desacopla continuidad del backend de la disponibilidad de terceros.",
        "evidence": "communication-service; RC-COMMUNICATION-FINAL-01.md; 16 pruebas de comunicación documentadas",
        "status": "Módulo crash-proof certificado frente a errores de proveedor, red y cola.",
    },
    {
        "title": "22. Migración y arquitectura de la aplicación móvil",
        "overview": "El cliente móvil activo migró desde una etapa basada en Expo hacia React Native CLI. La documentación histórica conserva referencias antiguas, pero el stack de build y despliegue vigente se encuentra en mobile/ con Android nativo y abstracciones multiplataforma.",
        "items": ["React Native 0.81 y React 19.", "Navegación centralizada en router y registro de rutas.", "Módulos nativos para ubicación, audio, llamadas, imágenes y almacenamiento seguro.", "Android Gradle, servicios foreground y tareas headless.", "Compatibilidad web selectiva mediante archivos .web y .native."],
        "integration": "La migración obligó a revisar permisos, deep links, teclado, background, build y distribución. El Portal no reemplaza la app: administra la organización, mientras Mobile ejecuta rutas, ubicación, incidencias y comunicación en campo.",
        "evidence": "mobile/README.md; docs/mobile-release.md; commit 4435681; mobile/android",
        "status": "React Native CLI es la fuente vigente; referencias Expo se consideran históricas.",
    },
    {
        "title": "23. Navegación móvil, deep links y políticas de acceso",
        "overview": "La navegación móvil se endureció mediante registro de rutas, políticas de acceso, linking y pruebas específicas. El router decide entre autenticación, gate de cuenta y superficies operativas según sesión, rol y estado comercial.",
        "items": ["Route registry con pruebas de cobertura.", "Deep links hacia chat, radio e incidencias.", "Política de navegación independiente de la UI.", "Protección contra rutas inválidas y estados de bootstrap.", "Infraestructura de inputs y teclado segura."],
        "integration": "Las notificaciones pueden dirigir a una intención, pero el router vuelve a comprobar acceso y disponibilidad. La navegación nunca debe permitir que un deep link salte autorización. Los estados de cuenta provenientes del backend influyen en el gate y enlazan directamente Ventas con la app operativa.",
        "evidence": "mobile/src/navigation; deep-linking.test.ts; navigation-policy.test.ts; account-routing.ts",
        "status": "Navegación centralizada y probada contra rutas, sesión y cuenta.",
    },
    {
        "title": "24. Perfil móvil, edición y seguridad de cuenta",
        "overview": "El perfil móvil presenta identidad, rol, organización, unidad y preferencias, y ofrece edición de datos permitidos. También concentra cierre de sesión, información legal y acceso a configuración relacionada.",
        "items": ["Pantallas separadas de consulta y edición.", "Avatar e información normalizada del usuario.", "Mensajes de éxito y error al guardar.", "Almacenamiento seguro de sesión y claves.", "Acceso a privacidad, términos y datos de cuenta."],
        "integration": "Los datos de perfil se consumen en chat, incidencias, radio y asignación. Por ello, la actualización debe refrescar el store sin duplicar identidades. La auditoría de alertas señaló confirmaciones faltantes en cierre de sesión y pérdida de cambios; estas observaciones forman parte de la mejora continua de UX.",
        "evidence": "mobile/src/screens/profile-screen.tsx; profile-edit-screen.tsx; secure-store.ts; RC-ALERTS-01.md",
        "status": "Perfil operativo implementado; acciones destructivas deben conservar confirmación coherente.",
    },
    {
        "title": "25. Operación offline, reconexión y recuperación",
        "overview": "La aplicación incorpora caché offline, banners de conexión, reintentos y preservación de trabajo para escenarios móviles. El objetivo no es fingir que todo está actualizado, sino distinguir información local, pendiente y confirmada.",
        "items": ["offline-cache con pruebas unitarias.", "ConnectionBanner para offline y sincronización.", "Estados pendientes en chat y acciones de jornada.", "Conservación del último mapa válido con frescura visible.", "Session epoch para invalidar respuestas de sesiones anteriores."],
        "integration": "Cada dominio define su política: un borrador puede guardarse localmente; una ubicación antigua no puede presentarse como actual; una acción operativa pendiente requiere reconciliación; y un mensaje debe reintentarse sin duplicarse. Esta diferenciación protege integridad y confianza.",
        "evidence": "mobile/src/api/offline-cache.ts; connection-banner.tsx; session-epoch.ts; realtime-state.ts",
        "status": "Resiliencia transversal implementada con políticas específicas por dato.",
    },
    {
        "title": "26. Ventas: landing pública y descubrimiento comercial",
        "overview": "La experiencia pública de Ventas presenta ManeComb, planes y beneficios, y conduce al checkout. La evolución incluyó rediseño visual, corrección de CORS, URL de API, comportamiento sin conexión y despliegue en Cloudflare.",
        "items": ["Catálogo obtenido desde GET /api/commercial/plans.", "Cinco planes reales renderizados desde backend.", "Estados de carga, error y reintento separados del vacío real.", "Diseño responsive y movimiento reducido.", "Rutas limpias y shims compatibles con hosting estático."],
        "integration": "La landing no contiene precios paralelos: consume el catálogo central. Si la API falla, no muestra falsamente que no existen planes. La selección conserva planId hacia checkout, orden, proveedor de pago, suscripción y Portal.",
        "evidence": "ventas/src y features/commercial; RC-COMERCIAL-FINAL-01.md; docs/deploy-ventas-cloudflare.md",
        "status": "Landing conectada a catálogo productivo y preparada para error de API.",
    },
    {
        "title": "27. Catálogo de planes y motor comercial",
        "overview": "El backend mantiene un catálogo estático versionado de planes por tamaño de flotilla. El motor comercial separa reglas, tipos, contratos y adaptadores para que UI y proveedor de pago consuman la misma definición.",
        "items": ["Planes para 2, 4, 6, 8 y 12 combis.", "Addon de Radio según elegibilidad o inclusión.", "Snapshot del plan guardado con la orden.", "Validación de planId y total en backend.", "Cache de planes y adaptadores de experiencia."],
        "integration": "El catálogo es fuente única para nombres, precios, badges y características. No se agregó CRUD ni modelo Mongo porque no existía necesidad confirmada. El snapshot protege órdenes históricas ante cambios futuros del catálogo.",
        "evidence": "backend/src/config/commercial-plans.js; ventas/features/commercial; RC-VENTAS-PLANES-FINAL-01.md",
        "status": "Catálogo central certificado; descuentos no forman parte del contrato actual.",
    },
    {
        "title": "28. Checkout, validación y creación de orden",
        "overview": "El checkout recopila información de empresa, contacto y facturación, valida la selección y crea una orden comercial. La UI usa hooks de experiencia y servicios de validación, mientras el backend recalcula y conserva datos autoritativos.",
        "items": ["Validación de datos antes de enviar.", "Orden con planId, addons, total y snapshot.", "Estados de envío y recuperación de error.", "Separación de reglas comerciales y componentes visuales.", "Continuidad hacia Mercado Pago o instrucciones aplicables."],
        "integration": "El cliente nunca es autoridad del importe. El backend valida plan y addons, crea la orden y genera la preferencia. El mismo identificador se mantiene en metadata y external_reference para reconciliar confirmación y webhook.",
        "evidence": "ventas/features/commercial/hooks/use-checkout-experience.ts; checkout-validation.ts; backend/src/modules/commercial",
        "status": "Checkout conectado extremo a extremo con validación del servidor.",
    },
    {
        "title": "29. Integración con Mercado Pago e idempotencia",
        "overview": "La integración comercial crea preferencias de Mercado Pago, procesa confirmación y webhook, y activa solo pagos aprobados. Las pruebas cubren credenciales, selección de entorno, metadata, importes y referencias.",
        "items": ["Preferencia con external_reference y metadata coherentes.", "Separación sandbox/producción.", "Webhook y confirmación idempotentes.", "Addon incluido en el total calculado.", "Protección frente a credenciales ausentes o ambiguas."],
        "integration": "El pago no activa directamente desde la interfaz. La API recibe evidencia del proveedor, actualiza la orden y deriva suscripción y onboarding. No se ejecutó un cobro real durante certificación para evitar una transacción; el contrato se verificó con proveedor simulado y pruebas de webhook.",
        "evidence": "backend/src/services/commercial-payment.js; test/mercado-pago.test.js; RC-COMERCIAL-FINAL-01.md",
        "status": "Contrato certificado con simulación; monitoreo de webhooks sigue siendo requisito productivo.",
    },
    {
        "title": "30. Suscripción, cambios y cancelación",
        "overview": "La suscripción representa el acceso comercial vigente de la organización. Ventas incluye una máquina de presentación para estados como trial, activo, pago fallido, cancelado y cambio programado, alineada con reglas del backend.",
        "items": ["Activación derivada de pago aprobado.", "Cambio de plan mediante acción real.", "Cancelación con protección contra repetición.", "Presentación de periodo e importe sin prometer renovación inexistente.", "Consistencia del estado activo compartido entre backend, Portal y Mobile."],
        "integration": "El estado de suscripción condiciona acceso y acciones. Una corrección reciente unificó la interpretación de suscripción activa y evitó que el Mobile derivara acceso de datos inconsistentes. Las transiciones se validan en backend; el frontend traduce a mensajes y acciones permitidas.",
        "evidence": "ventas/features/commercial/subscription-state.ts; subscription-validator.ts; commit a1c3daf; RC-SUBSCRIPTION-CONSISTENCY-01.md",
        "status": "Ciclo comercial real conectado; no se prometen reactivaciones o renovaciones no implementadas.",
    },
    {
        "title": "31. Activación, llaves y alta de conductores",
        "overview": "Después de la compra, la organización ingresa a un flujo de activación. El administrador genera llaves para que cada conductor complete su propio registro, evitando que el Portal cree credenciales o perfiles incompletos directamente.",
        "items": ["Generación, copia, compartir y revocación de activation keys.", "Canje por parte del conductor.", "Validación de unidad y suscripción activa.", "Timeline de eventos de activación.", "Navegación guiada hacia Equipo, Unidades y Rutas."],
        "integration": "La llave conecta identidad, organización y capacidad contratada. Las correcciones simplificaron el registro y validaron que una unidad no quede asociada de forma inválida. Portal muestra el progreso, mientras Mobile completa el alta del conductor.",
        "evidence": "portal-onboarding-screen.tsx; backend módulos de activación/comercial; RC-ACTIVATION-FLOW-CONSOLIDATION-01.md",
        "status": "Flujo consolidado y certificado en Portal, con responsabilidades separadas por actor.",
    },
    {
        "title": "32. Onboarding de empresa y puesta en marcha",
        "overview": "El onboarding guía desde la suscripción aprobada hasta una operación mínima: perfil de empresa, equipo, unidades, rutas y acceso móvil. El diseño evita presentar el Portal como terminado solo porque existe el pago.",
        "items": ["Wizard y progreso por pasos.", "Enlaces directos a pantallas administrativas.", "Starter fleet cuando aplica.", "Seguimiento de eventos y bloqueos.", "Mensajes diferenciados entre pendiente, activo y acción requerida."],
        "integration": "Cada paso consume datos reales del backend. El onboarding no mantiene una segunda lista de unidades o usuarios: consulta los stores existentes. La llave de conductor y la descarga de la app forman parte de la puesta en marcha, no del checkout mismo.",
        "evidence": "portal-onboarding-screen.tsx; commercial-activation.js; RC-PORTAL-ACTIVATION-FINAL-01.md",
        "status": "Ciclo de activación completo y enlazado con administración real.",
    },
    {
        "title": "33. Arquitectura del Portal web",
        "overview": "El Portal vive dentro del paquete Ventas y utiliza React Native Web/Vite para ofrecer administración de cuenta y operación. PortalLayout protege el acceso, organiza navegación y aplica permisos; stores separados mantienen datos operativos y comerciales.",
        "items": ["Layout responsive con sidebar y navegación por secciones.", "useAppStore para usuarios y unidades.", "usePortalStore para cuenta, plan, facturas, sesiones y onboarding.", "Cliente API compartido y normalizador de errores.", "Error boundary a nivel de aplicación."],
        "integration": "El Portal reutiliza endpoints existentes; no implementa lógica de negocio móvil ni crea fuentes paralelas. La certificación final verificó la cadena UI -> store -> API -> backend -> store -> render y el control de permisos por módulo.",
        "evidence": "ventas/features/portal/components/portal-layout.tsx; use-portal-store.ts; RC-PORTAL-FINAL-01.md",
        "status": "Portal certificado dentro de configuración y credenciales del entorno.",
    },
    {
        "title": "34. Dashboard de Operaciones",
        "overview": "La pantalla principal del Portal concentra métricas, jornadas, eventos, checkpoints, posiciones y accesos rápidos. Fue una de las superficies con más iteraciones de UI y refactor por su densidad funcional.",
        "items": ["Filtros por unidad, conductor, ruta y estado.", "Métricas derivadas de datos operativos reales.", "Paneles de historial, recorridos y detalle.", "Mapa OperationsMap conectado a ubicaciones y geometría.", "Acciones rápidas reutilizando PortalButton."],
        "integration": "El dashboard no utiliza fixtures: combina usuarios, unidades y endpoints de historial. Las mejoras sustituyeron UUID por etiquetas comprensibles, tradujeron estados de jornada y normalizaron errores. También se corrigieron fallos de runtime durante extracciones de componentes compartidos.",
        "evidence": "portal-dashboard-screen.tsx; operations-map.tsx; RC-OPERATIONS-UI-MASTER-01.md; RC-PORTAL-OPERATIONS-CERTIFICATION-01.md",
        "status": "Centro operativo consolidado y conectado a seguimiento real.",
    },
    {
        "title": "35. Mapa y seguimiento en el Portal",
        "overview": "El mapa web permite observar unidades, recorridos y rutas sin exponer configuración técnica al usuario. Si Mapbox falta o falla, la experiencia conserva una representación funcional basada en última ubicación real y datos de recorrido.",
        "items": ["OperationsMap con marcadores y geometría.", "Fallback operativo sin instrucciones para desarrolladores.", "Filtros completos sin truncar resultados.", "Frescura GPS compartida con Mobile.", "Detalle de jornada, ruta y conductor."],
        "integration": "Portal y Mobile usan la misma ubicación autoritativa. El mapa web no calcula una verdad alterna y respeta gpsFreshness. Las mejoras recientes rediseñaron seguimiento y polilíneas, con énfasis en estilos y lectura de recorrido.",
        "evidence": "ventas/features/portal/components/operations-map.tsx; utils/tracking.ts; RC-MAP-SALES-DEEP-AUDIT-01.md",
        "status": "Seguimiento web funcional con fallback y paridad de datos.",
    },
    {
        "title": "36. Administración de Equipo, Unidades y Rutas",
        "overview": "Tres pantallas cubren las entidades principales de puesta en marcha. Equipo administra usuarios habilitados; Unidades gestiona flotilla y asignaciones; Rutas define y vincula recorridos.",
        "items": ["Listas con estados de carga, vacío y error.", "Formularios y modales de edición.", "Confirmaciones para eliminar, liberar o revocar.", "Permisos users, vehicles y routes.", "Listas compartidas mediante PortalDataList y PortalDataRow."],
        "integration": "Las pantallas consumen la misma identidad de unidad y usuario que Operaciones. La extracción de PortalButton y primitivas de lista redujo repetición, pero se verificó runtime debido a un incidente previo donde estilos evaluados al cargar el módulo derribaron Operaciones aun cuando typecheck pasaba.",
        "evidence": "portal-users-screen.tsx; portal-units-screen.tsx; portal-routes-screen.tsx; portal-button.tsx; portal-data-list.tsx",
        "status": "Administración integrada con componentes compartidos y reglas de acceso.",
    },
    {
        "title": "37. Documentos e Incidencias en el Portal",
        "overview": "Una auditoría de release detectó que el backend ya soportaba documentos e incidencias mientras el Portal no ofrecía superficies administrativas. El estado posterior del repositorio incluye portal-documents-screen y portal-incidents-screen, por lo que el reporte registra la brecha como etapa histórica y la incorporación como evolución posterior.",
        "items": ["Listado y composición de filas con primitivas compartidas.", "Consulta de estados y acciones disponibles.", "Integración con contratos existentes del backend.", "Manejo de contenido real, carga y error.", "Separación entre funcionalidad consolidada y pantallas aún en afinamiento."],
        "integration": "Documentos se relaciona con perfil, revisión y vencimientos; incidencias con seguimiento, unidad y alertas. La incorporación al Portal cerró una brecha administrativa, aunque las reglas visuales y de carga documental requieren validación continua y no deben documentarse con supuestos no verificados.",
        "evidence": "RC-RELEASE-CANDIDATE-01.md; portal-documents-screen.tsx; portal-incidents-screen.tsx; backend modules",
        "status": "Pantallas presentes en el estado actual; su madurez debe evaluarse por flujo y pruebas runtime.",
    },
    {
        "title": "38. Plan, facturación, pagos y perfil de cuenta",
        "overview": "El Portal ofrece superficies para consultar suscripción, facturas, pagos, empresa, seguridad y soporte. Estas pantallas consumen la orden y suscripción persistidas en lugar de inventar una billetera local.",
        "items": ["Mi plan con estado, importe y acciones reales.", "Facturas y descarga mediante URL del backend.", "Pagos y reintento sobre el plan existente.", "Perfil de empresa y sesiones activas.", "Revocación de sesión remota con confirmación."],
        "integration": "La certificación comercial eliminó métodos de pago locales que Mercado Pago no consumía y mensajes que prometían renovaciones inexistentes. Los errores deben distinguir ausencia real de datos de fallo de carga para no alarmar a clientes con suscripción válida.",
        "evidence": "portal-plan-screen.tsx; portal-billing-screen.tsx; portal-payments-screen.tsx; portal-profile-screen.tsx",
        "status": "Cuenta y ciclo comercial conectados al backend; feedback de errores continúa como área sensible.",
    },
    {
        "title": "39. Mobile App Center, versiones y descarga",
        "overview": "El Mobile App Center conecta configuración de versión en backend, administración en Portal, aviso de actualización en Mobile y descarga pública de la aplicación. Los commits más recientes enlazaron frontend y backend para la descarga.",
        "items": ["GET y PATCH de información de versión.", "Estadísticas de versiones de dispositivos.", "UpdateBanner y updateInfo en store móvil.", "Pantalla Portal para administrar versión y artefacto.", "Descarga expuesta desde experiencia comercial."],
        "integration": "La fuente de verdad es store.getAppConfig. Portal refleja GET /app/info y Mobile recibe updateInfo durante login/refresh. Una certificación arquitectónica corrigió un import wrapErrors inexistente que podía causar crash en startup y alineó requireAdmin con el patrón de middleware.",
        "evidence": "RC-MOBILE-APP-CENTER-05.2-CERTIFICATION.md; portal-app-movil-screen.tsx; update-banner.tsx; commits 71f88fd-30a2052",
        "status": "Módulo certificado y conectado de administración a descarga y actualización.",
    },
    {
        "title": "40. Sistema visual, componentes compartidos y refactor UI",
        "overview": "Mobile y Portal evolucionaron desde implementaciones locales hacia componentes reutilizables. El Portal introdujo PortalButton y primitivas de lista; Mobile mantiene PrimaryButton, StatusPill, AppCard y componentes de shell.",
        "items": ["Tokens de tema y paleta por superficie.", "PortalButton con variantes, tamaños, iconos, loading y disabled.", "PortalDataList y PortalDataRow con zonas composables.", "Cards, badges y layouts compartidos.", "Revisión de runtime además de typecheck y build."],
        "integration": "La lección principal fue que compilación no garantiza seguridad de evaluación de módulo. PortalButton provocó un error runtime al resolver estilos antes de inicialización; la corrección trasladó resolución al render. Las extracciones posteriores mantuvieron esa restricción y migraron solo patrones que encajaban limpiamente.",
        "evidence": "ventas/features/portal/components; portal-theme.ts; mobile/src/components; historial de extracciones",
        "status": "Base compartida en crecimiento, con disciplina de verificar cada pantalla en ejecución.",
    },
    {
        "title": "41. Seguridad, privacidad y endurecimiento productivo",
        "overview": "La seguridad se trabajó como capa transversal: autenticación, autorización, aislamiento de organización, headers, CORS, rate limiting, secretos, sanitización y almacenamiento seguro. También se corrigieron problemas productivos de proxy y variables.",
        "items": ["Helmet, CORS configurable y express-rate-limit.", "JWT, bcrypt y política de contraseñas.", "requireAdmin y permisos por dominio.", "Aislamiento multi-tenant en rutas y sockets.", "E2EE en chat directo y secure store en móvil."],
        "integration": "El despliegue en Render requirió confiar correctamente en proxy y manejar secretos mínimos. Cloudflare y API deben compartir orígenes autorizados. La seguridad del frontend reduce exposición, pero la decisión final permanece en backend. Los datos sensibles se eliminan de respuestas sanitizadas.",
        "evidence": "backend/src/middlewares; config/env.js; RC-AUTHORIZATION-FINAL-01.md; RC-CLOUDFLARE-VITE-ENV-01.md",
        "status": "Controles principales activos; secretos, CORS y tenant deben auditarse por entorno.",
    },
    {
        "title": "42. Observabilidad, auditoría y trazabilidad",
        "overview": "El backend registra eventos de API, solicitudes lentas, errores, actividad comercial, push, RTC e incidencias. Los identificadores de traza permiten correlacionar solicitudes y el Portal administrativo consulta snapshots operativos.",
        "items": ["x-trace-id en solicitudes.", "Eventos estructurados por módulo, acción y estado.", "Auditoría de cambios administrativos.", "Métricas de socket y sesiones RTC.", "Historial comercial y de activación."],
        "integration": "La observabilidad se mantuvo proporcional: llamadas reutilizan logger y store en vez de introducir Prometheus o tracing distribuido específico. Communication-service registra fallos estructurados sin derribar procesos. La trazabilidad permite relacionar compra, activación, sesión y operación.",
        "evidence": "backend/src/services/telemetry.js; audit.js; ops module; RC-WEBRTC-CERTIFICATION-01.md",
        "status": "Trazabilidad integrada; conviene establecer retención y tableros productivos formales.",
    },
    {
        "title": "43. Pruebas automatizadas, typecheck y calidad",
        "overview": "El repositorio contiene pruebas unitarias e integración distribuidas principalmente en backend y Mobile. En el conteo actual se localizaron 29 archivos de prueba en backend, 27 en Mobile y una prueba en communication-service; Ventas depende especialmente de typecheck, build y smokes específicos.",
        "items": ["Backend: contratos, tracking, comercial, auth y servicios.", "Mobile: navegación, mapa, radio, offline, E2EE y utilidades.", "TypeScript estricto en Mobile y Ventas.", "Build Vite y Android assembleDebug en certificaciones.", "Smokes de Mercado Pago, API, CORS y comunicación."],
        "integration": "Los reportes distinguen pruebas automáticas de validación manual. Una RC de tracking documentó 21 suites y 99 pruebas Mobile verdes, suite backend verde y builds aprobados, pero mantuvo abierta la certificación por falta de dispositivo y acceso a Mongo productivo. Esta separación evita declarar producción solo porque compila.",
        "evidence": "backend/test; mobile/src/**/*.test.ts; RC-TRACKING-EXECUTION-01.md; RC-CI-REPAIR-01.md",
        "status": "Cobertura significativa en capas críticas; Portal necesita ampliar pruebas automatizadas de UI/runtime.",
    },
    {
        "title": "44. CI, builds y proceso de release",
        "overview": "El proceso de release combina typecheck, pruebas, builds web y Android, revisión de diferencias y smokes. Diversos ciclos repararon CI, URLs de API, configuración Android y compatibilidad de despliegue.",
        "items": ["npm run typecheck en clientes TypeScript.", "npm test en paquetes con suites definidas.", "Vite build para Ventas.", "Gradle assembleDebug/assembleRelease para Android.", "Checklist y reportes de release candidate."],
        "integration": "El release no es un único comando: requiere validar backend, Portal, Mobile y proveedores. Los errores concurrentes se documentan sin atribuirlos al cambio en evaluación. Las advertencias de chunk grande en Vite se registran como informativas, no como build fallido.",
        "evidence": ".github; package.json por paquete; RC-CI-REPAIR-01.md; RC-RELEASE-01.md",
        "status": "Proceso repetible con certificaciones; falta automatizar más pruebas end-to-end del Portal.",
    },
    {
        "title": "45. Despliegue e infraestructura",
        "overview": "La solución contempla despliegue de backend en Render, Ventas en Cloudflare, MongoDB para persistencia, Redis para colas/arbitraje cuando aplica y coturn para llamadas en redes restrictivas. Docker Compose documenta dependencias locales y productivas.",
        "items": ["Variables VITE_API_URL y orígenes CORS.", "REQUIRE_MONGO para impedir store volátil en producción.", "Redis para BullMQ y arbitraje de Radio según configuración.", "TURN_URLS, TURN_SECRET y TURN_REALM para coturn REST.", "Artefactos Android y descarga desde el ecosistema comercial."],
        "integration": "Cada dependencia debe degradar de forma explícita. Correo puede deshabilitarse sin crash; TURN ausente reduce llamadas a STUN; Mongo requerido debe detener el arranque si no está disponible; y el Portal debe mostrar fallback de mapa sin filtrar variables técnicas.",
        "evidence": "docker-compose.yml; docker-compose.prod.yml; docs/deployment.md; docs/deploy-ventas-cloudflare.md",
        "status": "Arquitectura de despliegue definida; certificación final exige credenciales y pruebas del entorno real.",
    },
]


def cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(para.add_run("MANECOMB"), size=20, bold=True, color=NAVY)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(para.add_run("Reporte integral del desarrollo"), size=28, bold=True, color=NAVY)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(para.add_run("Aplicación móvil, Ventas, Portal web, Backend y servicios de comunicación"), size=15, italic=True, color=GRAY)
    for _ in range(7):
        doc.add_paragraph()
    table(doc, ["Campo", "Detalle"], [
        ["Periodo documentado", "Mayo-julio de 2026, con estado del repositorio al 21/07/2026"],
        ["Alcance", "Historia, arquitectura, implementación, integración, pruebas, despliegue y riesgos"],
        ["Repositorio", "combis-app / ManeComb"],
        ["Preparado para", "Revisión técnica, académica, administrativa y de continuidad del proyecto"],
    ], [2200, 7160], size=10)
    next_page(doc)


def toc_page(doc):
    h(doc, "Contenido", 1)
    base.add_toc(doc.add_paragraph())
    p(doc, "El índice es un campo automático de Word. Al abrir el archivo, seleccione Actualizar tabla para recalcular números de página definitivos.")
    callout(doc, "Organización", "El informe está diseñado en 50 páginas planificadas: cinco páginas iniciales y de cierre, más 45 capítulos técnicos que recorren el desarrollo de principio a fin.")
    next_page(doc)


def executive_page(doc):
    h(doc, "Resumen ejecutivo", 1)
    p(doc, "ManeComb evolucionó en un periodo corto desde una base de gestión de flotilla hasta un ecosistema que combina operación móvil, administración web, venta de planes, activación de organizaciones y comunicación en tiempo real. El repositorio actual reúne cuatro áreas principales: backend, mobile, ventas y communication-service. El inventario local contabiliza 154 archivos en backend, 271 en Mobile, 87 en Ventas y 37 en el servicio de comunicación, además de documentación, scripts e infraestructura.")
    p(doc, "El backend concentra autenticación, usuarios, unidades, rutas, ubicación, jornadas, incidencias, documentos, notificaciones, comercial, chat, radio, llamadas, aplicación y observabilidad. Mobile ejecuta seguimiento, control, incidencias, chat, Radio PTT, llamadas, perfil y actualización. Ventas integra landing, catálogo, checkout, Mercado Pago, onboarding y Portal. communication-service protege el envío de correo frente a fallos de proveedores y colas.")
    table(doc, ["Frente", "Resultado principal", "Estado resumido"], [
        ["Operación", "Seguimiento, mapa, jornadas, incidencias y alertas conectados.", "Implementado; validación física pendiente en escenarios específicos."],
        ["Comunicación", "Chat E2EE, multimedia, Radio PTT, presencia y WebRTC.", "Implementado con certificaciones y pruebas adicionales requeridas en red real."],
        ["Comercial", "Planes, checkout, Mercado Pago, suscripción y activación.", "Certificado a nivel de contratos, pruebas y UI."],
        ["Portal", "Cuenta, equipo, unidades, rutas, operación, pagos y soporte.", "Certificado en estado posterior a auditorías intermedias."],
    ], [1800, 4760, 2800], size=9)
    p(doc, "El resultado global es una plataforma integrada con una base técnica sólida, acompañada por deuda explícita: pruebas manuales en dispositivo, verificación de Mongo productivo, cobertura UI del Portal, rendimiento de recuperación de audio y monitoreo real de proveedores.")
    next_page(doc)


def chronology_page(doc):
    h(doc, "Cronología consolidada", 1)
    table(doc, ["Etapa", "Trabajo predominante", "Resultado"], [
        ["31 mayo", "Baseline del monorepo y limpieza inicial.", "Estructura común para todos los frentes."],
        ["3-18 junio", "Migración móvil, Ventas, CORS, auth, producción y CI.", "Clientes conectados al backend y despliegues estabilizados."],
        ["Finales de junio", "Comunicación, radio, E2EE, presencia y llamadas.", "Tiempo real y seguridad de comunicación fortalecidos."],
        ["Primera mitad de julio", "RBAC, control, historial, rutas y operación.", "Integridad de datos y consistencia operacional."],
        ["15-17 julio", "Auditorías RC, Portal, comercial, tracking y App Center.", "Hallazgos, correcciones y certificaciones por módulo."],
        ["18-21 julio", "Seguimiento, polilíneas, UI y descarga de app.", "Afinamiento visual y conexión de distribución móvil."],
    ], [1700, 4840, 2820], size=9)
    h(doc, "Cómo leer los estados", 2)
    p(doc, "Los RC son fotografías de momentos distintos. Un documento puede declarar un bloqueo que después fue corregido. Este reporte mantiene ambos datos: identifica el hallazgo original, describe la respuesta posterior y utiliza el árbol actual como referencia final. No se considera resuelto aquello que únicamente compila cuando el propio RC exige prueba física o acceso a infraestructura productiva.")
    h(doc, "Escala revisada", 2)
    bullets(doc, ["131 revisiones Git en el historial analizado.", "Más de 500 archivos entre los cuatro frentes principales.", "56 archivos de prueba localizados entre Backend y Mobile, más pruebas de comunicación.", "Decenas de reportes RC y auditorías especializadas.", "Despliegues y dependencias externas que requieren verificación por entorno."])
    next_page(doc)


def topic_page(doc, topic, number, total):
    h(doc, topic["title"], 1)
    p(doc, topic["overview"])
    h(doc, "Desarrollo realizado", 2)
    bullets(doc, topic["items"])
    h(doc, "Integración entre capas", 2)
    p(doc, topic["integration"])
    table(doc, ["Dimensión", "Evidencia del reporte"], [
        ["Frontend", "La interfaz consume stores, hooks y contratos; no sustituye reglas autoritativas del servidor."],
        ["Backend", "Las rutas validan identidad, permisos y datos antes de delegar en servicios y store."],
        ["Persistencia / realtime", "Los cambios se conservan o distribuyen según el dominio, con segmentación y trazabilidad."],
        ["Fuentes verificadas", topic["evidence"]],
    ], [2300, 7060], size=8)
    callout(doc, "Estado al corte", topic["status"])
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(para.add_run(f"Capítulo técnico {number} de {total}"), size=9, italic=True, color=GRAY)
    next_page(doc)


def conclusion_page(doc):
    h(doc, "Conclusiones, riesgos y siguiente etapa", 1)
    p(doc, "El desarrollo realizado demuestra una transición desde funciones aisladas hacia un sistema con contratos compartidos. La compra se convierte en suscripción y activación; la activación habilita perfiles, unidades y rutas; la operación produce ubicación, jornadas e incidencias; y la comunicación conecta alertas, chat, radio y llamadas. El valor principal está en esa continuidad.")
    h(doc, "Logros consolidados", 2)
    bullets(doc, ["Backend modular con persistencia Mongo/embebida y Socket.IO.", "App React Native CLI con operación, mapas y comunicación.", "Ventas y Portal conectados a catálogo, pago, suscripción y administración reales.", "Comunicación resiliente frente a fallos de proveedores y red.", "Auditorías y RC que registran decisiones, correcciones y límites honestos."])
    h(doc, "Riesgos pendientes", 2)
    bullets(doc, ["Completar matrices manuales en dispositivos físicos y redes distintas.", "Verificar datos productivos que no pudieron auditarse por restricciones de acceso.", "Aumentar pruebas automatizadas de UI y runtime en el Portal.", "Monitorear rendimiento de radio, retención de eventos y disponibilidad de proveedores.", "Mantener documentación alineada con el estado actual, separando histórico de certificación final."])
    h(doc, "Recomendación", 2)
    p(doc, "La siguiente fase debe priorizar certificación operacional reproducible: escenarios end-to-end con cuentas y datos controlados, evidencias de dispositivo, observabilidad de producción y criterios de aceptación por rol. Las nuevas funciones deben integrarse a contratos existentes y demostrar su necesidad antes de crear otro store, servicio o fuente de verdad.")
    h(doc, "Fuentes internas principales", 2)
    p(doc, "Código vigente de backend/, mobile/, ventas/ y communication-service/; docs/project-master.md; docs/alcance-sistema-combis.md; documentación de despliegue y llamadas; historial Git; reportes RC de release, Portal, Comercial, Tracking, Comunicación, App Center, autorización, radio, presencia y UI.")


def build():
    doc = Document()
    configure(doc)
    cover(doc)
    toc_page(doc)
    executive_page(doc)
    chronology_page(doc)
    for index, topic in enumerate(TOPICS, start=1):
        topic_page(doc, topic, index, len(TOPICS))
    conclusion_page(doc)
    doc.core_properties.title = "Reporte integral del desarrollo de ManeComb"
    doc.core_properties.subject = "Aplicación móvil, Ventas, Portal, Backend y servicios de comunicación"
    doc.core_properties.author = "Proyecto ManeComb"
    doc.core_properties.keywords = "ManeComb, desarrollo, backend, frontend, mobile, ventas, portal, reporte"
    doc.save(OUT)
    base.update_fields(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
