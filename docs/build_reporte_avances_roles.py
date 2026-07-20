from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_documentacion_tecnica_combis import (
    BLUE,
    DARK_BLUE,
    MUTED,
    NAVY,
    add_body,
    add_bullet,
    add_heading,
    add_page_number,
    add_small_note,
    add_table,
    configure_document,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "reporte-avances-analista-backend-frontend-tester.docx"


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(56)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("REPORTE DE AVANCES")
    set_run_font(r, size=12, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Sistema Inteligente de Gestión de Combis")
    set_run_font(r, size=26, bold=True, color=NAVY)

    p = doc.add_paragraph()
    r = p.add_run("Analista · Backend · Frontend · Tester")
    set_run_font(r, size=15, bold=True, color=DARK_BLUE)

    add_small_note(
        doc,
        "Objetivo",
        "Consolidar el trabajo realizado, la evidencia técnica disponible, el estado de validación y las siguientes acciones por disciplina.",
    )
    add_table(
        doc,
        ["Dato", "Valor"],
        [
            ["Fecha de corte", date.today().strftime("%d/%m/%Y")],
            ["Producto", "Aplicación móvil, backend API y portal web de ventas"],
            ["Alcance", "Análisis funcional, implementación, integración y aseguramiento de calidad"],
            ["Base del reporte", "Código, configuración de ejecución, pruebas automatizadas e historial reciente del repositorio"],
        ],
        [2700, 6660],
        font_size=9,
    )
    doc.add_page_break()


def add_summary(doc: Document):
    add_heading(doc, "1. Resumen ejecutivo", 1)
    add_body(doc, "El sistema evolucionó de pantallas aisladas hacia una plataforma conectada para operación de transporte. La solución integra una aplicación móvil para trabajo en campo, una API central con persistencia y tiempo real, y un portal web para ventas, activación y administración de cuentas.")
    add_body(doc, "El periodo reciente se concentró en seguimiento mediante mapas y polilíneas, refinamiento de interfaz, corrección del flujo de compra, validación de activaciones y consistencia del estado de suscripción. En paralelo, el repositorio ya contiene bases para navegación, GPS, incidencias, documentos, chat, radio, llamadas WebRTC, correo comercial, pagos, onboarding y aislamiento multiempresa.")
    add_table(
        doc,
        ["Frente", "Resultado consolidado", "Estado"],
        [
            ["Analista", "Flujos, actores, reglas, dependencias y riesgos principales identificados.", "Avance alto"],
            ["Backend", "API modular, autenticación, tiempo real, operación, comercial, comunicaciones y auditoría.", "Implementado / evolutivo"],
            ["Frontend", "App móvil operativa y portal web con seguimiento, mapas, cuenta, ventas y administración.", "Implementado / refinamiento"],
            ["Tester", "Pruebas backend y móviles para arquitectura, permisos, tracking, navegación, RTC y flujos críticos.", "Cobertura activa"],
        ],
        [1700, 5660, 2000],
        font_size=8.8,
    )
    add_small_note(doc, "Lectura del estado", "Implementado no significa cerrado: llamadas, push, correo productivo, observabilidad y pruebas del portal requieren endurecimiento antes de una liberación de producción.")
    add_heading(doc, "1.1 Evolución observada", 2)
    add_body(doc, "La evolución reciente puede agruparse en cuatro líneas. Primero se fortaleció el seguimiento geográfico mediante mapas, geometrías y polilíneas. Después se refinó la presentación para que los datos operativos fueran más legibles en móvil y web. En tercer lugar se corrigieron puntos del proceso comercial y de activación. Finalmente se trabajó en consistencia de estado, especialmente en suscripciones y asociación de conductores con unidades.")
    add_table(
        doc,
        ["Etapa", "Trabajo representativo", "Valor aportado"],
        [
            ["Base funcional", "Autenticación, usuarios, unidades, rutas, GPS y administración.", "Permite operar los recursos principales del sistema."],
            ["Integración", "REST, Socket.IO, navegación, chat, radio, incidencias y documentos.", "Conecta operación, comunicación y supervisión."],
            ["Comercial", "Planes, compra, pago, activación, onboarding, facturación y cuenta.", "Convierte la plataforma en un producto contratabile y administrable."],
            ["Refinamiento", "Mapas, seguimiento, UI, consistencia de suscripción y validaciones.", "Reduce errores y mejora uso cotidiano."],
        ],
        [1800, 4260, 3300],
        font_size=8.6,
    )
    add_heading(doc, "1.2 Alcance de este reporte", 2)
    add_body(doc, "Este informe no pretende sustituir historias de usuario, bitácoras de commits ni resultados de ejecución de CI. Su función es resumir el trabajo comprobable por disciplina y explicar cómo se relaciona con el producto. Los estados se expresan de forma prudente: una capacidad se considera implementada cuando existe código funcional; se considera validada únicamente cuando también existe evidencia de prueba o verificación.")
    doc.add_page_break()


def add_analyst(doc: Document):
    add_heading(doc, "2. Reporte del Analista", 1)
    add_heading(doc, "2.1 Trabajo realizado", 2)
    for item in [
        "Identificación de actores y permisos: propietario, administrador, despachador, supervisor, facturación, soporte, consulta y conductor.",
        "Definición del ciclo comercial: selección de plan, checkout, confirmación, activación, onboarding, claves de acceso y administración de suscripción.",
        "Definición del ciclo operativo: alta de unidad, asignación de conductor y ruta, sesión de recorrido, GPS, incidencias, documentos y seguimiento.",
        "Separación de responsabilidades entre app móvil, portal de ventas y backend como fuente de verdad y coordinador multiempresa.",
        "Inventario y unificación inicial de patrones de interfaz compartidos en el portal, incluyendo botones y primitivas de listas.",
        "Documentación técnica de entidades, procesos, arquitectura, llamadas, correo e interconexión de aplicaciones.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 Reglas y decisiones clave", 2)
    add_table(
        doc,
        ["Decisión", "Razón", "Impacto"],
        [
            ["Backend como fuente de verdad", "Evita estados distintos entre app y ventas.", "Consistencia de unidad, ruta, GPS, cuenta y suscripción."],
            ["Aislamiento por organización", "La plataforma atiende varias empresas.", "Toda lectura o mutación debe respetar organizationId."],
            ["REST + Socket.IO", "REST resuelve comandos y consultas; eventos propagan cambios.", "Actualización inmediata sin duplicar lógica de negocio."],
            ["Estados explícitos", "Pagos, documentos, incidencias y rutas tienen ciclos propios.", "Permite auditoría, mensajes claros y pruebas por transición."],
        ],
        [2200, 3260, 3900],
        font_size=8.6,
    )
    add_heading(doc, "2.3 Flujos funcionales analizados", 2)
    add_heading(doc, "Flujo comercial y de activación", 3)
    add_body(doc, "El prospecto consulta planes y complementos, inicia checkout y obtiene una orden con referencia. La confirmación de pago —manual, simulada o por webhook— modifica el estado comercial. Al activarse la suscripción, el responsable entra al portal, completa datos de empresa, genera claves de activación e incorpora usuarios y conductores. Este flujo enlaza ventas con operación: la cuenta activa determina qué organización puede crear recursos y qué límites de plan deben aplicarse.")
    add_heading(doc, "Flujo de operación de transporte", 3)
    add_body(doc, "El administrador registra unidades, conductores y rutas; asigna recursos y habilita el trabajo del conductor. La app inicia una sesión de ruta, obtiene ubicación, reporta posiciones y permite registrar pausas, finalización e incidencias. El portal consulta el estado consolidado, la geometría recorrida, métricas y eventos sin modificar la fuente de datos del móvil.")
    add_heading(doc, "Flujo de comunicación", 3)
    add_body(doc, "Los usuarios autenticados pueden entrar a conversaciones, enviar texto, audio o medios y participar en radio o llamadas. Socket.IO distribuye presencia y señalización; WebRTC transporta audio o video. Las notificaciones se guardan, se emiten en tiempo real y pueden enviarse como push cuando hay una suscripción configurada.")
    add_heading(doc, "2.4 Entregables de análisis", 2)
    add_table(
        doc,
        ["Entregable", "Contenido", "Uso"],
        [
            ["Modelo funcional", "Actores, permisos, módulos y responsabilidades.", "Alinear alcance entre producto y desarrollo."],
            ["Modelo de datos", "Entidades, relaciones, claves y restricciones.", "Guiar persistencia, APIs y validaciones."],
            ["Modelo de procesos", "Login, registro, rutas, GPS, viajes, pagos y administración.", "Identificar entradas, transiciones y resultados."],
            ["Arquitectura", "App móvil, backend, portal, sockets, correo, RTC e integraciones.", "Explicar dependencias y puntos de falla."],
            ["Inventario UI", "Patrones duplicados y candidatos a componentes compartidos.", "Reducir mantenimiento repetido."],
        ],
        [2200, 4060, 3100],
        font_size=8.5,
    )
    add_heading(doc, "2.5 Pendientes del análisis", 2)
    for item in [
        "Formalizar criterios de aceptación por módulo y por rol.",
        "Priorizar una versión mínima de producción frente a capacidades posteriores.",
        "Definir SLA de ubicación, llamadas, correo y atención de incidencias.",
        "Completar matriz de trazabilidad requisito–endpoint–pantalla–prueba.",
        "Acordar política de retención para GPS, grabaciones, documentos y auditoría.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "2.6 Riesgos funcionales", 2)
    for item in [
        "Diferencias entre el estado comercial y el acceso operativo si un evento de pago o activación no se procesa de forma idempotente.",
        "Confusión entre datos de ubicación en vivo y la proyección operacional consolidada si una pantalla usa la fuente incorrecta.",
        "Permisos demasiado amplios en operaciones administrativas o mezcla de datos entre organizaciones.",
        "Expectativas de llamadas o notificaciones push en redes reales sin infraestructura TURN/FCM terminada.",
        "Crecimiento de funciones sin criterios de aceptación medibles ni prioridad de versión.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def add_backend(doc: Document):
    add_heading(doc, "3. Reporte de Backend", 1)
    add_body(doc, "El backend está construido con Node.js, Express, MongoDB/Mongoose y Socket.IO. La organización por módulos separa autenticación, usuarios, unidades, ubicaciones, navegación, incidencias, documentos, portal, cuenta, comercial, notificaciones y RTC.")
    add_heading(doc, "3.1 Funcionalidad entregada", 2)
    rows = [
        ["Identidad y acceso", "JWT, refresh, sesiones, roles, permisos y aislamiento por organización."],
        ["Operación", "Proyección de unidades, GPS, rutas, sesiones, posiciones, métricas, bitácoras e incidencias."],
        ["Tiempo real", "Presencia, ubicación, chat, radio, notificaciones y señalización RTC mediante Socket.IO."],
        ["Comercial", "Planes, checkout, confirmación, webhook, activación, suscripción, facturas y descargas."],
        ["Comunicaciones", "Correo comercial con plantillas y Resend; WhatsApp opcional; push y eventos persistidos."],
        ["Seguridad", "Helmet, CORS, rate limiting, validación, control de tenant y manejo centralizado de errores."],
        ["Observabilidad", "Eventos de aplicación, trazabilidad, estados de entrega y consulta administrativa de sesiones RTC."],
    ]
    add_table(doc, ["Área", "Trabajo realizado"], rows, [2200, 7160], font_size=8.8)
    add_heading(doc, "3.2 Correcciones y refinamientos recientes", 2)
    for item in [
        "Consistencia del estado de suscripción activa compartido entre módulos.",
        "Validación de activación y asignación de unidad para conductores.",
        "Ajustes del flujo de compra y activación comercial.",
        "Mejoras de geolocalización y seguimiento operacional.",
        "Protecciones de integridad para tracking, permisos y aislamiento entre empresas.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "3.3 Arquitectura e interconexión", 2)
    add_body(doc, "Las rutas HTTP están separadas por dominio y protegidas por middlewares de autenticación y acceso. Los servicios concentran reglas que deben reutilizarse en distintos endpoints. La persistencia abstrae consultas de usuarios, órdenes, unidades, sesiones, documentos, notificaciones y eventos. Socket.IO complementa las respuestas REST con salas por usuario, rol y organización, evitando que el frontend dependa de sondeos continuos.")
    add_body(doc, "El módulo comercial publica eventos como account:created, payment:confirmed, plan:active y subscription:updated. El portal usa esas señales para actualizar onboarding y cuenta. En operación, los eventos de ubicación, presencia, chat, radio, incidencias y RTC mantienen sincronizados a los clientes conectados. Esta separación permite que el backend sea la autoridad y que cada frontend conserve solo el estado de presentación.")
    add_heading(doc, "3.4 Manejo de datos y seguridad", 2)
    for item in [
        "Los tokens JWT transportan identidad y claims; los refresh tokens permiten recuperar sesión sin reutilizar indefinidamente el access token.",
        "El acceso multiempresa se restringe con organizationId y middlewares especializados para portal y operación.",
        "Las contraseñas se almacenan con bcrypt y los errores se normalizan para no exponer detalles internos.",
        "Las cargas y documentos pasan por límites y servicios de almacenamiento; los enlaces se resuelven desde el backend.",
        "La telemetría registra eventos, fallos de push, sesiones RTC, operaciones comerciales y trazas de solicitud.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "3.5 Integraciones externas", 2)
    add_table(
        doc,
        ["Integración", "Uso actual", "Condición de producción"],
        [
            ["MongoDB", "Persistencia principal y modelos operativos/comerciales.", "Usuarios, índices, respaldo y mínimos privilegios."],
            ["Mercado Pago", "Checkout, confirmación y webhook comercial.", "Secretos, firma, idempotencia y conciliación."],
            ["Resend", "Correo transaccional con plantillas.", "Dominio verificado, SPF/DKIM/DMARC y métricas."],
            ["Twilio WhatsApp", "Avisos comerciales opcionales.", "Credenciales, plantillas aprobadas y consentimiento."],
            ["STUN/TURN", "Conectividad WebRTC.", "TURN redundante, TLS y monitoreo."],
            ["Mapas", "Geocodificación, planeación y geometrías.", "Cuotas, caché y manejo de indisponibilidad."],
        ],
        [1800, 3860, 3700],
        font_size=8.4,
    )
    add_heading(doc, "3.6 Deuda técnica y próximos pasos", 2)
    add_table(
        doc,
        ["Prioridad", "Acción", "Resultado esperado"],
        [
            ["Alta", "Configurar TURN, correo productivo, dominio y credenciales push.", "Comunicación estable fuera del entorno local."],
            ["Alta", "Ejecutar pruebas contra persistencia e integraciones reales controladas.", "Menor diferencia entre mocks y producción."],
            ["Media", "Colas persistentes e idempotencia para correo, webhooks y notificaciones.", "Recuperación segura ante fallos temporales."],
            ["Media", "Paneles de métricas para GPS, RTC, pagos, errores y tiempos de respuesta.", "Diagnóstico operativo medible."],
            ["Media", "Contratos OpenAPI y versionado explícito.", "Integración más predecible con clientes."],
        ],
        [1500, 4560, 3300],
        font_size=8.5,
    )
    add_heading(doc, "3.7 Criterios de terminación del backend", 2)
    for item in [
        "Contratos de API documentados y estables para los flujos P0.",
        "Pruebas automatizadas verdes con persistencia y servicios externos simulados de forma controlada.",
        "Migraciones, índices, variables de entorno y procedimientos de respaldo documentados.",
        "Métricas y alertas para latencia, errores, sockets, webhooks, correo y RTC.",
        "Prueba negativa de aislamiento multiempresa y permisos por rol antes de liberar.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def add_frontend(doc: Document):
    add_heading(doc, "4. Reporte de Frontend", 1)
    add_heading(doc, "4.1 Aplicación móvil", 2)
    for item in [
        "Navegación por rol, autenticación, recuperación de sesión y rutas protegidas.",
        "Mapa operativo con seguimiento, polilíneas, selección de unidad, estado de ubicación y paneles de control.",
        "Planeación y ejecución de rutas, sesiones, paradas, historial y bitácora.",
        "Chat, audio, medios, radio, presencia, notificaciones y llamadas WebRTC.",
        "Soporte Android para ubicación y llamadas en segundo plano, cámara, archivos y almacenamiento seguro.",
        "Manejo de conectividad, caché offline y recuperación de solicitudes críticas.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "4.2 Portal web de ventas", 2)
    for item in [
        "Experiencia comercial de planes, checkout, confirmación y entrada al portal.",
        "Pantallas de operaciones, unidades, rutas, usuarios, documentos, incidencias y onboarding.",
        "Administración de plan, pagos, facturación, sesiones, perfil y claves de activación.",
        "Mapa de seguimiento y geometría de rutas para supervisión desde escritorio.",
        "Store compartido con cargas agrupadas, caché temporal y aplicación de eventos en tiempo real.",
        "Primeras extracciones visuales compartidas: PortalButton y primitivas PortalDataList/PortalDataRow.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "4.3 Integración con datos y tiempo real", 2)
    add_body(doc, "Ambos clientes utilizan Axios para las solicitudes y Socket.IO para eventos. La app móvil prioriza tolerancia a conectividad variable, caché y reanudación de tareas; el portal prioriza cargas agrupadas, estados administrativos y actualización inmediata después de eventos comerciales. En los dos casos, la interfaz debe representar loading, vacío, error, permiso insuficiente y datos parciales sin sustituir la regla del backend.")
    add_body(doc, "El trabajo de mapas separa la geometría planificada, el recorrido registrado y la posición actual. Esta distinción es importante para no dibujar rutas engañosas. Las pantallas de seguimiento se refinaron para mostrar polilíneas, estado de unidad, información de sesión y controles contextuales sin saturar la vista.")
    add_heading(doc, "4.4 Diseño y mantenibilidad", 2)
    for item in [
        "Uso de tokens de color, espaciado, tipografía, radio y opacidad para reducir literales visuales.",
        "Extracción de PortalButton para variantes, tamaños, iconos, loading, disabled y ancho completo.",
        "Extracción de PortalDataList y PortalDataRow para filas componibles sin acoplar el componente al dominio.",
        "Conservación de casos especializados —timelines, mapas y tarjetas operativas— cuando una abstracción genérica dañaría la lógica.",
        "Error boundaries por pantalla para aislar fallos, aunque deben complementarse con pruebas runtime de todas las rutas.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "4.5 Trabajo reciente y pendientes", 2)
    add_table(
        doc,
        ["Tema", "Realizado", "Siguiente paso"],
        [
            ["Seguimiento", "Mapas, polilíneas y afinamiento de paneles.", "Pruebas con rutas extensas, reconexión y datos incompletos."],
            ["UI", "Refinamientos visuales y componentes compartidos iniciales.", "Continuar tokens, listas, feedback, KPIs y estados vacíos."],
            ["Ventas", "Compra, activación, cuenta y administración operativa.", "Pruebas automatizadas del portal y accesibilidad web."],
            ["Runtime", "Manejo de errores por pantalla y estados de carga.", "Pruebas de navegación completa en build productivo."],
        ],
        [1800, 3780, 3780],
        font_size=8.6,
    )
    add_heading(doc, "4.6 Criterios de terminación del frontend", 2)
    add_table(
        doc,
        ["Dimensión", "Criterio"],
        [
            ["Funcional", "Cada acción P0 completa su operación y refleja el resultado real del backend."],
            ["Runtime", "Todas las rutas cargan sin error boundary y muestran datos reales, vacíos y errores controlados."],
            ["Responsive", "Portal usable en resoluciones de escritorio y ventanas reducidas sin solapamiento."],
            ["Accesibilidad", "Foco, contraste, etiquetas y navegación por teclado verificados."],
            ["Rendimiento", "Mapas, listas y paneles evitan renders o solicitudes redundantes."],
            ["Mantenibilidad", "Patrones repetidos usan componentes y tokens compartidos cuando encajan."],
        ],
        [2200, 7160],
        font_size=8.6,
    )
    doc.add_page_break()


def add_tester(doc: Document):
    add_heading(doc, "5. Reporte del Tester", 1)
    add_body(doc, "La estrategia combina pruebas unitarias y de integración, typecheck, build, smoke tests, validación de arquitectura y recorridos punto a punto. El backend declara una cadena de pruebas críticas; la app móvil dispone de Jest, Playwright y Detox. El portal de ventas cuenta con typecheck y build, pero todavía no declara una suite automatizada propia.")
    add_heading(doc, "5.1 Cobertura disponible", 2)
    add_table(
        doc,
        ["Capa", "Cobertura observada", "Estado"],
        [
            ["Backend", "29 archivos; arquitectura, snapshot operativo, tracking, RBAC, tenant, pagos, RTC, navegación, activación y comunicación.", "Fuerte"],
            ["Móvil", "Suites configuradas para navegación, cifrado, presencia, RTC, caché, ubicación, mapas, radio y checklist.", "Activa"],
            ["E2E web", "Playwright disponible en la aplicación móvil/web compatible.", "Disponible"],
            ["E2E móvil", "Detox con scripts de build y ejecución.", "Disponible"],
            ["Ventas", "TypeScript sin emisión y build Vite.", "Falta suite funcional"],
        ],
        [1500, 6060, 1800],
        font_size=8.5,
    )
    add_heading(doc, "5.2 Escenarios críticos validados o cubiertos", 2)
    for item in [
        "Autenticación, refresh, contexto de usuario y recuperación de contraseña.",
        "Permisos por rol, acceso operacional y aislamiento por organización.",
        "Integridad de tracking, proyección de unidad y sesiones de ruta.",
        "Mercado Pago, activación, claves y asignación de conductor/unidad.",
        "Entrega de notificaciones, comunicación por correo y registro RTC/CDR.",
        "Navegación móvil, deep links, mapas, estado en tiempo real y caché offline.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.3 Enfoque de validación", 2)
    add_body(doc, "Las pruebas deben organizarse por pirámide. Las reglas puras y transformaciones se validan con pruebas unitarias; endpoints, repositorios y sockets se cubren con integración; los recorridos de usuario se reservan para E2E. Typecheck y build detectan incompatibilidades de tipos o empaquetado, pero no sustituyen la ejecución runtime, especialmente en pantallas con mapas, módulos nativos o eventos.")
    add_table(
        doc,
        ["Nivel", "Objetivo", "Ejemplos"],
        [
            ["Unitario", "Validar una regla en aislamiento.", "Permisos, estados, geometría, validadores, reducers."],
            ["Integración", "Validar colaboración entre módulos.", "Endpoint + store, socket + sala, pago + activación."],
            ["Contrato", "Evitar ruptura entre backend y clientes.", "Payloads, errores, estados y compatibilidad."],
            ["E2E", "Confirmar el resultado visible para el usuario.", "Compra, activación, ruta, GPS, llamada y documento."],
            ["No funcional", "Medir calidad operativa.", "Rendimiento, seguridad, accesibilidad y resiliencia."],
        ],
        [1500, 3460, 4400],
        font_size=8.5,
    )
    add_heading(doc, "5.4 Plan de pruebas pendiente", 2)
    add_table(
        doc,
        ["Prioridad", "Prueba", "Criterio de salida"],
        [
            ["P0", "Compra → pago → activación → primer acceso.", "Flujo completo sin intervención manual indebida."],
            ["P0", "Aislamiento entre dos organizaciones.", "Ningún dato, socket o archivo cruza tenants."],
            ["P0", "Seguimiento GPS con pérdida y recuperación de red.", "Sin saltos inválidos ni pérdida de sesión."],
            ["P1", "Llamadas en Wi-Fi, datos móviles y NAT restrictivo.", "Establecimiento, reconexión y cierre auditables."],
            ["P1", "Portal web en nueve pantallas principales.", "Sin error boundary; datos reales y acciones funcionales."],
            ["P1", "Correo, rebotes, reintentos y duplicados.", "Una entrega lógica por evento, con trazabilidad."],
            ["P2", "Accesibilidad, rendimiento y compatibilidad.", "Métricas y umbrales acordados por producto."],
        ],
        [1200, 4760, 3400],
        font_size=8.4,
    )
    add_heading(doc, "5.5 Gestión de defectos y evidencia", 2)
    for item in [
        "Cada defecto debe incluir ambiente, versión, rol, datos previos, pasos, resultado esperado, resultado real y evidencia visual o de logs.",
        "Los fallos intermitentes de GPS, sockets o RTC requieren timestamp, trace id, estado de red y dispositivo.",
        "Una corrección debe incorporar prueba de regresión automatizada cuando el nivel lo permita.",
        "Los casos bloqueados por credenciales externas deben registrarse como dependencia, no como aprobados.",
        "El cierre de versión requiere una matriz con P0/P1, estado, evidencia y riesgo aceptado.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.6 Riesgos de calidad", 2)
    add_table(
        doc,
        ["Riesgo", "Efecto", "Mitigación"],
        [
            ["Build verde sin prueba runtime", "Pantallas pueden fallar al evaluar módulos o datos reales.", "Smoke test navegando todas las rutas."],
            ["Mocks optimistas", "Producción difiere en latencia, permisos o formatos.", "Ambiente integrado con datos controlados."],
            ["Cobertura desigual", "Ventas acumula regresiones visuales o funcionales.", "Añadir unitarias y E2E del portal."],
            ["Servicios externos", "Pagos, correo, mapas o TURN fallan por configuración.", "Pruebas de readiness y degradación."],
            ["Concurrencia", "Eventos duplicados o estados fuera de orden.", "Idempotencia y pruebas simultáneas."],
        ],
        [2200, 3360, 3800],
        font_size=8.4,
    )
    doc.add_page_break()


def add_cross_team(doc: Document):
    add_heading(doc, "6. Reporte transversal de integración", 1)
    add_body(doc, "El valor del trabajo no está en cada disciplina por separado, sino en la continuidad entre requisito, regla, interfaz y prueba. Un cambio comercial, por ejemplo, comienza como definición de estado, se implementa en backend, se representa en el portal y se valida con escenarios de pago y activación. El mismo patrón se aplica a rutas, GPS, llamadas, documentos e incidencias.")
    add_table(
        doc,
        ["Capacidad", "Analista", "Backend", "Frontend", "Tester"],
        [
            ["Compra", "Estados y reglas", "Orden, pago, webhook", "Checkout y confirmación", "Flujo P0 e idempotencia"],
            ["Activación", "Criterios de acceso", "Cuenta, key y tenant", "Onboarding y registro", "Rol, unidad y primer acceso"],
            ["Seguimiento", "Fuente de verdad", "GPS, sesión y eventos", "Mapa y polilínea", "Red, precisión y regresión"],
            ["Llamadas", "Casos y restricciones", "ICE, TURN y CDR", "WebRTC y servicio Android", "Redes, permisos y reconexión"],
            ["Correo", "Eventos y contenido", "Plantillas y entrega", "Estados y mensajes", "Reintento, rebote y duplicado"],
        ],
        [1500, 1900, 2200, 2000, 1760],
        font_size=7.9,
    )
    add_heading(doc, "6.1 Cronología resumida del trabajo reciente", 2)
    add_table(
        doc,
        ["Periodo", "Avance", "Resultado"],
        [
            ["17 de julio", "Afinamiento de seguimiento en UI y uso de mapa.", "Mayor claridad de la operación geográfica."],
            ["18 de julio", "Continuidad de mapas y corrección de compra.", "Seguimiento más completo y flujo comercial corregido."],
            ["19 de julio", "Geolocalización, activación de conductor y suscripción.", "Datos y acceso más consistentes."],
            ["20 de julio", "Polilínea, refinamiento de UI y seguimiento en ventas.", "Mejor lectura del recorrido desde móvil y portal."],
        ],
        [1600, 4560, 3200],
        font_size=8.6,
    )
    add_heading(doc, "6.2 Próxima iteración coordinada", 2)
    for item in [
        "Analista: cerrar criterios de aceptación de compra, activación, seguimiento y llamadas.",
        "Backend: readiness de integraciones, idempotencia y observabilidad de flujos P0.",
        "Frontend: pruebas runtime del portal y manejo homogéneo de errores, vacíos y permisos.",
        "Tester: automatizar portal, ejecutar matriz multiempresa y validar redes reales.",
        "Equipo: acordar una definición de terminado común con evidencia y riesgo residual.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def add_close(doc: Document):
    add_heading(doc, "7. Conclusión y próximos hitos", 1)
    add_body(doc, "El producto cuenta con una base funcional amplia y una integración coherente entre operación móvil, servicios backend y administración comercial. El avance reciente mejoró especialmente seguimiento, mapas, interfaz y activación. La siguiente fase debe priorizar confiabilidad de producción y evidencia de calidad, no únicamente nuevas pantallas.")
    for item in [
        "Cerrar criterios de aceptación y matriz de trazabilidad por rol.",
        "Automatizar pruebas del portal de ventas y ejecutar recorridos runtime de todas sus pantallas.",
        "Configurar infraestructura real de TURN, correo y push con observabilidad.",
        "Validar compra, activación, GPS y llamadas en ambientes cercanos a producción.",
        "Continuar la unificación visual sin alterar reglas de negocio.",
    ]:
        add_bullet(doc, item)
    add_small_note(doc, "Resultado esperado", "Una versión liberable, medible y reproducible, con trazabilidad desde el requisito hasta la prueba y evidencia de operación estable.")


def build():
    doc = Document()
    configure_document(doc)
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Sistema Inteligente de Gestión de Combis | Reporte de avances"
    section.footer.paragraphs[0].clear()
    add_page_number(section.footer.paragraphs[0])
    add_cover(doc)
    add_summary(doc)
    add_analyst(doc)
    add_backend(doc)
    add_frontend(doc)
    add_tester(doc)
    add_cross_team(doc)
    add_close(doc)
    doc.core_properties.title = "Reporte de avances por rol"
    doc.core_properties.subject = "Analista, backend, frontend y tester"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
