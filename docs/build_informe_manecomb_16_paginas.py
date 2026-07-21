from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import build_actividad_usabilidad_manecomb as base


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "informe-integral-usabilidad-asistencia-manecomb-16-paginas.docx"


def p(doc, text, lead=None):
    return base.paragraph(doc, text, bold_lead=lead)


def h(doc, text, level=1):
    return base.heading(doc, text, level)


def bullets(doc, items):
    for item in items:
        base.bullet(doc, item)


def steps(doc, items):
    for item in items:
        base.numbered(doc, item)


def table(doc, headers, rows, widths, size=9):
    return base.add_table(doc, headers, rows, widths, size=size)


def callout(doc, label, text, fill="EEF3F8"):
    return base.callout(doc, label, text, fill=fill)


def next_page(doc):
    doc.add_page_break()


def page_1_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.font(para.add_run("MANECOMB"), size=18, bold=True, color=base.BLUE)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.font(para.add_run("Informe integral de usabilidad, documentación y asistencia técnica"), size=22, bold=True, color=base.NAVY)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.font(para.add_run("Seguimiento, mapa, alertas, chat, Radio PTT, perfil y evolución funcional"), size=14, italic=True, color=base.GRAY)
    for _ in range(5):
        doc.add_paragraph()
    table(doc, ["Dato", "Información"], [
        ["Actividad", "Diseño de documentación y asistencia técnica centrada en el usuario"],
        ["Sistema", "ManeComb - Plataforma inteligente para la gestión de transporte colectivo"],
        ["Alumno", "____________________________________________"],
        ["Asignatura", "____________________________________________"],
        ["Docente", "____________________________________________"],
        ["Fecha", date.today().strftime("%d/%m/%Y")],
    ], [2300, 7060], size=11)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.font(para.add_run("Documento académico en formato APA 7"), italic=True, color=base.GRAY)
    next_page(doc)


def page_2_index(doc):
    h(doc, "Contenido", 1)
    base.add_toc(doc.add_paragraph())
    callout(doc, "Nota de lectura", "Este informe presenta ManeComb como un ecosistema operativo. Los módulos centrales son seguimiento, mapa, alertas, chat, Radio PTT y perfil. La gestión documental se analiza únicamente como actualización futura, pues sus pantallas y reglas todavía requieren definición funcional.")
    h(doc, "Estructura del informe", 2)
    table(doc, ["Bloque", "Contenido", "Páginas"], [
        ["Fundamentos", "Contexto, alcance, arquitectura funcional y usuarios", "3-5"],
        ["Diagnóstico", "Seguimiento, mapa, alertas, chat, Radio PTT y perfil", "6-10"],
        ["Propuesta", "Asistencia, guías operativas, prototipos y accesibilidad", "11-14"],
        ["Evolución", "Evaluación, actualización de Documentos, conclusiones y referencias", "15-16"],
    ], [2100, 5260, 2000], size=10)
    p(doc, "La tabla de contenido es un campo automático de Word. Para actualizar números de página, seleccione la tabla y utilice la opción Actualizar campo.")
    next_page(doc)


def page_3_intro(doc):
    h(doc, "1. Introducción", 1)
    p(doc, "ManeComb es una plataforma orientada a la coordinación del transporte colectivo. Su valor no depende de una pantalla aislada, sino de la continuidad entre la ubicación de las unidades, la interpretación del mapa, la detección de eventos, la comunicación entre personas y el seguimiento de decisiones. Una experiencia satisfactoria exige que cada módulo conserve el contexto del anterior: una alerta debe conducir a la unidad afectada; desde esa unidad debe abrirse el chat o la radio; y la conversación debe permitir volver al evento que originó la atención.")
    p(doc, "La documentación y la asistencia técnica forman parte de esa continuidad. No deben limitarse a explicar botones, sino ayudar a comprender estados, prioridades, consecuencias y rutas de recuperación. Cuando el sistema informa que una unidad está sin conexión, el usuario necesita saber desde cuándo, en qué ruta se encontraba, qué nivel de urgencia tiene el evento, quién lo está atendiendo y qué canal de comunicación está disponible. Este enfoque convierte la ayuda en una herramienta operacional, no en un manual separado de la actividad cotidiana.")
    h(doc, "1.1 Objetivo general", 2)
    p(doc, "Diseñar una propuesta integral de documentación y asistencia técnica para ManeComb que mejore la comprensión, prevención de errores y recuperación en seguimiento, mapa, alertas, chat, Radio PTT y perfil, incorporando accesibilidad y una hoja de ruta realista para funciones todavía en evolución.")
    h(doc, "1.2 Objetivos específicos", 2)
    bullets(doc, [
        "Identificar fricciones que afectan la supervisión y la coordinación en tiempo real.",
        "Definir mensajes, ayudas contextuales y mecanismos proactivos según el riesgo de cada tarea.",
        "Representar la interconexión entre eventos, unidades, rutas, alertas y conversaciones.",
        "Proponer criterios accesibles para estados dinámicos, navegación, contraste y teclado.",
        "Ubicar el módulo Documentos como actualización futura, sin atribuirle pantallas o comportamientos no consolidados.",
    ])
    h(doc, "1.3 Alcance y método", 2)
    p(doc, "El análisis adopta principios de diseño centrado en el usuario, heurísticas de visibilidad y retroalimentación, prevención de errores y criterios WCAG 2.2. Se revisan tareas representativas desde la perspectiva de administración, operación y conducción. Las propuestas son conceptuales y deben validarse con usuarios reales antes de convertirse en requisitos definitivos.")
    next_page(doc)


def page_4_system(doc):
    h(doc, "2. Visión integral del sistema", 1)
    p(doc, "ManeComb articula información estática —usuarios, perfiles, rutas y unidades— con información dinámica —posición, conexión, mensajes, transmisiones y alertas—. El mapa constituye la superficie principal para comprender la operación; el seguimiento organiza el historial y estado de cada recorrido; las alertas señalan desviaciones que requieren atención; el chat mantiene coordinación asíncrona; Radio PTT responde a comunicaciones rápidas; y el perfil determina identidad, permisos y preferencias.")
    h(doc, "2.1 Interconexión funcional", 2)
    table(doc, ["Origen", "Evento o dato", "Destino", "Resultado esperado"], [
        ["Seguimiento", "Cambio de posición, demora o pérdida de señal", "Mapa y alertas", "Actualizar la representación y priorizar una posible incidencia."],
        ["Mapa", "Selección de unidad o ruta", "Detalle, chat y Radio PTT", "Conservar el contexto al iniciar comunicación."],
        ["Alertas", "Evento crítico, alto, medio o informativo", "Unidad, incidencia o conversación", "Atender, asignar responsable, posponer o escalar."],
        ["Chat", "Mensaje, evidencia o respuesta", "Seguimiento e incidencias", "Registrar acuerdos y mantener trazabilidad."],
        ["Radio PTT", "Transmisión inmediata", "Canal y participantes", "Coordinar una acción urgente con retroalimentación clara."],
        ["Perfil", "Rol, permisos y preferencias", "Todos los módulos", "Mostrar funciones autorizadas y personalizar avisos."],
    ], [1500, 2780, 2020, 3060], size=8)
    h(doc, "2.2 Flujo operativo de referencia", 2)
    steps(doc, [
        "El sistema recibe la posición y el estado de conexión de una unidad.",
        "El mapa actualiza el marcador y el seguimiento registra el cambio.",
        "Una regla detecta una condición que merece atención y genera un evento.",
        "El centro de alertas asigna severidad, contexto y estado de atención.",
        "El operador abre la unidad, consulta su ruta e inicia chat o Radio PTT.",
        "La respuesta, responsable y resolución quedan vinculados al evento original.",
    ])
    callout(doc, "Principio rector", "Ningún cambio de módulo debe obligar al usuario a reconstruir manualmente qué unidad, ruta, alerta o conversación estaba atendiendo.")
    next_page(doc)


def page_5_users(doc):
    h(doc, "3. Usuarios y contexto de uso", 1)
    p(doc, "La calidad de la asistencia depende del rol y del contexto. Un administrador trabaja con configuraciones y supervisión global; un operador necesita priorizar eventos bajo presión; un conductor utiliza el sistema en movilidad, con atención dividida y conectividad variable. La misma explicación no funciona para los tres perfiles.")
    table(doc, ["Perfil", "Metas", "Condiciones", "Necesidades de asistencia"], [
        ["Administrador", "Configurar cuentas, rutas, unidades, permisos y políticas.", "Uso frecuente en escritorio; decisiones con impacto global.", "Confirmaciones, historial de cambios, lenguaje consistente y prevención de acciones irreversibles."],
        ["Operador", "Monitorear unidades, atender alertas y coordinar respuestas.", "Varias tareas simultáneas; alto volumen de información.", "Priorización, filtros persistentes, contexto inmediato y acciones rápidas sin perder trazabilidad."],
        ["Conductor", "Consultar ruta, reportar situación y comunicarse.", "Movilidad, ruido, conectividad inestable y pantalla pequeña.", "Controles grandes, mensajes breves, estados claros, recuperación automática y mínima escritura."],
        ["Analista o supervisor", "Revisar tendencias, tiempos de respuesta e incidencias.", "Consulta histórica y comparación entre periodos.", "Definiciones, filtros explicados, exportación y relación entre indicador y evidencia."],
    ], [1700, 2350, 2400, 2910], size=8)
    h(doc, "3.1 Escenarios críticos", 2)
    bullets(doc, [
        "Una unidad deja de transmitir ubicación durante un recorrido activo.",
        "Varias alertas similares se generan en pocos minutos y compiten por atención.",
        "Un mensaje de chat queda pendiente por falta de red y el usuario no sabe si reenviarlo.",
        "Dos participantes intentan usar Radio PTT al mismo tiempo.",
        "El usuario cambia de teléfono o sesión y necesita recuperar preferencias y conversaciones.",
    ])
    h(doc, "3.2 Criterios de éxito", 2)
    p(doc, "La propuesta se considera útil si reduce el tiempo para localizar una unidad, permite distinguir rápidamente la prioridad de una alerta, evita mensajes duplicados, aclara el estado de una transmisión, conserva el contexto al cambiar de módulo y comunica límites de permiso sin exponer términos internos. También debe funcionar con teclado, lector de pantalla, aumento de texto y conectividad intermitente.")
    next_page(doc)


def page_6_tracking_map(doc):
    h(doc, "4. Diagnóstico: seguimiento y mapa", 1)
    p(doc, "Seguimiento y mapa forman una sola experiencia cognitiva. El primero explica el estado y la evolución temporal; el segundo sitúa ese estado en el territorio. Si los filtros, leyendas o tiempos de actualización no son claros, el usuario puede confundir una posición antigua con una ubicación actual o interpretar una pérdida de señal como inmovilidad.")
    h(doc, "4.1 Fricciones principales", 2)
    table(doc, ["Fricción", "Impacto", "Asistencia recomendada"], [
        ["Marca de tiempo poco visible", "Se toman decisiones con datos antiguos.", "Mostrar Última actualización y cambiar a Estado desactualizado después del umbral definido."],
        ["Marcadores similares", "Cuesta distinguir unidades activas, detenidas o sin conexión.", "Combinar forma, icono, texto y color; incluir leyenda accesible."],
        ["Filtros que se reinician", "El operador pierde su vista de trabajo y repite configuraciones.", "Conservar filtros por sesión y ofrecer Restablecer con explicación."],
        ["Mapa separado del detalle", "Se pierde la relación entre ubicación, ruta, conductor y evento.", "Panel contextual persistente con accesos a seguimiento, alertas, chat y radio."],
        ["Carga o red sin estado", "El usuario no sabe si el mapa está actualizando o falló.", "Indicador no bloqueante, hora de último dato y acción Reintentar."],
    ], [2400, 2860, 4100], size=9)
    h(doc, "4.2 Ayuda contextual propuesta", 2)
    callout(doc, "Unidad C-24", "En ruta · Última ubicación hace 35 s\nRuta Centro-Terminal · Conductor asignado\n[Ver seguimiento] [Abrir chat] [Radio PTT]")
    p(doc, "El panel evita que el usuario interprete el marcador sin contexto. Si la ubicación envejece, el texto cambia progresivamente: Actualizada, Demorada y Sin datos recientes. Estos estados deben basarse en reglas conocidas por operación y mostrarse en la leyenda. No se debe animar el marcador para simular movimiento cuando no existen datos nuevos.")
    h(doc, "4.3 Recuperación", 2)
    p(doc, "Cuando falla la actualización, el mapa conserva el último dato válido y lo identifica como histórico. La acción Reintentar no elimina filtros ni selección. Si la unidad permanece sin señal, se ofrece crear o abrir una alerta vinculada. Esta recuperación protege la continuidad y evita que una falla técnica borre la evidencia que el operador necesita.")
    next_page(doc)


def page_7_alerts(doc):
    h(doc, "5. Diagnóstico: centro de alertas", 1)
    p(doc, "El centro de alertas traduce eventos automáticos en trabajo humano. Su problema principal no es únicamente mostrar avisos, sino ayudar a decidir qué atender primero, quién es responsable y cuándo un evento puede considerarse resuelto. Una lista cronológica sin jerarquía produce fatiga, especialmente cuando varias reglas reaccionan al mismo incidente.")
    h(doc, "5.1 Modelo de severidad", 2)
    table(doc, ["Nivel", "Uso", "Presentación", "Respuesta"], [
        ["Crítica", "Riesgo inmediato o pérdida grave de continuidad.", "Texto CRÍTICA, icono, contraste alto y persistencia.", "Confirmación explícita, responsable y escalamiento."],
        ["Alta", "Situación que puede afectar una ruta activa.", "Prioridad visible y tiempo transcurrido.", "Atender en periodo definido o justificar posposición."],
        ["Media", "Desviación que requiere revisión, no interrupción.", "Centro de alertas y notificación moderada.", "Revisar, agrupar o programar seguimiento."],
        ["Informativa", "Cambio esperado o confirmación.", "Sin interrupción; disponible en historial.", "Lectura opcional y cierre automático cuando proceda."],
    ], [1300, 2730, 2780, 2550], size=8)
    h(doc, "5.2 Ciclo de vida", 2)
    steps(doc, [
        "Detectada: el evento se registra con origen, unidad, ruta y hora.",
        "Nueva: todavía no existe confirmación humana.",
        "Reconocida: una persona acepta revisar el evento.",
        "En atención: existe responsable y acción en curso.",
        "Resuelta: se documenta resultado y hora de cierre.",
        "Reabierta: una nueva condición demuestra que el problema continúa.",
    ])
    h(doc, "5.3 Prevención de fatiga", 2)
    bullets(doc, [
        "Agrupar eventos equivalentes por unidad, regla y periodo sin ocultar su cantidad.",
        "Permitir preferencias por rol, turno y gravedad; las críticas no se silencian sin autorización.",
        "Evitar repetir la misma notificación en web y móvil cuando ya fue reconocida.",
        "Medir falsos positivos, tiempo de reconocimiento, resolución y reapertura.",
    ])
    callout(doc, "Mensaje recomendado", "Unidad C-24 sin ubicación desde hace 4 min. Último punto: Av. Central. Ruta Centro-Terminal. [Ver unidad] [Abrir chat] [Atender]", fill="FFF2CC")
    next_page(doc)


def page_8_chat(doc):
    h(doc, "6. Diagnóstico: chat operativo", 1)
    p(doc, "El chat sostiene conversaciones que no requieren ocupar el canal de radio y conserva evidencia consultable. Para funcionar en transporte debe tolerar interrupciones, cambios de red y respuestas tardías. La interfaz debe distinguir con claridad el estado local del mensaje del estado confirmado por el servidor.")
    h(doc, "6.1 Estados del mensaje", 2)
    table(doc, ["Estado", "Significado para el usuario", "Acción disponible"], [
        ["Escribiendo", "El contenido todavía no se ha enviado.", "Editar, adjuntar o descartar."],
        ["Pendiente", "El dispositivo no logró confirmar el envío.", "Esperar reintento automático o Reintentar."],
        ["Enviado", "El servidor recibió el mensaje.", "Continuar; no implica entrega al destinatario."],
        ["Entregado", "El dispositivo destinatario lo recibió.", "Esperar respuesta o escalar según urgencia."],
        ["Leído", "El destinatario abrió el mensaje.", "No asumir resolución; registrar respuesta si se requiere."],
        ["Fallido", "El envío no pudo completarse tras los reintentos.", "Editar, copiar, reenviar o cambiar de canal."],
    ], [1600, 4920, 2840], size=9)
    h(doc, "6.2 Contexto y continuidad", 2)
    p(doc, "Una conversación puede nacer desde una unidad, ruta, alerta o incidencia. El encabezado debe mostrar ese vínculo y permitir volver al origen. Al compartir una alerta se inserta una tarjeta con identificador, severidad, estado y hora, no una copia de texto que pueda quedar obsoleta. Cuando el evento cambia, la tarjeta se actualiza sin modificar el mensaje histórico.")
    h(doc, "6.3 Conectividad y privacidad", 2)
    bullets(doc, [
        "Conservar borradores y mensajes pendientes de forma segura en el dispositivo.",
        "Evitar duplicados mediante identificadores locales al reintentar.",
        "Mostrar el estado de conexión sin culpar al usuario.",
        "Limitar conversaciones y adjuntos según rol; registrar accesos administrativos.",
        "Definir retención, eliminación, exportación y tratamiento de contenido sensible.",
    ])
    callout(doc, "Mensaje de recuperación", "Sin conexión. Tu mensaje permanece pendiente y se enviará automáticamente al recuperar señal. [Reintentar ahora] [Copiar mensaje]")
    next_page(doc)


def page_9_ptt(doc):
    h(doc, "7. Diagnóstico: Radio PTT", 1)
    p(doc, "Radio PTT responde a situaciones donde escribir sería lento. Su modelo de mantener presionado para hablar no equivale a una llamada: el canal puede estar ocupado, la transmisión atraviesa varias etapas y el usuario necesita confirmar si el audio fue capturado y distribuido. La ayuda debe enseñar el gesto sin interferir con el uso habitual.")
    h(doc, "7.1 Estados comprensibles", 2)
    table(doc, ["Estado técnico", "Texto visible", "Retroalimentación"], [
        ["Listo", "Mantén presionado para hablar", "Control disponible y canal identificado."],
        ["Solicitando canal", "Conectando…", "Progreso breve; no iniciar grabación todavía."],
        ["Transmitiendo", "Hablando · suelta para enviar", "Color, texto, temporizador y vibración inicial."],
        ["Procesando", "Enviando audio…", "Mantener la transmisión visible y evitar doble envío."],
        ["Entregado", "Audio enviado", "Confirmación discreta y registro en la conversación."],
        ["Canal ocupado", "Otra persona está hablando", "Identificar participante y permitir escuchar."],
        ["Sin conexión", "No se pudo transmitir", "Reintentar, guardar como audio o abrir chat."],
    ], [1960, 3100, 4300], size=8)
    h(doc, "7.2 Tutorial progresivo", 2)
    steps(doc, [
        "Primer acceso: explicar el propósito del canal y solicitar permiso de micrófono con razón clara.",
        "Primer intento: resaltar el gesto Mantén presionado y confirmar con vibración.",
        "Primera transmisión: explicar Enviando y Audio enviado.",
        "Primer conflicto: enseñar Canal ocupado sin convertirlo en error.",
        "Usos posteriores: reducir instrucciones y conservar ayuda bajo un botón accesible.",
    ])
    h(doc, "7.3 Relación con chat y alertas", 2)
    p(doc, "Una alerta crítica puede ofrecer Radio PTT como acción, pero no debe iniciar audio automáticamente. Si una transmisión falla, el usuario puede convertirla en mensaje de voz dentro del chat. El historial indica que la comunicación surgió desde una alerta y conserva participantes, hora y resultado. Esta relación evita canales paralelos sin trazabilidad.")
    next_page(doc)


def page_10_profile(doc):
    h(doc, "8. Diagnóstico: perfil, cuenta y preferencias", 1)
    p(doc, "El perfil no es únicamente información personal. Es el punto donde el usuario comprende su identidad dentro de ManeComb, el rol asignado, los permisos disponibles, las preferencias de comunicación y la seguridad de la sesión. Una configuración confusa puede producir notificaciones excesivas o hacer que una función parezca averiada cuando en realidad está restringida.")
    h(doc, "8.1 Contenido recomendado", 2)
    table(doc, ["Área", "Información", "Asistencia"], [
        ["Identidad", "Nombre, fotografía, contacto y organización.", "Explicar qué datos son visibles para otros participantes."],
        ["Rol y permisos", "Administrador, operador, conductor u otro rol.", "Mostrar capacidades con lenguaje funcional y canal para solicitar cambios."],
        ["Notificaciones", "Gravedad, canal, horario y excepciones.", "Vista previa de efectos y protección para alertas críticas."],
        ["Comunicación", "Micrófono, sonido, vibración, estado y privacidad.", "Prueba de audio y diagnóstico de permisos."],
        ["Seguridad", "Sesiones, dispositivos, contraseña y cierre remoto.", "Confirmaciones, historial y pasos de recuperación."],
        ["Accesibilidad", "Tamaño de texto, movimiento y preferencias visuales.", "Persistencia entre dispositivos cuando sea posible."],
    ], [1800, 3820, 3740], size=9)
    h(doc, "8.2 Mensajes de permiso", 2)
    p(doc, "En lugar de mostrar Acceso denegado, ManeComb debe explicar: Tu rol de conductor puede consultar esta ruta, pero no modificarla. Si necesitas un cambio, contacta a Operaciones. La explicación reduce incertidumbre y ofrece una salida. La interfaz tampoco debe mostrar controles que nunca pueden utilizarse, salvo cuando sea necesario enseñar una capacidad y explicar cómo obtenerla.")
    h(doc, "8.3 Seguridad y recuperación", 2)
    bullets(doc, [
        "Confirmar cambios de contraseña y cierre de otras sesiones.",
        "Mostrar dispositivos activos con fecha aproximada y permitir revocación.",
        "No revelar si un correo pertenece a otra cuenta durante recuperación.",
        "Notificar cambios críticos por un canal alterno y mantener registro.",
        "Explicar claramente cuándo una preferencia es local y cuándo se sincroniza.",
    ])
    next_page(doc)


def page_11_assistance(doc):
    h(doc, "9. Modelo integral de asistencia técnica", 1)
    p(doc, "La asistencia de ManeComb debe aparecer en el momento y con la intensidad adecuados. La ayuda contextual explica un control; la ayuda proactiva anticipa una dificultad de alto riesgo; la asistencia inteligente interpreta una condición; y la base de conocimiento permite profundizar. Ninguna modalidad sustituye a las demás.")
    table(doc, ["Módulo", "Modalidad principal", "Ejemplo"], [
        ["Seguimiento y mapa", "Contextual", "Leyenda, vigencia del dato, filtros y panel de unidad."],
        ["Alertas", "Proactiva y priorizada", "Notificación según severidad, agrupación y escalamiento."],
        ["Chat", "Contextual e inteligente", "Estados de entrega, conectividad, reintento y vínculos."],
        ["Radio PTT", "Tutorial progresivo", "Gesto, canal ocupado, transmisión y recuperación."],
        ["Perfil", "Contextual y preventiva", "Permisos, seguridad, preferencias y consecuencias."],
        ["Documentos (futuro)", "Por definir mediante investigación", "No fijar restricciones ni flujo antes de validar pantallas y reglas."],
    ], [1900, 2500, 4960], size=9)
    h(doc, "9.1 Arquitectura de ayuda", 2)
    bullets(doc, [
        "Microcopy dentro de la interfaz para estados y restricciones inmediatas.",
        "Ayuda contextual enlazada desde controles complejos y mensajes de error.",
        "Centro de ayuda por tareas, no por nombres internos de pantallas.",
        "Diagnóstico guiado para conexión, permisos, micrófono y notificaciones.",
        "Escalamiento a soporte con contexto técnico adjunto y consentimiento del usuario.",
    ])
    h(doc, "9.2 Principios editoriales", 2)
    p(doc, "Los mensajes deben comenzar con el resultado observable, explicar la causa cuando sea conocida y terminar con una acción. Se evita lenguaje como socket, payload, token o error 500 en la interfaz. El detalle técnico puede conservarse en un identificador copiable para soporte. La misma condición debe mantener el mismo nombre en web, móvil, correo y documentación.")
    callout(doc, "Patrón de mensaje", "Qué ocurrió + qué significa + qué puede hacer el usuario. Ejemplo: Sin conexión. La ubicación mostrada es de hace 6 min. Reintenta o abre el chat para coordinar con la unidad.")
    next_page(doc)


def page_12_guide_tracking(doc):
    h(doc, "10. Guía operativa: seguimiento y mapa", 1)
    p(doc, "Esta guía está dirigida a personal de operación que necesita localizar unidades y evaluar la vigencia de la información antes de tomar decisiones.")
    h(doc, "10.1 Consultar una unidad", 2)
    steps(doc, [
        "Abra Seguimiento y confirme que el indicador general muestre conexión activa.",
        "Utilice los filtros de ruta, estado o unidad; revise cuántos resultados permanecen visibles.",
        "Seleccione el marcador o la fila de la unidad. El panel debe conservar la misma selección en ambas vistas.",
        "Compruebe Última actualización antes de interpretar la ubicación.",
        "Revise ruta, conductor, estado de conexión y alertas activas.",
        "Si requiere coordinación, use Abrir chat o Radio PTT desde el panel contextual.",
        "Al finalizar, registre o cierre la alerta si la tarea se originó en un evento.",
    ])
    h(doc, "10.2 Interpretación de estados", 2)
    table(doc, ["Indicador", "Interpretación", "Acción"], [
        ["Actualizada", "El dato está dentro del umbral operativo.", "Continuar seguimiento normal."],
        ["Demorada", "La actualización tarda más de lo esperado.", "Verificar conexión y observar tendencia."],
        ["Sin datos recientes", "La posición no debe tratarse como actual.", "Abrir alerta, chat o Radio PTT según severidad."],
        ["Fuera de ruta", "La posición se separó del corredor definido.", "Confirmar precisión, contexto y causa antes de escalar."],
    ], [1900, 4280, 3180], size=9)
    h(doc, "10.3 Errores que deben evitarse", 2)
    bullets(doc, [
        "No asumir que un marcador inmóvil representa una unidad detenida sin revisar la hora.",
        "No quitar filtros para localizar una unidad si puede buscarse por identificador.",
        "No cerrar una alerta solo porque la posición volvió; confirmar la condición de resolución.",
        "No compartir capturas sin hora ni identificador cuando exista un enlace contextual.",
    ])
    next_page(doc)


def page_13_guide_comms(doc):
    h(doc, "11. Guía operativa: alertas, chat y Radio PTT", 1)
    h(doc, "11.1 Atender una alerta", 2)
    steps(doc, [
        "Lea severidad, unidad, ruta, hora y condición antes de abrir la acción.",
        "Pulse Atender para registrar responsabilidad y evitar trabajo duplicado.",
        "Abra el contexto y compruebe seguimiento, última ubicación y eventos relacionados.",
        "Elija chat para coordinación documentada o Radio PTT para comunicación inmediata.",
        "Registre la decisión y cambie el estado a Resuelta solo cuando exista evidencia.",
    ])
    h(doc, "11.2 Continuar por chat", 2)
    steps(doc, [
        "Confirme que el encabezado corresponde a la unidad o alerta correcta.",
        "Redacte un mensaje breve con acción solicitada y tiempo esperado.",
        "Observe si queda Pendiente, Enviado, Entregado o Leído.",
        "Si falla, utilice Reintentar una sola vez; el sistema debe impedir duplicados.",
        "Adjunte evidencia solo cuando sea necesaria y esté autorizada.",
    ])
    h(doc, "11.3 Cambiar a Radio PTT", 2)
    steps(doc, [
        "Compruebe canal y participantes.",
        "Mantenga presionado; espere la vibración o confirmación de Transmitiendo.",
        "Hable con una instrucción concreta y suelte para enviar.",
        "Verifique Audio enviado. Si no se confirma, use chat o reintento guiado.",
    ])
    callout(doc, "Regla de canal", "Chat conserva detalle y evidencia; Radio PTT reduce tiempo de respuesta. Una alerta organiza la prioridad. El usuario debe poder cambiar de canal sin perder el contexto.", fill="E7F0FA")
    h(doc, "11.4 Escalamiento", 2)
    p(doc, "Si no existe respuesta dentro del tiempo definido para la severidad, ManeComb propone escalar al responsable siguiente. El sistema no debe marcar la alerta como resuelta por el simple hecho de enviar un mensaje. La resolución requiere estado verificable, responsable y nota breve.")
    next_page(doc)


def page_14_prototype_a11y(doc):
    h(doc, "12. Prototipo integrado y accesibilidad", 1)
    callout(doc, "Vista de operación", "MAPA · 18 unidades visibles · Actualizado hace 12 s\n\nUnidad C-24 · Sin datos recientes · hace 4 min\nRuta Centro-Terminal · Alerta AL-208 · ALTA\n[Ver seguimiento] [Atender] [Abrir chat] [Radio PTT]\n\nChat C-24 · Conexión inestable\nOperaciones 10:42 · Entregado: Confirma tu ubicación.\nConductor 10:44 · Pendiente: Estoy cerca de Terminal.", fill="EAF0F6")
    h(doc, "12.1 Comportamiento", 2)
    p(doc, "La selección de C-24 permanece al abrir seguimiento, alerta o conversación. Atender asigna al operador sin cerrar la tarjeta. El mensaje pendiente conserva su texto y se reintenta con el mismo identificador. Radio PTT solicita permiso solo si falta y vuelve al mismo contexto al terminar. La interfaz utiliza carga progresiva para evitar que el mapa completo se bloquee cuando falla un panel secundario.")
    h(doc, "12.2 Aplicación de WCAG 2.2", 2)
    table(doc, ["Criterio", "Aplicación en ManeComb"], [
        ["Perceptible", "Severidad y estados combinan texto, icono y contraste; el mapa incluye lista equivalente."],
        ["Operable", "Filtros, marcadores, mensajes y acciones funcionan con teclado y foco visible."],
        ["Comprensible", "Los mismos estados conservan nombre y comportamiento en todos los módulos."],
        ["Robusto", "Regiones vivas anuncian alertas y mensajes sin mover el foco inesperadamente."],
        ["Objetivo táctil", "Acciones móviles mantienen área suficiente y separación para evitar activaciones accidentales."],
        ["Autenticación accesible", "La recuperación no depende solo de memoria, acertijos o transcripción compleja."],
    ], [2300, 7060], size=9)
    h(doc, "12.3 Movimiento, sonido y tiempo", 2)
    p(doc, "El parpadeo no se utiliza para urgencia. Las animaciones pueden reducirse; el sonido y la vibración respetan preferencias salvo políticas justificadas para alertas críticas. Los límites de tiempo se informan y permiten extensión cuando no comprometen seguridad. Los mensajes nuevos se anuncian con remitente y conversación, pero no interrumpen la lectura de contenido actual.")
    next_page(doc)


def page_15_validation_roadmap(doc):
    h(doc, "13. Validación y evolución futura", 1)
    h(doc, "13.1 Plan de evaluación", 2)
    table(doc, ["Prueba", "Participantes", "Indicador de éxito"], [
        ["Localizar unidad", "Operadores nuevos y frecuentes", "Tiempo, errores y comprensión de vigencia."],
        ["Atender alerta", "Operadores y supervisores", "Priorización, asignación y cierre correcto."],
        ["Enviar con red inestable", "Conductores y operadores", "Ausencia de duplicados y comprensión de estados."],
        ["Usar Radio PTT", "Conductores", "Comprensión del gesto, canal y confirmación."],
        ["Configurar perfil", "Todos los roles", "Comprensión de permisos, seguridad y preferencias."],
        ["Accesibilidad", "Usuarios de teclado y tecnologías de apoyo", "Finalización sin barreras críticas."],
    ], [2400, 3100, 3860], size=9)
    h(doc, "13.2 Actualización futura: Documentos", 2)
    p(doc, "Documentos se considera una capacidad nueva y no una pantalla consolidada. Por esa razón, este informe no fija un flujo de carga, formatos, límites, estados de aprobación ni diseño visual definitivo. Inventar esos detalles produciría documentación que podría contradecir la implementación posterior. Antes de desarrollar ayuda se requiere definir usuarios, tipos documentales, responsables de revisión, seguridad, retención, vencimientos, notificaciones y relación con perfiles.")
    h(doc, "13.3 Trabajo previo a Documentos", 3)
    bullets(doc, [
        "Investigar quién carga, consulta, aprueba, rechaza y reemplaza cada documento.",
        "Diseñar prototipos de baja fidelidad y probarlos con conductores y administración.",
        "Definir reglas del backend, auditoría, permisos, cifrado, retención y eliminación.",
        "Acordar estados y mensajes antes de escribir manuales o preguntas frecuentes.",
        "Integrar vencimientos con alertas y permitir conversación contextual sin exponer archivos indebidamente.",
        "Validar carga móvil bajo conectividad limitada y accesibilidad de selección, progreso y errores.",
    ])
    h(doc, "13.4 Futuras mejoras generales", 2)
    p(doc, "La evolución de ManeComb puede incluir métricas de calidad de alertas, búsqueda unificada, resúmenes de conversación, reglas configurables por operación, integración de correo para comunicaciones administrativas y llamadas como escalamiento controlado. Cualquier automatización debe explicar su origen, permitir corrección humana y conservar privacidad. El correo no reemplaza alertas críticas ni chat en tiempo real; sirve para resúmenes, invitaciones, recuperación y comunicaciones que toleran demora.")
    next_page(doc)


def page_16_conclusion_refs(doc):
    h(doc, "14. Conclusiones", 1)
    p(doc, "ManeComb debe entenderse como una red de decisiones conectadas. Seguimiento y mapa proporcionan conciencia situacional; alertas convierten eventos en prioridades; chat conserva coordinación y evidencia; Radio PTT acelera la comunicación; y perfil establece identidad, permisos y preferencias. La calidad de la experiencia depende de mantener el contexto entre estos módulos, no de perfeccionar cada pantalla por separado.")
    p(doc, "La propuesta coloca la asistencia dentro del flujo real. Los estados explican qué ocurre, las acciones indican cómo continuar y los mecanismos de recuperación protegen el trabajo ante fallas de conexión. La accesibilidad amplía esta continuidad mediante navegación por teclado, alternativas al mapa, anuncios controlados de contenido dinámico, contraste y lenguaje comprensible. También se reduce la fatiga de alertas al diferenciar severidad, agrupar duplicados y exigir trazabilidad de atención.")
    p(doc, "Documentos permanece correctamente identificado como actualización futura. Su documentación deberá escribirse después de validar pantallas, actores y reglas; mientras tanto, el diseño puede preparar puntos de integración con perfil, alertas y chat sin presentar supuestos como hechos. Esta separación permite que el informe represente con profundidad las funciones existentes y, al mismo tiempo, proponga una evolución responsable.")
    h(doc, "15. Referencias", 1)
    refs = [
        "International Organization for Standardization. (2019). ISO 9241-210:2019 Ergonomics of human-system interaction - Part 210: Human-centred design for interactive systems.",
        "Lidwell, W., Holden, K., & Butler, J. (2010). Universal principles of design. Rockport Publishers.",
        "Nielsen, J. (1994). Usability engineering. Morgan Kaufmann.",
        "Norman, D. A. (2013). The design of everyday things (Rev. ed.). Basic Books.",
        "Shneiderman, B., Plaisant, C., Cohen, M., Jacobs, S., Elmqvist, N., & Diakopoulos, N. (2016). Designing the user interface (6th ed.). Pearson.",
        "World Wide Web Consortium. (2023). Web Content Accessibility Guidelines (WCAG) 2.2. https://www.w3.org/TR/WCAG22/",
    ]
    for ref in refs:
        para = p(doc, ref)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
    h(doc, "Cierre", 2)
    callout(doc, "Resultado esperado", "Una operación que permita comprender qué ocurre, dónde ocurre, a quién afecta, quién responde y cuál es la siguiente acción, incluso cuando la conectividad o el contexto cambian.")


def build():
    doc = Document()
    base.configure(doc)
    page_1_cover(doc)
    page_2_index(doc)
    page_3_intro(doc)
    page_4_system(doc)
    page_5_users(doc)
    page_6_tracking_map(doc)
    page_7_alerts(doc)
    page_8_chat(doc)
    page_9_ptt(doc)
    page_10_profile(doc)
    page_11_assistance(doc)
    page_12_guide_tracking(doc)
    page_13_guide_comms(doc)
    page_14_prototype_a11y(doc)
    page_15_validation_roadmap(doc)
    page_16_conclusion_refs(doc)
    doc.core_properties.title = "Informe integral de usabilidad y asistencia técnica de ManeComb"
    doc.core_properties.subject = "Seguimiento, mapa, alertas, chat, Radio PTT, perfil y evolución funcional"
    doc.core_properties.author = "ManeComb"
    doc.save(OUT)
    base.update_fields(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
