from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "reporte-integral-desarrollo-manecomb-app-ventas-backend-frontend.docx"
OUTPUT = ROOT / "docs" / "reporte-integral-desarrollo-manecomb-estilo-rojo.docx"

RED = RGBColor(192, 0, 0)
DARK_RED = RGBColor(139, 0, 0)
RED_HEX = "C00000"
DARK_RED_HEX = "8B0000"
LIGHT_RED_HEX = "FDECEC"
ALT_RED_HEX = "FFF5F5"

ORIGINAL_ACCENTS = {
    "1F374E",  # azul marino del reporte
    "2E74B5",  # azul de títulos y subtítulos
    "1F4D78",  # azul oscuro de tercer nivel
    "1F3750",
}


def set_fill(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def run_color_hex(run):
    color = run.font.color.rgb
    return str(color).upper() if color is not None else None


def recolor_runs(paragraphs):
    for paragraph in paragraphs:
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        for run in paragraph.runs:
            current = run_color_hex(run)
            if is_heading:
                run.font.color.rgb = RED if paragraph.style.name != "Heading 1" else DARK_RED
            elif current in ORIGINAL_ACCENTS:
                run.font.color.rgb = DARK_RED


def recolor_document():
    document = Document(SOURCE)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.color.rgb = DARK_RED if style_name == "Heading 1" else RED

    recolor_runs(document.paragraphs)
    for section in document.sections:
        recolor_runs(section.header.paragraphs)
        recolor_runs(section.footer.paragraphs)
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = DARK_RED

    for table in document.tables:
        first_row_properties = table.rows[0]._tr.get_or_add_trPr()
        is_data_table = first_row_properties.find(qn("w:tblHeader")) is not None

        if is_data_table:
            for cell in table.rows[0].cells:
                set_fill(cell, DARK_RED_HEX)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
            for row_index, row in enumerate(table.rows[1:], start=1):
                for cell in row.cells:
                    set_fill(cell, ALT_RED_HEX if row_index % 2 == 0 else "FFFFFF")
        else:
            for row in table.rows:
                for cell in row.cells:
                    set_fill(cell, LIGHT_RED_HEX)

        for row in table.rows:
            for cell in row.cells:
                recolor_runs(cell.paragraphs)

    document.core_properties.title = "Reporte integral del desarrollo de ManeComb - edición roja"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    recolor_document()
